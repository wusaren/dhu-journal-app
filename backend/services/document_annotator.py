"""
文档批注功能模块
用于为论文格式检测结果添加批注
"""

import os
import shutil
import re
import logging
from docx import Document
from typing import Dict, Any, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)

# 检测模块执行顺序
DETECTION_ORDER = ['Title', 'Abstract', 'Keywords', 'Content', 'Formula', 'Figure', 'Table']


def create_document_copy(original_path: str, output_dir: str) -> Optional[str]:
    """
    创建文档副本
    
    参数：
        original_path: 原始文档路径
        output_dir: 输出目录
    
    返回：
        副本文件路径
    """
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        # annotate_filename = f"{timestamp}_annotated.docx"
        base_name = os.path.splitext(os.path.basename(original_path))[0].split('_')[-1]
        copy_filename = f"{timestamp}_{base_name}_annotated.docx"
        copy_path = os.path.join(output_dir, copy_filename)
        
        shutil.copy(original_path, copy_path)
        logger.info(f"文档副本已创建: {copy_path}")
        return copy_path
    except Exception as e:
        logger.error(f"创建文档副本失败: {e}")
        return None


def find_paragraph_by_keyword(doc: Document, keyword: str, case_sensitive: bool = False) -> Optional[Any]:
    """
    通过关键字查找段落
    """
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
        
        text = paragraph.text if case_sensitive else paragraph.text.lower()
        search_key = keyword if case_sensitive else keyword.lower()
        
        if search_key in text:
            return paragraph
    
    return None


def find_paragraph_by_index(doc: Document, index: int) -> Optional[Any]:
    """
    通过索引查找段落
    """
    try:
        if 0 <= index < len(doc.paragraphs):
            return doc.paragraphs[index]
    except Exception:
        pass
    
    return None


def find_paragraph_by_text(doc: Document, text_fragment: str, threshold: float = 0.7) -> Optional[Any]:
    """
    通过文本片段查找段落（模糊匹配）
    """
    best_match = None
    best_score = 0
    
    search_text = text_fragment.lower().strip()
    if len(search_text) < 5:
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text.lower():
                return paragraph
        return None
    
    for paragraph in doc.paragraphs:
        if not paragraph.text or len(paragraph.text.strip()) < 5:
            continue
        
        para_text = paragraph.text.lower().strip()
        
        if search_text in para_text:
            score = len(search_text) / len(para_text)
            if score > best_score:
                best_score = score
                best_match = paragraph
    
    if best_score >= threshold:
        return best_match
    
    return None


def find_content_paragraph_by_number(doc: Document, para_number: int, hierarchy_report: Dict) -> Optional[Any]:
    """
    根据段落编号定位具体的正文段落
    """
    title_paragraph_indices = set()
    introduction_index = None
    
    if hierarchy_report and hierarchy_report.get('titles'):
        for title_info in hierarchy_report['titles']:
            if 'paragraph_index' in title_info:
                idx = title_info['paragraph_index']
                title_paragraph_indices.add(idx)
                if title_info.get('level') == 0 and title_info.get('text') == 'Introduction':
                    introduction_index = idx
    
    if introduction_index is None:
        for i, para in enumerate(doc.paragraphs):
            if 'Introduction' in para.text:
                introduction_index = i
                break
    
    if introduction_index is None:
        return None
    
    content_paragraph_count = 0
    
    for i in range(introduction_index + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        
        if i in title_paragraph_indices:
            continue
        
        text = paragraph.text.strip() if paragraph.text else ""
        if not text or len(text) <= 20:
            continue
        
        text_lower = text.lower()
        if (text_lower.startswith('table ') or 
            text_lower.startswith('figure ') or 
            text_lower.startswith('fig.') or
            text_lower.startswith('图 ') or 
            text_lower.startswith('表 ') or
            re.match(r'^(表|图|Table|Figure)\s*\d+', text, re.IGNORECASE)):
            continue
        
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        alignment = paragraph.alignment
        if alignment is not None and alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and len(text) < 80:
            continue
        
        content_paragraph_count += 1
        
        if content_paragraph_count == para_number:
            return paragraph
    
    return None


def add_comment_to_paragraph(doc: Document, paragraph: Any, comment_text: str, 
                            author: str = "论文检测系统", initials: str = "PDS") -> bool:
    """
    为段落添加批注
    """
    try:
        if not paragraph.runs:
            paragraph.add_run("")
        
        runs_to_annotate = paragraph.runs if len(paragraph.runs) > 1 else paragraph.runs[0]
        
        doc.add_comment(runs=runs_to_annotate, text=comment_text, author=author, initials=initials)
        return True
    except Exception as e:
        try:
            run = paragraph.runs[0]
            doc.add_comment(runs=run, text=comment_text, author=author, initials=initials)
            logger.warning(f"批注仅覆盖部分文本（备用方案）")
            return True
        except Exception as e2:
            logger.error(f"添加批注失败: {e2}")
            return False


def extract_paragraph_number_from_message(message: str) -> Optional[int]:
    """
    从错误消息中提取段落编号
    """
    match = re.search(r'正文段落\s*(\d+)', message)
    if match:
        return int(match.group(1))
    
    match = re.search(r'段落\s*(\d+)', message)
    if match:
        return int(match.group(1))
    
    return None


def group_messages_by_paragraph(messages: List[str]) -> Dict[Optional[int], List[str]]:
    """
    将消息按段落分组
    """
    grouped = {}
    
    for msg in messages:
        para_num = extract_paragraph_number_from_message(msg)
        if para_num not in grouped:
            grouped[para_num] = []
        grouped[para_num].append(msg)
    
    return grouped


def group_messages_by_title(messages: List[str], titles: List[Dict]) -> Dict[Optional[str], List[str]]:
    """
    将标题格式/大小写消息按标题文本分组
    """
    grouped = {}
    
    for msg in messages:
        title_match = re.search(r"标题\s+'([^']+)'", msg)
        if title_match:
            title_text = title_match.group(1)
            if title_text not in grouped:
                grouped[title_text] = []
            grouped[title_text].append(msg)
        else:
            if None not in grouped:
                grouped[None] = []
            grouped[None].append(msg)
    
    if None in grouped:
        del grouped[None]
    
    return grouped


def parse_issues_from_reports(all_reports: Dict[str, Any]) -> List[Dict]:
    """
    从报告中提取问题列表和定位信息
    
    返回：
        问题列表，每个问题包含模块名、检测项、消息、定位方法和定位数据
    """
    issues = []
    
    for module_name in DETECTION_ORDER:
        if module_name not in all_reports:
            continue
        
        report = all_reports[module_name]
        
        if report.get('error', False):
            continue
        
        # Title模块处理
        if module_name == 'Title':
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if not messages:
                        continue
                    
                    if 'title' in section_key.lower():
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'index',
                            'locate_data': 0
                        })
                    elif 'author' in section_key.lower():
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'index',
                            'locate_data': 1
                        })
                    elif 'affiliation' in section_key.lower():
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': 'College'
                        })
                    elif 'format' in section_key.lower():
                        # 按类型分组
                        title_msgs = []
                        author_msgs = []
                        affiliation_msgs = []
                        current_group = None
                        
                        for msg in messages:
                            msg_lower = msg.lower()
                            msg_stripped = msg.strip()
                            is_header = not msg_stripped.startswith('- ')
                            
                            if is_header:
                                if '标题格式' in msg or 'title' in msg_lower:
                                    current_group = 'title'
                                    title_msgs.append(msg)
                                elif '作者格式' in msg or 'author' in msg_lower:
                                    current_group = 'author'
                                    author_msgs.append(msg)
                                elif '单位格式' in msg or '单位段落' in msg or 'affiliation' in msg_lower:
                                    current_group = 'affiliation'
                                    affiliation_msgs.append(msg)
                                else:
                                    current_group = 'title'
                                    title_msgs.append(msg)
                            else:
                                if current_group == 'title':
                                    title_msgs.append(msg)
                                elif current_group == 'author':
                                    author_msgs.append(msg)
                                elif current_group == 'affiliation':
                                    affiliation_msgs.append(msg)
                                else:
                                    title_msgs.append(msg)
                        
                        if title_msgs:
                            issues.append({
                                'module': module_name,
                                'section': section_key + '_title',
                                'messages': title_msgs,
                                'locate_method': 'index',
                                'locate_data': 0
                            })
                        
                        if author_msgs:
                            issues.append({
                                'module': module_name,
                                'section': section_key + '_authors',
                                'messages': author_msgs,
                                'locate_method': 'index',
                                'locate_data': 1
                            })
                        
                        if affiliation_msgs:
                            affiliation_by_para = {}
                            general_affiliation_msgs = []
                            current_para_idx = None
                            
                            for msg in affiliation_msgs:
                                msg_stripped = msg.strip()
                                is_header = not msg_stripped.startswith('- ')
                                
                                if is_header:
                                    para_match = re.search(r'第(\d+)段', msg)
                                    if para_match:
                                        current_para_idx = int(para_match.group(1)) - 1
                                        if current_para_idx not in affiliation_by_para:
                                            affiliation_by_para[current_para_idx] = []
                                        affiliation_by_para[current_para_idx].append(msg)
                                    else:
                                        current_para_idx = None
                                        general_affiliation_msgs.append(msg)
                                else:
                                    if current_para_idx is not None:
                                        affiliation_by_para[current_para_idx].append(msg)
                                    else:
                                        general_affiliation_msgs.append(msg)
                            
                            for para_idx, msgs in affiliation_by_para.items():
                                issues.append({
                                    'module': module_name,
                                    'section': section_key + f'_affiliation_para{para_idx+1}',
                                    'messages': msgs,
                                    'locate_method': 'index',
                                    'locate_data': para_idx
                                })
                            
                            if general_affiliation_msgs:
                                issues.append({
                                    'module': module_name,
                                    'section': section_key + '_affiliations',
                                    'messages': general_affiliation_msgs,
                                    'locate_method': 'keyword',
                                    'locate_data': 'College'
                                })
                    else:
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'index',
                            'locate_data': 0
                        })
        
        # Abstract模块处理
        elif module_name == 'Abstract':
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if messages:
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': 'Abstract:'
                        })
        
        # Keywords模块处理
        elif module_name == 'Keywords':
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if messages:
                        if 'clc' in section_key.lower():
                            locate_data = 'CLC number'
                        elif 'footnote' in section_key.lower():
                            locate_data = 'CLC number'
                        else:
                            locate_data = 'Keywords:'
                        
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': locate_data
                        })
        
        # Content模块处理
        elif module_name == 'Content':
            hierarchy_report = report.get('hierarchy', {})
            titles = report.get('titles', [])
            
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details', 'hierarchy', 'titles']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if not messages:
                        continue
                    
                    if section_key == 'content_format':
                        grouped_messages = group_messages_by_paragraph(messages)
                        
                        for para_num, para_messages in grouped_messages.items():
                            if para_num is not None:
                                issues.append({
                                    'module': module_name,
                                    'section': f'{section_key}_para{para_num}',
                                    'messages': para_messages,
                                    'locate_method': 'content_paragraph',
                                    'locate_data': para_num,
                                    'extra': {
                                        'hierarchy_report': hierarchy_report
                                    }
                                })
                    
                    elif section_key in ['format', 'case']:
                        grouped_title_messages = group_messages_by_title(messages, titles)
                        
                        for title_text, title_messages in grouped_title_messages.items():
                            if title_text and title_messages:
                                title_info = next((t for t in titles if t.get('text') == title_text), None)
                                if title_info and 'paragraph_index' in title_info:
                                    issues.append({
                                        'module': module_name,
                                        'section': section_key,
                                        'messages': title_messages,
                                        'locate_method': 'index',
                                        'locate_data': title_info.get('paragraph_index')
                                    })
                                elif title_text:
                                    full_title_text = title_info.get('full_text', title_text) if title_info else title_text
                                    issues.append({
                                        'module': module_name,
                                        'section': section_key,
                                        'messages': title_messages,
                                        'locate_method': 'text',
                                        'locate_data': full_title_text
                                    })
                    else:
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': 'Introduction'
                        })
        
        # Formula模块处理
        elif module_name == 'Formula':
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if not messages:
                        continue
                    
                    details = report.get('details', {})
                    formula_paragraphs = details.get('formula_paragraphs', [])
                    
                    if formula_paragraphs:
                        first_formula = formula_paragraphs[0]
                        text_preview = first_formula.get('text_preview', '')
                        
                        number_match = re.search(r'\((\d+)\)\s*$', text_preview)
                        
                        if number_match:
                            formula_number = number_match.group(0).strip()
                            issues.append({
                                'module': module_name,
                                'section': section_key,
                                'messages': messages,
                                'locate_method': 'formula_number',
                                'locate_data': formula_number,
                                'extra': {
                                    'text_preview': text_preview
                                }
                            })
                        else:
                            locate_text = text_preview[-20:].strip() if len(text_preview) > 20 else text_preview
                            
                            if locate_text and len(locate_text) >= 3:
                                issues.append({
                                    'module': module_name,
                                    'section': section_key,
                                    'messages': messages,
                                    'locate_method': 'text',
                                    'locate_data': locate_text
                                })
        
        # Table模块处理
        elif module_name == 'Table':
            numbering_report = report.get('numbering', {})
            if isinstance(numbering_report, dict) and not numbering_report.get('ok', False):
                messages = numbering_report.get('messages', [])
                if messages:
                    details = report.get('details', {})
                    caption_info_list = details.get('caption_info', [])
                    if caption_info_list:
                        first_caption = caption_info_list[0]
                        caption_text = first_caption.get('text', 'Table')
                        issues.append({
                            'module': module_name,
                            'section': 'numbering',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]
                        })
            
            tables_report = report.get('tables', [])
            for i, table_report in enumerate(tables_report):
                caption_info = table_report.get('caption', {})
                caption_text = caption_info.get('text', f'Table {i+1}')
                
                for check_key in ['caption_format', 'table_style', 'table_alignment']:
                    check_result = table_report.get(check_key, {})
                    if isinstance(check_result, dict) and not check_result.get('ok', False):
                        messages = check_result.get('messages', [])
                        if messages:
                            issues.append({
                                'module': module_name,
                                'section': f'table{i+1}_{check_key}',
                                'messages': messages,
                                'locate_method': 'keyword',
                                'locate_data': caption_text[:20]
                            })
        
        # Figure模块处理
        elif module_name == 'Figure':
            numbering_report = report.get('numbering', {})
            if isinstance(numbering_report, dict) and not numbering_report.get('ok', False):
                messages = numbering_report.get('messages', [])
                if messages:
                    captions = report.get('captions', [])
                    if captions:
                        first_caption = captions[0]
                        caption_text = first_caption.get('full_text', 'Fig.')
                        issues.append({
                            'module': module_name,
                            'section': 'numbering',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]
                        })
            
            figures = report.get('figures', [])
            for fig_report in figures:
                fig_idx = fig_report.get('figure_index', 0)
                para_idx = fig_report.get('paragraph_index', 0)
                
                has_caption = fig_report.get('has_caption', False)
                caption_info = fig_report.get('caption_info') if has_caption else None
                
                caption_messages = []
                format_check = fig_report.get('format_check', {})
                if isinstance(format_check, dict) and not format_check.get('ok', False):
                    caption_messages.extend(format_check.get('messages', []))
                
                if caption_messages and has_caption:
                    caption_text = caption_info.get('full_text', f'Fig.{caption_info.get("number", fig_idx)}')[:30]
                    issues.append({
                        'module': module_name,
                        'section': f'figure{fig_idx}_caption',
                        'messages': caption_messages,
                        'locate_method': 'keyword',
                        'locate_data': caption_text
                    })
                
                picture_messages = []
                
                picture_check = fig_report.get('picture_check', {})
                if isinstance(picture_check, dict) and not picture_check.get('ok', False):
                    picture_messages.extend(picture_check.get('messages', []))
                
                content_check = fig_report.get('content_check', {})
                if isinstance(content_check, dict) and not content_check.get('ok', False):
                    if content_check.get('is_chart', False):
                        picture_messages.extend(content_check.get('messages', []))
                
                if picture_messages:
                    issues.append({
                        'module': module_name,
                        'section': f'figure{fig_idx}_picture',
                        'messages': picture_messages,
                        'locate_method': 'index',
                        'locate_data': para_idx
                    })
    
    return issues


def add_all_comments(doc_path: str, copy_path: str, issues_list: List[Dict]) -> int:
    """
    在文档副本上添加所有批注
    
    返回：
        成功添加的批注数量
    """
    try:
        doc = Document(copy_path)
        comment_count = 0
        
        logger.info(f"正在添加批注，共有 {len(issues_list)} 个问题")
        
        for issue in issues_list:
            module_name = issue['module']
            section_name = issue['section']
            messages = issue['messages']
            locate_method = issue['locate_method']
            locate_data = issue['locate_data']
            extra = issue.get('extra', {})
            
            paragraph = None
            if locate_method == 'keyword' and locate_data:
                paragraph = find_paragraph_by_keyword(doc, locate_data)
            elif locate_method == 'index':
                paragraph = find_paragraph_by_index(doc, locate_data)
            elif locate_method == 'text':
                paragraph = find_paragraph_by_text(doc, locate_data)
            elif locate_method == 'content_paragraph':
                hierarchy_report = extra.get('hierarchy_report', {})
                paragraph = find_content_paragraph_by_number(doc, locate_data, hierarchy_report)
            elif locate_method == 'formula_number':
                paragraph = find_paragraph_by_keyword(doc, locate_data)
            
            if paragraph:
                comment_text = f"[{module_name}-{section_name}]\n"
                for msg in messages:
                    comment_text += f"• {msg}\n"
                
                if add_comment_to_paragraph(doc, paragraph, comment_text.strip()):
                    comment_count += 1
                    logger.info(f"已添加批注: {module_name}-{section_name}")
            else:
                logger.warning(f"无法定位段落: {module_name}-{section_name} (方法:{locate_method}, 数据:{locate_data})")
        
        doc.save(copy_path)
        logger.info(f"成功添加 {comment_count} 个批注")
        
        return comment_count
    
    except Exception as e:
        logger.error(f"添加批注过程出错: {e}")
        return 0


def generate_annotated_document(docx_path: str, all_reports: Dict[str, Any], output_dir: str) -> Optional[str]:
    """
    生成带批注的文档
    
    参数：
        docx_path: 原始文档路径
        all_reports: 所有检测报告
        output_dir: 输出目录
    
    返回：
        批注文档的路径，失败返回None
    """
    try:
        # 创建文档副本
        copy_path = create_document_copy(docx_path, output_dir)
        if not copy_path:
            return None
        
        # 从报告中提取问题
        issues_list = parse_issues_from_reports(all_reports)
        logger.info(f"共识别出 {len(issues_list)} 个问题")
        
        # 添加批注
        if issues_list:
            comment_count = add_all_comments(docx_path, copy_path, issues_list)
            if comment_count > 0:
                logger.info(f"批注添加完成！共添加 {comment_count} 个批注")
                return copy_path
            else:
                logger.warning("未能添加任何批注")
                return copy_path
        else:
            logger.info("未发现问题，无需添加批注")
            return copy_path
    
    except Exception as e:
        logger.error(f"生成批注文档失败: {e}")
        return None

