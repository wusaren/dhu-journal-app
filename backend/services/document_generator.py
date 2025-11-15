import os
import re
import json
from datetime import datetime
from pathlib import Path
from typing import List, Any, Optional, Dict
import logging

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook

logger = logging.getLogger(__name__)

# 默认配置文件路径
DEFAULT_CONFIG_FILE = Path(__file__).parent.parent / 'configs' / 'default_stats_columns.json'
DEFAULT_TUIWEN_CONFIG_FILE = Path(__file__).parent.parent / 'configs' / 'default_tuiwen_fields.json'

def load_default_columns_config() -> List[Dict]:
    """
    从JSON文件加载默认列配置
    如果文件不存在或加载失败，返回硬编码的默认配置作为后备
    
    Returns:
        列配置列表，格式: [{'key': 'manuscript_id', 'order': 1}, ...]
    """
    try:
        if DEFAULT_CONFIG_FILE.exists():
            with open(DEFAULT_CONFIG_FILE, 'r', encoding='utf-8') as f:
                config_data = json.load(f)
                columns = config_data.get('columns', [])
                if columns:
                    logger.info(f"从配置文件加载了默认列配置: {len(columns)} 列")
                    return columns
        
        logger.warning(f"默认配置文件不存在或为空，使用硬编码默认配置: {DEFAULT_CONFIG_FILE}")
    except Exception as e:
        logger.error(f"加载默认列配置失败: {str(e)}，使用硬编码默认配置")
    
    # 后备硬编码配置
    return [
        {'key': 'manuscript_id', 'order': 1},
        {'key': 'pdf_pages', 'order': 2},
        {'key': 'first_author', 'order': 3},
        {'key': 'corresponding', 'order': 4},
        {'key': 'issue', 'order': 5},
        {'key': 'is_dhu', 'order': 6},
    ]

def load_default_tuiwen_fields_config() -> List[Dict]:
    """
    从JSON文件加载默认推文字段配置
    如果文件不存在或加载失败，返回硬编码的默认配置作为后备
    
    Returns:
        字段配置列表，格式: [{'field': 'title', 'label': '标题', 'order': 1}, ...]
    """
    try:
        if DEFAULT_TUIWEN_CONFIG_FILE.exists():
            with open(DEFAULT_TUIWEN_CONFIG_FILE, 'r', encoding='utf-8') as f:
                config_data = json.load(f)
                fields = config_data.get('fields', [])
                if fields:
                    logger.info(f"从配置文件加载了默认推文字段配置: {len(fields)} 个字段")
                    return fields
        
        logger.warning(f"默认推文配置文件不存在或为空，使用硬编码默认配置: {DEFAULT_TUIWEN_CONFIG_FILE}")
    except Exception as e:
        logger.error(f"加载默认推文字段配置失败: {str(e)}，使用硬编码默认配置")
    
    # 后备硬编码配置
    return [
        {'field': 'chinese_title', 'label': '中文标题', 'required': False, 'order': 1},
        {'field': 'chinese_authors', 'label': '中文作者', 'required': False, 'order': 2},
        {'field': 'second_image', 'label': '论文配图', 'required': False, 'order': 3},
        {'field': 'title', 'label': '标题', 'required': True, 'order': 4},
        {'field': 'authors', 'label': '作者', 'required': True, 'order': 5},
        {'field': 'doi', 'label': 'DOI', 'required': False, 'order': 6},
        {'field': 'citation', 'label': '引用信息', 'required': True, 'order': 7},
        {'field': 'first_image', 'label': '作者说链接/OSID', 'required': False, 'order': 8},
    ]

def _get(a: Any, key: str):
    """获取对象属性，兼容dict和object - 完全照搬参考代码"""
    if isinstance(a, dict): 
        return a.get(key)
    return getattr(a, key)

def format_authors_for_citation(authors: str, max_authors: int = 3) -> str:
    """
    格式化作者信息用于citation
    规则：
    1. 最多显示3个作者
    2. 作者的名只取首字母（如 "HUANG Jiacui" -> "HUANG J"）
    3. 如果超过3个作者，在第三个作者后加上 ", et al."
    
    示例：
    - "HUANG Jiacui, ZHAO Mingbo, ZHANG Hongtao, WANG Li" 
      -> "HUANG J, ZHAO M, ZHANG H, et al."
    - "HUANG Jiacui, ZHAO Mingbo, ZHANG Hongtao" 
      -> "HUANG J, ZHAO M, ZHANG H"
    """
    if not authors:
        return ""
    
    # 分割作者（支持逗号或中文逗号分隔）
    author_list = [a.strip() for a in re.split(r'[,，]', authors) if a.strip()]
    
    if not author_list:
        return ""
    
    # 格式化每个作者：名字只取首字母
    formatted_authors = []
    for author in author_list[:max_authors]:
        # 分割姓和名（通常格式：姓 名，如 "HUANG Jiacui"）
        parts = author.split()
        if len(parts) >= 2:
            # 有姓和名：姓 + 名的首字母
            last_name = parts[0]
            first_name = parts[1]
            formatted_author = f"{last_name} {first_name[0] if first_name else ''}"
            formatted_authors.append(formatted_author.strip())
        else:
            # 只有一个部分，可能是只有姓或者格式不对，保持原样
            formatted_authors.append(author)
    
    # 如果超过max_authors个作者，添加 ", et al."
    if len(author_list) > max_authors:
        formatted_authors.append("et al.")
    
    # 用逗号和空格连接
    return ", ".join(formatted_authors)

def generate_toc_docx(papers: List[Any], journal: Any) -> str:
    """
    生成目录Word文档 - 导出期刊内所有论文
    """
    try:
        # 创建Word文档
        doc = Document()
        
        # 设置样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(11)
        
        # 添加期刊标题
        title_para = doc.add_paragraph()
        title_run = title_para.runs[0] if title_para.runs else title_para.add_run()
        title_run.text = f"{journal.title} - {journal.issue}"
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # 添加空行
        doc.add_paragraph()
        
        # 按页码排序
        items = sorted([a for a in papers if _get(a, 'page_start') is not None], 
                      key=lambda x: _get(x, 'page_start'))
        
        # 添加目录内容
        for i, a in enumerate(items, 1):
            # 论文标题和页码
            doc.add_paragraph(f"{_get(a,'page_start')} {_get(a,'title') or ''}")
            # 作者信息
            doc.add_paragraph(_get(a, 'authors') or '')
            # 空行分隔
            if i < len(items):
                doc.add_paragraph()
        
        # 保存文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"目录_{journal.issue}_{timestamp}.docx"
        output_path = os.path.join('uploads', filename)
        
        # 确保目录存在
        os.makedirs('uploads', exist_ok=True)
        
        doc.save(output_path)
        logger.info(f"目录文档已生成: {output_path}")
        
        return output_path
        
    except Exception as e:
        logger.error(f"生成目录文档失败: {str(e)}")
        raise Exception(f"生成目录文档失败: {str(e)}")

def generate_excel_stats(articles, journal, columns_config=None) -> str:
    """
    生成统计表Excel - 导出期刊内所有论文的统计信息
    支持自定义列配置
    
    Args:
        articles: 论文列表
        journal: 期刊对象
        columns_config: 列配置列表，格式: [{'key': 'manuscript_id', 'order': 1}, ...]
                       如果为None，使用默认配置
    """
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = '期刊统计表'
        
        # 如果没有提供列配置，从JSON文件加载默认配置
        if not columns_config:
            columns_config = load_default_columns_config()
        
        # 按order排序
        columns_config = sorted(columns_config, key=lambda x: x.get('order', 999))
        
        # 定义所有可用列的映射（key -> 显示标签和数据获取函数）
        column_definitions = {
            'manuscript_id': {
                'label': '稿件号',
                'get_value': lambda a: _get(a, 'manuscript_id') or ''
            },
            'pdf_pages': {
                'label': '页数',
                'get_value': lambda a: _get(a, 'pdf_pages') or 0
            },
            'first_author': {
                'label': '一作',
                'get_value': lambda a: _get(a, 'first_author') or ''
            },
            'corresponding': {
                'label': '通讯',
                'get_value': lambda a: _get(a, 'corresponding') or ''
            },
            'authors': {
                'label': '作者',
                'get_value': lambda a: _get(a, 'authors') or ''
            },
            'issue': {
                'label': '刊期',
                'get_value': lambda a: _get(a, 'issue') or ''
            },
            'is_dhu': {
                'label': '是否东华大学',
                'get_value': lambda a: '是' if _get(a, 'is_dhu') else '否'
            },
            'title': {
                'label': '标题',
                'get_value': lambda a: _get(a, 'title') or ''
            },
            'chinese_title': {
                'label': '中文标题',
                'get_value': lambda a: _get(a, 'chinese_title') or ''
            },
            'chinese_authors': {
                'label': '中文作者',
                'get_value': lambda a: _get(a, 'chinese_authors') or ''
            },
            'doi': {
                'label': 'DOI',
                'get_value': lambda a: _get(a, 'doi') or ''
            },
            'page_start': {
                'label': '起始页码',
                'get_value': lambda a: _get(a, 'page_start') or ''
            },
            'page_end': {
                'label': '结束页码',
                'get_value': lambda a: _get(a, 'page_end') or ''
            },
        }
        
        # 生成表头（只包含配置中存在的列）
        valid_columns = []
        for col_config in columns_config:
            key = col_config.get('key')
            if key in column_definitions:
                valid_columns.append({
                    'key': key,
                    'label': column_definitions[key]['label'],
                    'get_value': column_definitions[key]['get_value']
                })
        
        headers = [col['label'] for col in valid_columns]
        ws.append(headers)
        
        # 添加所有论文数据
        r = 2  # 从第2行开始（第1行是表头）
        for a in articles:
            col_idx = 1
            for col in valid_columns:
                value = col['get_value'](a)
                ws.cell(row=r, column=col_idx, value=value)
                col_idx += 1
            r += 1

        # 保存文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"统计表_{journal.issue}_{timestamp}.xlsx"
        output_path = os.path.join('uploads', filename)
        
        os.makedirs('uploads', exist_ok=True)
        wb.save(output_path)
        
        logger.info(f"统计表已生成: {output_path}")
        logger.info(f"使用的列配置: {[col['key'] for col in valid_columns]}")
        
        # 记录生成的数据用于调试
        debug_data = []
        for a in articles:
            row_data = {}
            for col in valid_columns:
                row_data[col['label']] = col['get_value'](a)
            debug_data.append(row_data)
        logger.info(f"生成的数据（前3条）: {debug_data[:3]}")
        
        return output_path
        
    except Exception as e:
        logger.error(f"生成统计表Excel失败: {str(e)}")
        raise Exception(f"生成统计表Excel失败: {str(e)}")

def generate_excel_stats_from_template(articles, journal, template_file_path: str, column_mapping: List[Dict]) -> str:
    """
    基于模板生成统计表Excel
    
    Args:
        articles: 论文列表
        journal: 期刊对象
        template_file_path: 模板文件路径
        column_mapping: 列映射配置，格式: [
            {'template_header': '稿件号', 'system_key': 'manuscript_id', 'order': 1, 'is_custom': False},
            ...
        ]
    """
    try:
        from services.template_service import TemplateService
        
        # 加载模板文件
        wb = load_workbook(template_file_path)
        ws = wb.active
        
        # 找到表头行（第一个有内容的行）
        template_service = TemplateService()
        headers = template_service.extract_headers_from_excel(template_file_path)
        
        if not headers:
            raise Exception("无法从模板文件中提取表头")
        
        # 找到表头行号
        header_row = None
        for row_idx in range(1, min(10, ws.max_row + 1)):
            row = ws[row_idx]
            has_content = False
            for cell in row:
                if cell.value and str(cell.value).strip():
                    has_content = True
                    break
            if has_content:
                header_row = row_idx
                break
        
        if header_row is None:
            raise Exception("无法找到表头行")
        
        # 按 order 排序 column_mapping
        column_mapping = sorted(column_mapping, key=lambda x: x.get('order', 999))

        # 需要保留的表头集合
        mapped_headers = set()
        for mapping in column_mapping:
            template_header = mapping.get('template_header')
            if template_header:
                mapped_headers.add(str(template_header).strip())

        # 删除模板中不在映射里的列，避免表头残留
        columns_to_remove = []
        for col_idx, cell in enumerate(ws[header_row], start=1):
            header_text = str(cell.value).strip() if cell.value else ''
            if header_text and header_text not in mapped_headers:
                columns_to_remove.append(col_idx)
        for col_idx in sorted(columns_to_remove, reverse=True):
            ws.delete_cols(col_idx)

        # 删除后重新建立表头映射，并记录最后一个有内容的列
        header_to_col = {}
        last_content_col = 0
        for col_idx, cell in enumerate(ws[header_row], start=1):
            header_text = str(cell.value).strip() if cell.value else ''
            if header_text:
                header_to_col[header_text] = col_idx
                last_content_col = col_idx  # 更新最后一个有内容的列

        # 保存原始表头样式（用于新列）
        source_header_cell = None
        for cell in ws[header_row]:
            if cell.value and str(cell.value).strip():
                source_header_cell = cell
                break
        
        # 检查 column_mapping 中是否有模板文件中不存在的字段，如果有则添加
        # 从最后一个有内容的列之后开始添加，避免中间有空列
        next_col = last_content_col + 1
        for mapping in column_mapping:
            template_header = mapping.get('template_header')
            if template_header and template_header not in header_to_col:
                # 在表头行添加新列
                cell = ws.cell(row=header_row, column=next_col)
                cell.value = template_header
                # 复制表头行的样式（如果有）
                if source_header_cell:
                    try:
                        if source_header_cell.font:
                            from openpyxl.styles import Font
                            cell.font = Font(
                                name=source_header_cell.font.name,
                                size=source_header_cell.font.size,
                                bold=source_header_cell.font.bold,
                                color=source_header_cell.font.color
                            )
                        if source_header_cell.fill:
                            from openpyxl.styles import PatternFill
                            cell.fill = PatternFill(
                                start_color=source_header_cell.fill.start_color,
                                end_color=source_header_cell.fill.end_color,
                                fill_type=source_header_cell.fill.fill_type
                            )
                        if source_header_cell.alignment:
                            from openpyxl.styles import Alignment
                            cell.alignment = Alignment(
                                horizontal=source_header_cell.alignment.horizontal,
                                vertical=source_header_cell.alignment.vertical
                            )
                        if source_header_cell.border:
                            from openpyxl.styles import Border
                            cell.border = source_header_cell.border
                    except Exception as e:
                        logger.warning(f"复制表头样式失败: {str(e)}")
                header_to_col[template_header] = next_col
                next_col += 1
                logger.info(f"添加新列到模板: {template_header} (列 {header_to_col[template_header]})")
        
        # 删除表头行之后的所有数据行（保留表头）
        data_start_row = header_row + 1
        if ws.max_row >= data_start_row:
            ws.delete_rows(data_start_row, ws.max_row - header_row)
        
        # 定义所有可用列的映射
        column_definitions = {
            'manuscript_id': lambda a: _get(a, 'manuscript_id') or '',
            'pdf_pages': lambda a: _get(a, 'pdf_pages') or 0,
            'first_author': lambda a: _get(a, 'first_author') or '',
            'corresponding': lambda a: _get(a, 'corresponding') or '',
            'authors': lambda a: _get(a, 'authors') or '',
            'issue': lambda a: _get(a, 'issue') or '',
            'is_dhu': lambda a: '是' if _get(a, 'is_dhu') else '否',
            'title': lambda a: _get(a, 'title') or '',
            'chinese_title': lambda a: _get(a, 'chinese_title') or '',
            'chinese_authors': lambda a: _get(a, 'chinese_authors') or '',
            'doi': lambda a: _get(a, 'doi') or '',
            'page_start': lambda a: _get(a, 'page_start') or '',
            'page_end': lambda a: _get(a, 'page_end') or '',
        }
        
        # 填充数据
        for article_idx, article in enumerate(articles, start=0):
            row_num = data_start_row + article_idx
            
            for mapping in column_mapping:
                template_header = mapping.get('template_header')
                system_key = mapping.get('system_key')
                col_idx = header_to_col.get(template_header)
                
                if col_idx:
                    cell = ws.cell(row=row_num, column=col_idx)
                    
                    if system_key and system_key in column_definitions:
                        value = column_definitions[system_key](article)
                        cell.value = value
                    elif mapping.get('is_custom', False):
                        cell.value = ''
                    # 数据行保持默认格式，不复制表头样式
        
        # 保存文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"统计表_{journal.issue}_{timestamp}.xlsx"
        output_path = os.path.join('uploads', filename)
        
        os.makedirs('uploads', exist_ok=True)
        wb.save(output_path)
        
        logger.info(f"基于模板生成统计表: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"基于模板生成统计表Excel失败: {str(e)}")
        raise Exception(f"基于模板生成统计表Excel失败: {str(e)}")

def get_local_image_path(image_path: str, temp_dir: str) -> Optional[str]:
    """
    获取本地图片路径，如果图片存在则返回路径，否则返回None
    返回: 本地图片路径，如果图片不存在则返回None
    """
    try:
        # 检查图片路径是否存在
        if image_path and os.path.exists(image_path):
            logger.info(f"本地图片存在: {image_path}")
            return image_path
        else:
            logger.warning(f"本地图片不存在: {image_path}")
            return None
            
    except Exception as e:
        logger.error(f"检查本地图片时出错: {str(e)}")
        return None

def generate_tuiwen_content(papers, journal):
    """
    生成秀米推文内容 - Word文档格式
    使用默认字段配置（从JSON文件加载或硬编码后备）
    """
    try:
        # 从JSON文件加载默认字段配置
        default_fields = load_default_tuiwen_fields_config()
        
        # 使用 generate_tuiwen_from_fields 函数生成推文
        return generate_tuiwen_from_fields(papers, journal, default_fields)
        
    except Exception as e:
        logger.error(f"使用默认配置生成推文失败: {str(e)}，尝试使用原始硬编码方式")
        # 如果使用配置方式失败，回退到原始硬编码方式
        return _generate_tuiwen_content_legacy(papers, journal)

def _generate_tuiwen_content_legacy(papers, journal):
    """生成秀米推文内容 - 原始硬编码方式（作为后备）"""
    try:
        # 创建临时目录用于存储下载的图片
        temp_dir = os.path.join('temp_images', datetime.now().strftime('%Y%m%d_%H%M%S'))
        os.makedirs(temp_dir, exist_ok=True)
        
        # 创建Word文档
        doc = Document()
        
        # 设置样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(11)
        
        # 获取当前日期
        current_date = datetime.now().strftime('%Y年%m月%d日')
        issue = journal.issue
        
        # 转换期刊号格式：从"2025,42(3)"转换为"2025年第42卷第3期"
        def convert_issue_format(issue_str):
            """
            将期刊号格式从"2025,42(3)"转换为"2025年第42卷第3期"
            支持多种可能的格式：
            - "2025,42(3)" -> "2025年第42卷第3期"
            - "2025,42(3-4)" -> "2025年第42卷第3-4期"
            - "2025,42(3,4)" -> "2025年第42卷第3-4期"
            """
            try:
                # 匹配格式：年份,卷号(期号)
                match = re.match(r'(\d{4}),\s*(\d+)\(([^)]+)\)', issue_str)
                if match:
                    year = match.group(1)  # 年份
                    volume = match.group(2)  # 卷号
                    issue_num = match.group(3)  # 期号
                    
                    # 处理期号中的特殊字符
                    issue_num = issue_num.replace(',', '-')  # 将逗号替换为连字符
                    
                    return f"{year}年第{volume}卷第{issue_num}期"
                else:
                    # 如果格式不匹配，返回原格式
                    logger.warning(f"无法解析期刊号格式: {issue_str}")
                    return issue_str
            except Exception as e:
                logger.error(f"转换期刊号格式失败: {str(e)}")
                return issue_str
        
        # 转换期刊号格式
        issue_info = convert_issue_format(issue)
        
        # 添加期刊标题
        title_para = doc.add_paragraph()
        title_run = title_para.runs[0] if title_para.runs else title_para.add_run()
        title_run.text = "东华大学学报"
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = None  # 黑色
        
        # 添加期刊号
        issue_para = doc.add_paragraph()
        issue_run = issue_para.runs[0] if issue_para.runs else issue_para.add_run()
        issue_run.text = issue_info
        issue_run.font.size = Pt(12)
        
        # 添加分隔线
        doc.add_paragraph("─" * 30)  # 使用字符作为分隔线
        
        # 添加编辑信息
        editor_para = doc.add_paragraph()
        editor_run = editor_para.runs[0] if editor_para.runs else editor_para.add_run()
        editor_run.text = f"本期责编: 编辑部 | {current_date}"
        editor_run.font.size = Pt(10)
        editor_run.font.italic = True
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加论文内容
        for i, paper in enumerate(papers, 1):
            # 处理数据库Paper对象或字典
            if hasattr(paper, 'title'):
                # 数据库对象
                title = paper.title
                authors = paper.authors
                chinese_title = getattr(paper, 'chinese_title', '') or ''
                chinese_authors = getattr(paper, 'chinese_authors', '') or ''
                doi = paper.doi or ''
                page_start = paper.page_start
                page_end = paper.page_end
                # 获取图片URL
                first_image_url = getattr(paper, 'first_image_url', '') or ''
                second_image_url = getattr(paper, 'second_image_url', '') or ''
                # 格式化作者用于citation（最多3个，名字只取首字母）
                formatted_authors = format_authors_for_citation(authors, max_authors=3)
                # 生成引用信息
                citation = f"{formatted_authors}. {title} [J]. Journal of Donghua University (English Edition), 2025, 42(3): {page_start}-{page_end}."
            
            # 添加中文标题（如果有）
            if chinese_title:
                chinese_title_para = doc.add_paragraph()
                chinese_title_run = chinese_title_para.runs[0] if chinese_title_para.runs else chinese_title_para.add_run()
                chinese_title_run.text = f"{i}. {chinese_title}"
                chinese_title_run.font.size = Pt(12)
                chinese_title_run.font.bold = True
            
            # 添加中文作者（如果有）
            if chinese_authors:
                chinese_authors_para = doc.add_paragraph()
                chinese_authors_run = chinese_authors_para.runs[0] if chinese_authors_para.runs else chinese_authors_para.add_run()
                chinese_authors_run.text = chinese_authors
                chinese_authors_run.font.size = Pt(11)
            
            # 添加第二张图片（在中文标题和中文作者下面）
            if second_image_url:
                try:
                    logger.info(f"开始处理第二张图片: {second_image_url}")
                    # 获取本地图片路径
                    local_image_path = get_local_image_path(second_image_url, temp_dir)
                    logger.info(f"本地图片路径检查结果: {local_image_path}, 文件存在: {os.path.exists(local_image_path) if local_image_path else False}")
                    
                    if local_image_path and os.path.exists(local_image_path):
                        # 添加图片描述
                        # image_desc_para = doc.add_paragraph()
                        # image_desc_run = image_desc_para.runs[0] if image_desc_para.runs else image_desc_para.add_run()
                        # image_desc_run.text = "论文配图:"
                        # image_desc_run.font.size = Pt(10)
                        # image_desc_run.font.italic = True
                        
                        # 插入图片到Word文档 - 添加异常处理
                        try:
                            logger.info(f"尝试插入图片到Word: {local_image_path}")
                            doc.add_picture(local_image_path, width=Pt(300))  # 设置图片宽度为300磅
                            logger.info(f"✅ 为论文 {i} 成功插入第二张图片: {second_image_url}")
                        except Exception as picture_error:
                            logger.error(f"❌ 插入图片到Word失败: {str(picture_error)}")
                            # 如果插入失败，添加图片路径作为替代
                            url_para = doc.add_paragraph()
                            url_run = url_para.runs[0] if url_para.runs else url_para.add_run()
                            url_run.text = f"图片路径: {second_image_url}"
                            url_run.font.size = Pt(9)
                    else:
                        logger.warning(f"❌ 无法找到第二张图片: {second_image_url}")
                except Exception as img_error:
                    logger.error(f"❌ 处理第二张图片失败: {str(img_error)}")
            
            # 添加论文标题
            paper_title_para = doc.add_paragraph()
            paper_title_run = paper_title_para.runs[0] if paper_title_para.runs else paper_title_para.add_run()
            paper_title_run.text = title
            paper_title_run.font.size = Pt(12)
            paper_title_run.font.bold = True
            
            # 添加作者信息（去掉"作者："前缀）
            doc.add_paragraph(authors)
            
            # 添加DOI信息（保留"DOI："前缀）
            if doi:
                doc.add_paragraph(f"DOI: {doi}")
            
            # 添加引用信息（添加"Citation:"前缀，调整字体大小）
            citation_para = doc.add_paragraph()
            citation_run = citation_para.runs[0] if citation_para.runs else citation_para.add_run()
            citation_run.text = f"Citation: {citation}"
            citation_run.font.size = Pt(11)  # 调整字体大小和其他内容一样
            
            # 在论文版块最后添加"作者说链接/OSID"文字和第一张图片
            # 添加分隔线
            doc.add_paragraph("─" * 20)
            
            # 添加"作者说链接/OSID"文字
            osid_para = doc.add_paragraph()
            osid_run = osid_para.runs[0] if osid_para.runs else osid_para.add_run()
            osid_run.text = "作者说链接/OSID"
            osid_run.font.size = Pt(10)
            osid_run.font.bold = True
            
            # 添加第一张图片（QRcode）
            if first_image_url:
                try:
                    logger.info(f"开始处理第一张图片(QRcode): {first_image_url}")
                    # 获取本地图片路径
                    local_image_path = get_local_image_path(first_image_url, temp_dir)
                    logger.info(f"本地图片路径检查结果: {local_image_path}, 文件存在: {os.path.exists(local_image_path) if local_image_path else False}")
                    
                    if local_image_path and os.path.exists(local_image_path):
                        # 插入图片到Word文档 - 添加异常处理
                        try:
                            logger.info(f"尝试插入第一张图片到Word: {local_image_path}")
                            doc.add_picture(local_image_path, width=Pt(100))  # 设置较小的宽度用于二维码
                            logger.info(f"✅ 为论文 {i} 成功插入第一张图片(QRcode): {first_image_url}")
                        except Exception as picture_error:
                            logger.error(f"❌ 插入第一张图片到Word失败: {str(picture_error)}")
                            # 如果插入失败，添加图片路径作为替代
                            url_para = doc.add_paragraph()
                            url_run = url_para.runs[0] if url_para.runs else url_para.add_run()
                            url_run.text = f"二维码路径: {first_image_url}"
                            url_run.font.size = Pt(9)
                    else:
                        logger.warning(f"❌ 无法找到第一张图片(QRcode): {first_image_url}")
                except Exception as qrcode_error:
                    logger.error(f"❌ 处理第一张图片(QRcode)失败: {str(qrcode_error)}")
            
            # 添加空行分隔
            doc.add_paragraph()
        
        # 添加页脚信息
        doc.add_paragraph("─" * 30)  # 分隔线
        footer_para = doc.add_paragraph()
        footer_run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        footer_run.text = "感谢您的阅读！欢迎引用本文内容"
        footer_run.font.size = Pt(10)
        
        copyright_para = doc.add_paragraph()
        copyright_run = copyright_para.runs[0] if copyright_para.runs else copyright_para.add_run()
        copyright_run.text = "© 2025 东华大学学报 版权所有"
        copyright_run.font.size = Pt(9)
        
        # 保存Word文档
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"推文_{journal.issue}_{timestamp}.docx"
        output_path = os.path.join('uploads', filename)
        
        # 确保目录存在
        os.makedirs('uploads', exist_ok=True)
        
        doc.save(output_path)
        logger.info(f"推文Word文档已生成: {output_path}")
        
        # 返回文件路径
        return output_path
    
    except Exception as e:
        logger.error(f"生成推文Word文档失败: {str(e)}")
        raise Exception(f"生成推文Word文档失败: {str(e)}")

def generate_tuiwen_from_fields(papers, journal, fields_config: List[Dict]) -> str:
    """
    根据字段配置生成推文内容
    
    Args:
        papers: 论文列表
        journal: 期刊对象
        fields_config: 字段配置列表，格式: [{'key': 'title', 'label': '标题', 'order': 1}, ...]
    """
    try:
        # 创建Word文档
        doc = Document()
        
        # 设置样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(11)
        
        # 获取当前日期
        current_date = datetime.now().strftime('%Y年%m月%d日')
        issue = journal.issue
        
        # 转换期刊号格式
        def convert_issue_format(issue_str):
            try:
                match = re.match(r'(\d{4}),\s*(\d+)\(([^)]+)\)', issue_str)
                if match:
                    year = match.group(1)
                    volume = match.group(2)
                    issue_num = match.group(3).replace(',', '-')
                    return f"{year}年第{volume}卷第{issue_num}期"
                return issue_str
            except:
                return issue_str
        
        issue_info = convert_issue_format(issue)
        
        # 添加期刊标题
        title_para = doc.add_paragraph()
        title_run = title_para.runs[0] if title_para.runs else title_para.add_run()
        title_run.text = journal.title or "东华大学学报"
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # 添加期刊号
        issue_para = doc.add_paragraph()
        issue_run = issue_para.runs[0] if issue_para.runs else issue_para.add_run()
        issue_run.text = issue_info
        issue_run.font.size = Pt(12)
        
        # 添加分隔线
        doc.add_paragraph("─" * 30)
        
        # 添加编辑信息
        editor_para = doc.add_paragraph()
        editor_run = editor_para.runs[0] if editor_para.runs else editor_para.add_run()
        editor_run.text = f"本期责编: 编辑部 | {current_date}"
        editor_run.font.size = Pt(10)
        editor_run.font.italic = True
        
        # 添加空行
        doc.add_paragraph()
        
        # 创建临时目录用于存储图片（在整个函数中共享）
        temp_dir = os.path.join('temp_images', datetime.now().strftime('%Y%m%d_%H%M%S'))
        os.makedirs(temp_dir, exist_ok=True)
        
        # 按配置的字段顺序生成每篇论文的内容
        for paper_idx, paper in enumerate(papers, 1):
            # 处理数据库Paper对象或字典
            if hasattr(paper, 'title'):
                title = paper.title
                authors = paper.authors
                chinese_title = getattr(paper, 'chinese_title', '') or ''
                chinese_authors = getattr(paper, 'chinese_authors', '') or ''
                doi = paper.doi or ''
                page_start = paper.page_start
                page_end = paper.page_end
                formatted_authors = format_authors_for_citation(authors, max_authors=3)
                citation = f"{formatted_authors}. {title} [J]. Journal of Donghua University (English Edition), 2025, 42(3): {page_start}-{page_end}."
            else:
                title = paper.get('title', '')
                authors = paper.get('authors', '')
                chinese_title = paper.get('chinese_title', '') or ''
                chinese_authors = paper.get('chinese_authors', '') or ''
                doi = paper.get('doi', '')
                page_start = paper.get('page_start')
                page_end = paper.get('page_end')
                formatted_authors = format_authors_for_citation(authors, max_authors=3)
                citation = f"{formatted_authors}. {title} [J]. Journal of Donghua University (English Edition), 2025, 42(3): {page_start}-{page_end}."
            
            # 获取图片URL
            if hasattr(paper, 'first_image_url'):
                first_image_url = getattr(paper, 'first_image_url', '') or ''
                second_image_url = getattr(paper, 'second_image_url', '') or ''
            else:
                first_image_url = paper.get('first_image_url', '') or ''
                second_image_url = paper.get('second_image_url', '') or ''
            
            # 字段值映射
            field_values = {
                'chinese_title': chinese_title,
                'chinese_authors': chinese_authors,
                'title': title,
                'authors': authors,
                'doi': doi,
                'citation': citation,
                'page_start': str(page_start) if page_start else '',
                'page_end': str(page_end) if page_end else '',
                'first_image': first_image_url,  # 图片URL
                'second_image': second_image_url,  # 图片URL
            }
            
            # 按配置的字段顺序添加内容
            for field_config in sorted(fields_config, key=lambda x: x.get('order', 999)):
                field_key = field_config.get('field') or field_config.get('key')  # 支持两种格式
                field_label = field_config.get('label', '')
                
                if field_key in field_values:
                    value = field_values[field_key]
                    
                    # 处理图片字段
                    if field_key == 'second_image' and value:
                        try:
                            local_image_path = get_local_image_path(value, temp_dir)
                            if local_image_path and os.path.exists(local_image_path):
                                try:
                                    doc.add_picture(local_image_path, width=Pt(300))
                                    logger.info(f"✅ 为论文 {paper_idx} 成功插入第二张图片")
                                except Exception as picture_error:
                                    logger.error(f"❌ 插入第二张图片失败: {str(picture_error)}")
                                    url_para = doc.add_paragraph()
                                    url_run = url_para.runs[0] if url_para.runs else url_para.add_run()
                                    url_run.text = f"图片路径: {value}"
                                    url_run.font.size = Pt(9)
                        except Exception as img_error:
                            logger.error(f"❌ 处理第二张图片失败: {str(img_error)}")
                    
                    elif field_key == 'first_image' and value:
                        # 添加分隔线
                        doc.add_paragraph("─" * 20)
                        # 添加"作者说链接/OSID"文字
                        osid_para = doc.add_paragraph()
                        osid_run = osid_para.runs[0] if osid_para.runs else osid_para.add_run()
                        osid_run.text = "作者说链接/OSID"
                        osid_run.font.size = Pt(10)
                        osid_run.font.bold = True
                        # 插入图片
                        try:
                            local_image_path = get_local_image_path(value, temp_dir)
                            if local_image_path and os.path.exists(local_image_path):
                                try:
                                    doc.add_picture(local_image_path, width=Pt(100))
                                    logger.info(f"✅ 为论文 {paper_idx} 成功插入第一张图片(QRcode)")
                                except Exception as picture_error:
                                    logger.error(f"❌ 插入第一张图片失败: {str(picture_error)}")
                                    url_para = doc.add_paragraph()
                                    url_run = url_para.runs[0] if url_para.runs else url_para.add_run()
                                    url_run.text = f"二维码路径: {value}"
                                    url_run.font.size = Pt(9)
                        except Exception as img_error:
                            logger.error(f"❌ 处理第一张图片失败: {str(img_error)}")
                    
                    # 处理文本字段
                    elif value and field_key not in ['first_image', 'second_image']:
                        # 添加字段标签和值
                        para = doc.add_paragraph()
                        run = para.runs[0] if para.runs else para.add_run()
                        
                        # 根据字段类型设置格式
                        if field_key in ['chinese_title', 'title']:
                            run.text = f"{paper_idx}. {value}"
                            run.font.size = Pt(12)
                            run.font.bold = True
                        elif field_key in ['chinese_authors', 'authors']:
                            run.text = value
                            run.font.size = Pt(11)
                        elif field_key == 'doi':
                            run.text = f"DOI: {value}"
                            run.font.size = Pt(11)
                        elif field_key == 'citation':
                            run.text = f"Citation: {value}"
                            run.font.size = Pt(11)
                        else:
                            run.text = f"{field_label}: {value}"
                            run.font.size = Pt(11)
            
            # 添加空行分隔
            doc.add_paragraph()
        
        # 添加页脚信息
        doc.add_paragraph("─" * 30)
        footer_para = doc.add_paragraph()
        footer_run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        footer_run.text = "感谢您的阅读！欢迎引用本文内容"
        footer_run.font.size = Pt(10)
        
        copyright_para = doc.add_paragraph()
        copyright_run = copyright_para.runs[0] if copyright_para.runs else copyright_para.add_run()
        copyright_run.text = "© 2025 东华大学学报 版权所有"
        copyright_run.font.size = Pt(9)
        
        # 保存文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"推文_{journal.issue}_{timestamp}.docx"
        output_path = os.path.join('uploads', filename)
        
        os.makedirs('uploads', exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"根据字段配置生成推文: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"根据字段配置生成推文失败: {str(e)}")
        raise Exception(f"根据字段配置生成推文失败: {str(e)}")

def generate_tuiwen_from_template(papers, journal, template_file_path: str) -> str:
    """
    基于Word模板生成推文内容
    
    Args:
        papers: 论文列表
        journal: 期刊对象
        template_file_path: 模板文件路径
    """
    try:
        import re
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 创建临时目录用于存储下载的图片
        temp_dir = os.path.join('temp_images', datetime.now().strftime('%Y%m%d_%H%M%S'))
        os.makedirs(temp_dir, exist_ok=True)
        
        # 加载模板文件
        doc = Document(template_file_path)
        
        # 获取当前日期
        current_date = datetime.now().strftime('%Y年%m月%d日')
        issue = journal.issue
        
        # 转换期刊号格式
        def convert_issue_format(issue_str):
            try:
                match = re.match(r'(\d{4}),\s*(\d+)\(([^)]+)\)', issue_str)
                if match:
                    year = match.group(1)
                    volume = match.group(2)
                    issue_num = match.group(3).replace(',', '-')
                    return f"{year}年第{volume}卷第{issue_num}期"
                return issue_str
            except:
                return issue_str
        
        issue_info = convert_issue_format(issue)
        
        # 定义占位符映射
        placeholders = {
            '{journal_title}': journal.title or '东华大学学报',
            '{issue}': issue_info,
            '{current_date}': current_date,
        }
        
        # 替换文档中的占位符（全局替换）
        def replace_placeholders_in_document(doc, placeholders):
            """在文档中替换占位符"""
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    text = run.text
                    for placeholder, value in placeholders.items():
                        if placeholder in text:
                            text = text.replace(placeholder, str(value))
                    run.text = text
            
            # 也处理表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                text = run.text
                                for placeholder, value in placeholders.items():
                                    if placeholder in text:
                                        text = text.replace(placeholder, str(value))
                                run.text = text
        
        # 先替换全局占位符
        replace_placeholders_in_document(doc, placeholders)
        
        # 查找论文循环标记（例如：{paper_loop_start} 和 {paper_loop_end}）
        # 如果没有找到循环标记，则在整个文档中查找论文相关的占位符
        paper_loop_start_idx = None
        paper_loop_end_idx = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if '{paper_loop_start}' in text or '{paper_start}' in text:
                paper_loop_start_idx = i
            if '{paper_loop_end}' in text or '{paper_end}' in text:
                paper_loop_end_idx = i
        
        # 如果找到了循环标记，在循环区域内为每篇论文生成内容
        if paper_loop_start_idx is not None and paper_loop_end_idx is not None and paper_loop_end_idx > paper_loop_start_idx:
            # 保存循环模板段落的内容（在删除前）
            loop_template_data = []
            for para_idx in range(paper_loop_start_idx + 1, paper_loop_end_idx):
                template_para = doc.paragraphs[para_idx]
                para_data = {
                    'text': template_para.text,
                    'style': template_para.style.name if template_para.style else None,
                    'runs': []
                }
                # 保存每个run的信息
                for run in template_para.runs:
                    run_data = {
                        'text': run.text,
                        'font_name': run.font.name if run.font and run.font.name else None,
                        'font_size': run.font.size if run.font and run.font.size else None,
                        'bold': run.font.bold if run.font else None,
                        'italic': run.font.italic if run.font else None,
                        'color': run.font.color.rgb if run.font and run.font.color and run.font.color.rgb else None,
                    }
                    para_data['runs'].append(run_data)
                loop_template_data.append(para_data)
            
            # 删除循环标记段落和模板段落（从后往前删除，避免索引变化）
            # 先收集要删除的段落元素
            paras_to_remove = []
            for para_idx in range(paper_loop_start_idx, paper_loop_end_idx + 1):
                paras_to_remove.append(doc.paragraphs[para_idx]._element)
            
            # 删除段落元素
            for para_element in paras_to_remove:
                para_element.getparent().remove(para_element)
            
            # 为每篇论文生成内容
            for paper_idx, paper in enumerate(papers):
                # 处理数据库Paper对象或字典
                if hasattr(paper, 'title'):
                    title = paper.title
                    authors = paper.authors
                    chinese_title = getattr(paper, 'chinese_title', '') or ''
                    chinese_authors = getattr(paper, 'chinese_authors', '') or ''
                    doi = paper.doi or ''
                    page_start = paper.page_start
                    page_end = paper.page_end
                    formatted_authors = format_authors_for_citation(authors, max_authors=3)
                    citation = f"{formatted_authors}. {title} [J]. Journal of Donghua University (English Edition), 2025, 42(3): {page_start}-{page_end}."
                else:
                    # 字典格式
                    title = paper.get('title', '')
                    authors = paper.get('authors', '')
                    chinese_title = paper.get('chinese_title', '') or ''
                    chinese_authors = paper.get('chinese_authors', '') or ''
                    doi = paper.get('doi', '')
                    page_start = paper.get('page_start')
                    page_end = paper.get('page_end')
                    formatted_authors = format_authors_for_citation(authors, max_authors=3)
                    citation = f"{formatted_authors}. {title} [J]. Journal of Donghua University (English Edition), 2025, 42(3): {page_start}-{page_end}."
                
                # 论文占位符映射
                paper_placeholders = {
                    '{title}': title or '',
                    '{chinese_title}': chinese_title or '',
                    '{authors}': authors or '',
                    '{chinese_authors}': chinese_authors or '',
                    '{doi}': doi or '',
                    '{citation}': citation or '',
                    '{page_start}': str(page_start) if page_start else '',
                    '{page_end}': str(page_end) if page_end else '',
                    '{paper_index}': str(paper_idx + 1),
                }
                
                # 根据模板数据创建新段落并替换占位符
                for para_data in loop_template_data:
                    new_para = doc.add_paragraph()
                    
                    # 设置段落样式
                    if para_data['style']:
                        try:
                            new_para.style = para_data['style']
                        except:
                            pass
                    
                    # 创建runs并替换占位符
                    text = para_data['text']
                    # 替换占位符
                    for placeholder, value in paper_placeholders.items():
                        if placeholder in text:
                            text = text.replace(placeholder, str(value))
                    
                    # 如果有多个runs，尝试保持原有格式
                    if para_data['runs']:
                        # 简化处理：将所有文本合并为一个run
                        new_run = new_para.add_run(text)
                        # 使用第一个run的样式（如果有）
                        first_run_data = para_data['runs'][0]
                        try:
                            if first_run_data['font_name']:
                                new_run.font.name = first_run_data['font_name']
                            if first_run_data['font_size']:
                                new_run.font.size = first_run_data['font_size']
                            if first_run_data['bold'] is not None:
                                new_run.font.bold = first_run_data['bold']
                            if first_run_data['italic'] is not None:
                                new_run.font.italic = first_run_data['italic']
                            if first_run_data['color']:
                                new_run.font.color.rgb = first_run_data['color']
                        except Exception as style_error:
                            logger.warning(f"应用run样式失败: {str(style_error)}")
                    else:
                        # 没有runs，直接添加文本
                        new_para.add_run(text)
                
                # 论文之间添加空行（可选）
                if paper_idx < len(papers) - 1:
                    doc.add_paragraph()
                
        else:
            # 没有找到循环标记，在整个文档中替换论文占位符（只替换第一篇论文）
            if papers:
                paper = papers[0]
                if hasattr(paper, 'title'):
                    title = paper.title
                    authors = paper.authors
                    chinese_title = getattr(paper, 'chinese_title', '') or ''
                    chinese_authors = getattr(paper, 'chinese_authors', '') or ''
                    doi = paper.doi or ''
                    page_start = paper.page_start
                    page_end = paper.page_end
                    formatted_authors = format_authors_for_citation(authors, max_authors=3)
                    citation = f"{formatted_authors}. {title} [J]. Journal of Donghua University (English Edition), 2025, 42(3): {page_start}-{page_end}."
                
                paper_placeholders = {
                    '{title}': title or '',
                    '{chinese_title}': chinese_title or '',
                    '{authors}': authors or '',
                    '{chinese_authors}': chinese_authors or '',
                    '{doi}': doi or '',
                    '{citation}': citation or '',
                    '{page_start}': str(page_start) if page_start else '',
                    '{page_end}': str(page_end) if page_end else '',
                }
                
                replace_placeholders_in_document(doc, paper_placeholders)
        
        # 保存文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"推文_{journal.issue}_{timestamp}.docx"
        output_path = os.path.join('uploads', filename)
        
        os.makedirs('uploads', exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"基于模板生成推文: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"基于模板生成推文失败: {str(e)}")
        raise Exception(f"基于模板生成推文失败: {str(e)}")
