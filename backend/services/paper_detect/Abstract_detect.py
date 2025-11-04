#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

"""
=== 论文格式检测系统 - 摘要检测器 ===

【摘要检测 (Abstract Detection)】

1. 【结构检测】
   - Abstract字段后紧接冒号，冒号后紧接摘要内容
   - 摘要不分段（单一段落）
   - 内容长度验证（50-2000字符）

2. 【格式检测】
   - 字体为Times New Roman，加粗，1.5倍行间距
   - 字号为小四(12pt)
   - 无斜体要求，无行前行后间距要求
   - 两端对齐，无缩进（顶格书写）

3. 【技术特性】
   - 支持段落样式和直接格式的混合检测
   - 智能段落对齐检测（处理复杂的样式继承）
   - 完整的结构、段落、格式三维度检测
   - 详细的格式错误诊断和建议
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 简化的字体检测 ----------
def detect_font_for_run(run, paragraph=None):
    """
    简化的字体检测函数（专用于摘要检测）
    返回 (font_size_pt, font_name, is_bold, is_italic, line_spacing)
    """
    if not run:
        return 12.0, "Unknown", False, False, 1.0
    
    # 1. 字体名称检测
    font_name = "Unknown"
    try:
        if getattr(run.font, 'name', None):
            font_name = run.font.name
        else:
            font_name = "Times New Roman"  # 默认使用学术常用字体
    except Exception:
        font_name = "Times New Roman"
    
    # 2. 字号检测
    font_size = 12.0  # 默认小四
    try:
        if getattr(run.font, 'size', None) and getattr(run.font.size, 'pt', None):
            font_size = float(run.font.size.pt)
        elif hasattr(run._element, 'rPr') and run._element.rPr is not None:
            sz = run._element.rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
            if sz is not None and sz.get('val'):
                font_size = float(int(sz.get('val'))) / 2.0
    except Exception:
        pass
    
    # 3. 加粗检测
    is_bold = False
    try:
        if getattr(run.font, 'bold', None) is not None:
            is_bold = bool(run.font.bold)
        elif hasattr(run._element, 'rPr') and run._element.rPr is not None:
            b = run._element.rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
            is_bold = b is not None
    except Exception:
        pass
    
    # 4. 斜体检测
    is_italic = False
    try:
        if getattr(run.font, 'italic', None) is not None:
            is_italic = bool(run.font.italic)
        elif hasattr(run._element, 'rPr') and run._element.rPr is not None:
            i = run._element.rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i')
            is_italic = i is not None
    except Exception:
        pass
    
    # 5. 行间距检测
    line_spacing = 1.0  # 默认单倍行距
    try:
        if paragraph and paragraph.paragraph_format.line_spacing:
            line_spacing = float(paragraph.paragraph_format.line_spacing)
    except Exception:
        pass
    
    return font_size, font_name, is_bold, is_italic, line_spacing

def get_font_size(pt_size, tpl=None):
    """字体大小转换为中文字号"""
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        size_map = {
            9: "小五", 10.5: "五号", 12: "小四", 14: "四号",
            16: "三号", 18: "小二", 22: "二号", 24: "小一", 26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

def get_line_spacing_name(spacing, tpl=None):
    """行间距转换为中文描述"""
    if tpl and 'check_rules' in tpl and 'line_spacing_mapping' in tpl['check_rules']:
        spacing_config = tpl['check_rules']['line_spacing_mapping']
        spacing_map = {}
        for key, value in spacing_config.items():
            try:
                spacing_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        spacing_map = {
            1.0: "单倍行距", 1.15: "1.15倍行距", 
            1.5: "1.5倍行距", 2.0: "双倍行距"
        }
    
    if not spacing_map:
        return f"{spacing}倍行距"
    
    closest_spacing = min(spacing_map.keys(), key=lambda x: abs(x - spacing))
    return spacing_map[closest_spacing]

def get_alignment_name(alignment_value, tpl=None):
    """段落对齐方式转换为中文描述"""
    alignment_map = {
        0: "左对齐", 1: "居中对齐", 2: "右对齐", 3: "两端对齐"
    }
    return alignment_map.get(alignment_value, f"未知对齐({alignment_value})")

def get_style_alignment(doc, style_id):
    """从文档的styles.xml中查找并返回指定样式的对齐方式。"""
    try:
        styles = doc.styles.element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        style_path = f".//w:style[@w:styleId='{style_id}']//w:jc"
        jc_element = styles.find(style_path, ns)
        if jc_element is not None:
            val = jc_element.get(ns['w'] + 'val')
            mapping = {'left': 0, 'center': 1, 'right': 2, 'both': 3, 'justify': 3}
            return mapping.get(val)
    except Exception:
        return None
    return None

# ---------- 请将您代码中旧的 detect_paragraph_alignment 函数替换为这个新版本 ----------

def detect_paragraph_alignment(paragraph):
    """
    改进的段落对齐检测，按照明确的优先级处理直接格式和样式格式
    
    检测逻辑：
    1. 如果直接格式不为None，直接使用直接格式结果
    2. 如果直接格式为None，但样式格式为JUSTIFY，归为两端对齐
    3. 如果直接格式不是两端对齐，就不是两端对齐
    4. 如果样式格式不是两端对齐，也不是两端对齐
    5. 最终默认为左对齐
    """
    
    # 优先级1: 检查直接格式
    direct_alignment = paragraph.paragraph_format.alignment
    if direct_alignment is not None:
        # 如果有直接格式设置，直接使用
        return int(direct_alignment)
    
    # 优先级2: 如果直接格式为None，检查样式格式
    if paragraph.style:
        try:
            style_alignment = paragraph.style.paragraph_format.alignment
            if style_alignment is not None:
                # 如果样式有对齐设置，使用样式对齐
                return int(style_alignment)
        except Exception:
            pass
    
    # 优先级3: 检查XML样式定义（处理复杂的样式继承）
    if paragraph.style and paragraph.style.style_id:
        try:
            xml_alignment = get_style_alignment(paragraph.part.document, paragraph.style.style_id)
            if xml_alignment is not None:
                return xml_alignment
        except Exception:
            pass
            
    # 默认值: 左对齐
    return WD_PARAGRAPH_ALIGNMENT.LEFT

def detect_paragraph_indent(paragraph):
    """检测段落缩进"""
    try:
        fmt = paragraph.paragraph_format
        first_line_indent = fmt.first_line_indent.pt if fmt.first_line_indent else 0.0
        left_indent = fmt.left_indent.pt if fmt.left_indent else 0.0
        right_indent = fmt.right_indent.pt if fmt.right_indent else 0.0
        return first_line_indent, left_indent, right_indent
    except Exception:
        return 0.0, 0.0, 0.0

# ---------- 摘要检测逻辑 ----------
def check_abstract_structure(text, tpl):
    """
    检查摘要结构（Abstract:冒号格式）
    返回 {'ok': bool, 'messages': [], 'content': str}
    """
    report = {'ok': True, 'messages': [], 'content': ''}
    
    # 检查Abstract:格式
    structure_rules = tpl.get('structure_rules', {})
    header_pattern = structure_rules.get('header_pattern', r'^\\s*Abstract\\s*:\\s*(.+)$')
    
    try:
        match = re.match(header_pattern, text.strip(), re.DOTALL | re.IGNORECASE)
        if match:
            report['content'] = match.group(1).strip()
            ok_msg = tpl.get('messages', {}).get('structure_header_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
        else:
            report['ok'] = False
            error_msg = tpl.get('messages', {}).get('structure_header_error')
            if error_msg:
                report['messages'].append(error_msg)
    except Exception:
        report['ok'] = False
        report['messages'].append("摘要格式检查出错")
    
    # 检查内容长度
    if report['content']:
        content_length = len(report['content'])
        min_length = structure_rules.get('min_content_length', 50)
        max_length = structure_rules.get('max_content_length', 2000)
        
        if content_length < min_length:
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('structure_length_short')
            if msg_tpl:
                try:
                    report['messages'].append(msg_tpl.format(min=min_length))
                except:
                    report['messages'].append(msg_tpl)
        elif content_length > max_length:
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('structure_length_long')
            if msg_tpl:
                try:
                    report['messages'].append(msg_tpl.format(max=max_length))
                except:
                    report['messages'].append(msg_tpl)
        else:
            ok_msg = tpl.get('messages', {}).get('structure_length_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
    
    return report

def check_abstract_paragraphs(doc, tpl):
    """
    检查摘要是否分段
    返回 {'ok': bool, 'messages': [], 'abstract_paragraph': paragraph}
    """
    report = {'ok': True, 'messages': [], 'abstract_paragraph': None}
    
    # 查找包含Abstract的段落
    abstract_paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text and re.search(r'\bAbstract\s*:', paragraph.text, re.IGNORECASE):
            abstract_paragraphs.append(paragraph)
    
    if len(abstract_paragraphs) == 1:
        report['abstract_paragraph'] = abstract_paragraphs[0]
        ok_msg = tpl.get('messages', {}).get('structure_no_paragraph_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
    elif len(abstract_paragraphs) > 1:
        report['ok'] = False
        error_msg = tpl.get('messages', {}).get('structure_paragraph_error')
        if error_msg:
            report['messages'].append(error_msg)
    else:
        report['ok'] = False
        report['messages'].append("未找到Abstract段落")
    
    return report

def check_abstract_format(paragraph, tpl):
    """
    检查摘要格式（字体、加粗、行间距等）
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not paragraph or not paragraph.runs:
        report['ok'] = False
        report['messages'].append("摘要段落没有文本内容")
        return report
    
    # 取第一个非空 run
    main_run = None
    for run in paragraph.runs:
        if run.text.strip():
            main_run = run
            break
    
    if not main_run:
        report['ok'] = False
        report['messages'].append("摘要段落没有有效文本")
        return report
    
    # 获取格式规则
    format_rules = tpl.get('format_rules', {}).get('abstract', {})
    
    # 检测实际格式
    actual_size_pt, actual_font_name, actual_bold, actual_italic, actual_line_spacing = detect_font_for_run(main_run, paragraph)
    
    issues = []
    
    # 字体大小检查
    if 'font_size_pt' in format_rules:
        expected_size_pt = float(format_rules['font_size_pt'])
        actual_size_name = get_font_size(actual_size_pt, tpl)
        expected_size_name = get_font_size(expected_size_pt, tpl)
        print(f"字体大小: {actual_size_name}（{actual_size_pt}pt）(期望: {expected_size_name}（{expected_size_pt}pt）)")
        if abs(actual_size_pt - expected_size_pt) > 0.5:
            issues.append(f"字体大小应为{expected_size_name}（{expected_size_pt}pt），实际为{actual_size_name}（{actual_size_pt}pt）")
    
    # 字体名称检查
    if 'font_name' in format_rules:
        expected_font_name = str(format_rules['font_name'])
        print(f"字体名称: {actual_font_name} (期望: {expected_font_name})")
        if expected_font_name.lower() not in actual_font_name.lower():
            issues.append(f"字体应为{expected_font_name}，实际为{actual_font_name}")
    
    # 加粗检查
    if 'bold' in format_rules:
        expected_bold = bool(format_rules['bold'])
        print(f"加粗: {'是' if actual_bold else '否'} (期望: {'是' if expected_bold else '否'})")
        if actual_bold != expected_bold:
            bold_status = "加粗" if expected_bold else "不加粗"
            actual_status = "加粗" if actual_bold else "不加粗"
            issues.append(f"字体应为{bold_status}，实际为{actual_status}")
    
    # 斜体检查
    if 'italic' in format_rules:
        expected_italic = bool(format_rules['italic'])
        print(f"斜体: {'是' if actual_italic else '否'} (期望: {'是' if expected_italic else '否'})")
        if actual_italic != expected_italic:
            italic_status = "斜体" if expected_italic else "正体"
            actual_status = "斜体" if actual_italic else "正体"
            issues.append(f"字体应为{italic_status}，实际为{actual_status}")
    
    # 行间距检查
    if 'line_spacing' in format_rules:
        expected_line_spacing = float(format_rules['line_spacing'])
        actual_spacing_name = get_line_spacing_name(actual_line_spacing, tpl)
        expected_spacing_name = get_line_spacing_name(expected_line_spacing, tpl)
        print(f"行间距: {actual_spacing_name}（{actual_line_spacing}倍）(期望: {expected_spacing_name}（{expected_line_spacing}倍）)")
        if abs(actual_line_spacing - expected_line_spacing) > 0.1:
            issues.append(f"行间距应为{expected_spacing_name}（{expected_line_spacing}倍），实际为{actual_spacing_name}（{actual_line_spacing}倍）")
    
    # 段落对齐检查
    if 'alignment' in format_rules:
        expected_alignment_str = str(format_rules['alignment'])
        alignment_map = {"left": 0, "center": 1, "right": 2, "justify": 3}
        expected_alignment = alignment_map.get(expected_alignment_str, 0)
        
        actual_alignment = detect_paragraph_alignment(paragraph)
        actual_alignment_name = get_alignment_name(actual_alignment, tpl)
        expected_alignment_name = get_alignment_name(expected_alignment, tpl)
        print(f"段落对齐: {actual_alignment_name} (期望: {expected_alignment_name})")
        if actual_alignment != expected_alignment:
            issues.append(f"段落应为{expected_alignment_name}，实际为{actual_alignment_name}")
    
    # 段落缩进检查
    if 'first_line_indent' in format_rules or 'left_indent' in format_rules or 'right_indent' in format_rules:
        first_line_indent, left_indent, right_indent = detect_paragraph_indent(paragraph)
        
        if 'first_line_indent' in format_rules:
            expected_first_indent = float(format_rules['first_line_indent'])
            print(f"首行缩进: {first_line_indent:.1f}pt (期望: {expected_first_indent}pt)")
            if abs(first_line_indent - expected_first_indent) > 1.0:  # 1pt容差
                issues.append(f"首行缩进应为{expected_first_indent}pt，实际为{first_line_indent:.1f}pt")
        
        if 'left_indent' in format_rules:
            expected_left_indent = float(format_rules['left_indent'])
            print(f"左缩进: {left_indent:.1f}pt (期望: {expected_left_indent}pt)")
            if abs(left_indent - expected_left_indent) > 1.0:
                issues.append(f"左缩进应为{expected_left_indent}pt，实际为{left_indent:.1f}pt")
        
        if 'right_indent' in format_rules:
            expected_right_indent = float(format_rules['right_indent'])
            print(f"右缩进: {right_indent:.1f}pt (期望: {expected_right_indent}pt)")
            if abs(right_indent - expected_right_indent) > 1.0:
                issues.append(f"右缩进应为{expected_right_indent}pt，实际为{right_indent:.1f}pt")
    
    print(f"发现 {len(issues)} 个格式问题")
    print("---")
    
    if issues:
        report['ok'] = False
        header = tpl.get('messages', {}).get('format_abstract_issue_header')
        if header:
            report['messages'].append(header)
        report['messages'].extend([f"  - {i}" for i in issues])
    else:
        ok_msg = tpl.get('messages', {}).get('format_abstract_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
    
    return report

def check_abstract_with_template(doc_path, template_identifier):
    """
    主检查函数：检查摘要格式
    """
    tpl = load_template(template_identifier)
    doc = Document(doc_path)
    
    # 查找摘要段落
    abstract_text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text and re.search(r'\bAbstract\s*:', paragraph.text, re.IGNORECASE):
            abstract_text = paragraph.text.strip()
            break
    print(f"abstract_text: {abstract_text}")
    if not abstract_text:
        return {
            'structure': {'ok': False, 'messages': ['未找到Abstract段落']},
            'paragraphs': {'ok': False, 'messages': ['未找到Abstract段落']},
            'format': {'ok': False, 'messages': ['未找到Abstract段落']},
            'summary': ['摘要检查失败：未找到Abstract段落']
        }
    
    # 执行各项检查
    structure_report = check_abstract_structure(abstract_text, tpl)
    paragraphs_report = check_abstract_paragraphs(doc, tpl)
    
    if paragraphs_report['abstract_paragraph']:
        format_report = check_abstract_format(paragraphs_report['abstract_paragraph'], tpl)
    else:
        format_report = {'ok': False, 'messages': ['无法检测摘要格式']}
    
    # 组装报告
    report = {
        'structure': structure_report,
        'paragraphs': paragraphs_report,
        'format': format_report,
        'summary': []
    }
    
    # 生成总结
    all_ok = (structure_report['ok'] and paragraphs_report['ok'] and format_report['ok'])
    summary_tpl = tpl.get('messages', {}).get('summary_overall')
    if summary_tpl:
        try:
            report['summary'].append(summary_tpl.format(ok=all_ok))
        except Exception:
            report['summary'].append(str(summary_tpl))
    
    return report

# ---------- 报告输出 ----------
def print_abstract_report(report):
    """打印摘要检查报告"""
    print("=== Abstract Check Report ===")
    
    sections = [
        ('structure', 'STRUCTURE'),
        ('paragraphs', 'PARAGRAPHS'), 
        ('format', 'FORMAT')
    ]
    
    for sec_key, sec_name in sections:
        info = report[sec_key]
        print(f"--- {sec_name} ---")
        print(" OK:", info['ok'])
        for m in info['messages']:
            print("  -", m)
    
    print("--- SUMMARY ---")
    for s in report['summary']:
        print(" ", s)

def print_help():
    print("Usage:")
    print("  python Abstract_detect.py check <paper.docx> <template.json_or_name>")

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print_help()
        sys.exit(0)
    
    cmd = sys.argv[1]
    if cmd == 'check':
        paper_path = sys.argv[2]
        tpl_id = sys.argv[3]
        if not os.path.isfile(paper_path):
            print(f"论文文件不存在: {paper_path}")
            sys.exit(1)
        try:
            print("=== 开始摘要格式检查 ===")
            report = check_abstract_with_template(paper_path, tpl_id)
        except Exception as e:
            print("检查时出错:", e)
            sys.exit(1)
        print_abstract_report(report)
    else:
        print_help()
        sys.exit(0)
'''
python paper_detect\Abstract_detect.py check template\test_abstract.docx templates\Abstract.json                               
'''