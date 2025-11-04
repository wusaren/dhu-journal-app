#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
from pathlib import Path

# 添加项目根目录到路径
if __name__ == '__main__':
    project_root = Path(__file__).parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

"""
=== 论文格式检测系统 - 中文部分检测器 ===

【中文部分检测 (Chinese Section Detection)】

检测参考文献后的中文部分，包括：
1. 中文标题 - 格式要求与英文标题一致
2. 中文作者 - 格式要求与英文作者一致
3. 中文单位 - 格式要求与英文单位一致
4. 中文摘要 - 格式要求与英文摘要一致
5. 中文关键词 - 格式要求与英文关键词一致

特殊要求：
- 英文部分使用 Times New Roman
- 中文部分使用 宋体
- 字体大小、行间距等与英文部分的检测规则一致
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    """解析模板路径，支持文件路径和模板名称"""
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    """加载JSON模板文件"""
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 字体检测函数 ----------
def detect_font_for_run(run, paragraph=None):
    """
    检测run的字体信息，包括中英文字体
    返回: (font_size, font_ascii, font_eastasia, is_bold, is_italic)
    """
    font_size = None
    font_name_ascii = None
    font_name_eastasia = None
    is_bold = False
    is_italic = False
    
    if not run:
        return 12.0, "Times New Roman", "宋体", False, False
    
    # 1. 检测字号
    try:
        if run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        if font_size is None and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass
    
    font_size = font_size if font_size is not None else 12.0
    
    # 2. 检测英文字体（ASCII）和中文字体（EastAsia）
    try:
        # 优先从run.font.name读取
        if run.font and run.font.name:
            font_name_ascii = run.font.name
        
        # 从XML中读取w:rFonts
        if hasattr(run._element, 'rPr'):
            rpr = run._element.rPr
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    # 读取各种字体属性
                    xml_ascii = rfonts.get(qn('w:ascii'))
                    xml_hansi = rfonts.get(qn('w:hAnsi'))
                    xml_eastasia = rfonts.get(qn('w:eastAsia'))
                    
                    # ASCII字体（英文）
                    if xml_ascii:
                        font_name_ascii = xml_ascii
                    elif xml_hansi and font_name_ascii is None:
                        font_name_ascii = xml_hansi
                    
                    # EastAsia字体（中文）
                    if xml_eastasia:
                        font_name_eastasia = xml_eastasia
                    elif xml_hansi and font_name_eastasia is None:
                        font_name_eastasia = xml_hansi
    except Exception as e:
        print(f"  字体检测异常: {e}")
    
    font_name_ascii = font_name_ascii if font_name_ascii else "Times New Roman"
    font_name_eastasia = font_name_eastasia if font_name_eastasia else "宋体"
    
    # 4. 检测加粗和斜体
    try:
        if run.font:
            is_bold = run.font.bold if run.font.bold is not None else False
            is_italic = run.font.italic if run.font.italic is not None else False
        
        is_bold = bool(is_bold)
        is_italic = bool(is_italic)
    except Exception:
        pass
    
    return font_size, font_name_ascii, font_name_eastasia, is_bold, is_italic

def get_paragraph_spacing(paragraph):
    """
    获取段落的段前段后间距
    返回: (space_before_lines, space_after_lines)
    """
    space_before = 0.0
    space_after = 0.0
    
    try:
        # 段前间距
        if paragraph.paragraph_format.space_before is not None:
            # 转换为行数（假设1行 = 12pt）
            space_before_pt = paragraph.paragraph_format.space_before.pt
            space_before = round(space_before_pt / 12.0, 1)
        
        # 段后间距
        if paragraph.paragraph_format.space_after is not None:
            space_after_pt = paragraph.paragraph_format.space_after.pt
            space_after = round(space_after_pt / 12.0, 1)
    except Exception:
        pass
    
    return space_before, space_after

def get_font_size(pt_size, tpl=None):
    """获取字体大小的中文名称"""
    size_map = {
        9: "小五",
        10.5: "五号",
        12: "小四",
        14: "四号",
        16: "三号",
        18: "小二",
        22: "二号",
        24: "小一",
        26: "一号"
    }
    
    if pt_size in size_map:
        return size_map[pt_size]
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

# ---------- 对齐检测函数 ----------
def detect_paragraph_alignment(paragraph):
    """检测段落对齐方式"""
    direct_alignment = paragraph.paragraph_format.alignment
    if direct_alignment is not None:
        return int(direct_alignment)
    
    if paragraph.style:
        try:
            style_alignment = paragraph.style.paragraph_format.alignment
            if style_alignment is not None:
                return int(style_alignment)
        except Exception:
            pass
    
    return 0  # 默认左对齐

# ---------- 中文部分定位 ----------
def find_references_section(doc):
    """查找参考文献部分的位置"""
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        # 参考文献可能的标识：References、参考文献、REFERENCES等
        if re.match(r'^(References|参考文献|REFERENCES)\s*$', text, re.IGNORECASE):
            return idx
    return None

def find_chinese_section(doc):
    """
    查找中文部分的起始位置
    返回：{
        'start_index': 起始段落索引,
        'title_index': 中文标题索引,
        'author_index': 中文作者索引,
        'affiliation_index': 中文单位索引,
        'abstract_index': 中文摘要索引,
        'keywords_index': 中文关键词索引
    }
    """
    references_idx = find_references_section(doc)
    
    if references_idx is None:
        return None
    
    result = {
        'start_index': references_idx,
        'title_index': None,
        'author_index': None,
        'affiliation_index': None,
        'abstract_index': None,
        'keywords_index': None
    }
    
    # 从参考文献后开始查找中文部分
    for idx in range(references_idx + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        # 判断是否包含中文字符
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', text))
        
        if has_chinese and result['title_index'] is None:
            # 第一个包含中文的段落视为中文标题
            result['title_index'] = idx
        elif result['title_index'] is not None and result['author_index'] is None:
            # 标题后的第一个段落视为作者
            if text and len(text) > 0:
                result['author_index'] = idx
        elif result['author_index'] is not None and result['affiliation_index'] is None:
            # 作者后的第一个段落视为单位
            if text and len(text) > 0:
                result['affiliation_index'] = idx
        
        # 查找摘要（包含"摘要"、"摘  要"等）
        if re.search(r'摘\s*要', text):
            result['abstract_index'] = idx
        
        # 查找关键词（包含"关键词"、"关键字"等）
        if re.search(r'关键(词|字)', text):
            result['keywords_index'] = idx
    
    return result

# ---------- 检测函数 ----------
def check_paragraph_full_format(paragraph, section_name, expected_format, tpl):
    """
    完整的段落格式检查（遍历所有run）
    """
    report = {'ok': True, 'messages': []}
    
    # 只打印一次段落的总体格式
    print(f"  检测Run数量: {len(paragraph.runs)}")
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_size, font_ascii, font_eastasia, is_bold, is_italic = detect_font_for_run(first_run, paragraph)
        print(f"  段落总体格式 - 中文: {font_eastasia}, 英文: {font_ascii}, 字号: {get_font_size(font_size)}({font_size}pt)")
    
    # 遍历段落中的所有run进行详细检查
    for run in paragraph.runs:
        if not run.text.strip():
            continue  # 跳过空白run
        
        # 1. 检测字体
        font_size, font_ascii, font_eastasia, is_bold, is_italic = detect_font_for_run(run, paragraph)
        
        # 检查run中的中文字符
        if re.search(r'[\u4e00-\u9fff]', run.text):
            expected_chinese_font = expected_format.get('chinese_font', '宋体')
            if font_eastasia != expected_chinese_font:
                msg = f"中文字体应为{expected_chinese_font}（当前：{font_eastasia}）"
                if msg not in report['messages']:
                    report['ok'] = False
                    report['messages'].append(msg)
        
        # 检查run中的英文字符
        if re.search(r'[a-zA-Z]', run.text):
            expected_english_font = expected_format.get('english_font', 'Times New Roman')
            if font_ascii != expected_english_font:
                msg = f"英文字体应为{expected_english_font}（当前：{font_ascii}）"
                if msg not in report['messages']:
                    report['ok'] = False
                    report['messages'].append(msg)
        
        # 检查字号
        expected_size = expected_format.get('font_size_pt', 12)
        if abs(font_size - expected_size) > 0.5:
            msg = f"字号应为{get_font_size(expected_size)}({expected_size}pt)，当前为{get_font_size(font_size)}({font_size}pt)"
            if msg not in report['messages']:
                report['ok'] = False
                report['messages'].append(msg)
        
        # 检查加粗
        expected_bold = expected_format.get('bold', False)
        if is_bold != expected_bold:
            if expected_bold:
                msg = f"{section_name}应加粗"
            else:
                msg = f"{section_name}不应加粗"
            if msg not in report['messages']:
                report['ok'] = False
                report['messages'].append(msg)
        
        # 检查斜体
        expected_italic = expected_format.get('italic', False)
        if is_italic != expected_italic:
            # 如果期望是正体，但检测到斜体，并且内容是单个字符，则忽略
            if not expected_italic and len(run.text.strip()) == 1:
                pass  # 允许单个字母为斜体，跳过报错
            else:
                if expected_italic:
                    msg = f"{section_name}应为斜体"
                else:
                    msg = f"{section_name}应为正体（不应为斜体，问题文本：'{run.text}'）"
                if msg not in report['messages']:
                    report['ok'] = False
                    report['messages'].append(msg)
    
    # 2. 检查对齐方式
    alignment = detect_paragraph_alignment(paragraph)
    alignment_names = {0: '左对齐', 1: '居中对齐', 2: '右对齐', 3: '两端对齐'}
    print(f"  对齐方式: {alignment_names.get(alignment)}")
    
    expected_alignment = expected_format.get('alignment', 'justify')
    if expected_alignment == 'justify' and alignment != 3:
        report['ok'] = False
        report['messages'].append(f"应两端对齐（当前：{alignment_names.get(alignment)}）")
    elif expected_alignment == 'center' and alignment != 1:
        report['ok'] = False
        report['messages'].append(f"应居中对齐（当前：{alignment_names.get(alignment)}）")
    
    # 3. 检查段前段后间距
    space_before, space_after = get_paragraph_spacing(paragraph)
    print(f"  段前间距: {space_before}行, 段后间距: {space_after}行")
    
    expected_before = expected_format.get('space_before', 0)
    expected_after = expected_format.get('space_after', 0)
    
    if abs(space_before - expected_before) > 0.2:
        report['ok'] = False
        report['messages'].append(f"段前间距应为{expected_before}行（当前：{space_before}行）")
    
    if abs(space_after - expected_after) > 0.2:
        report['ok'] = False
        report['messages'].append(f"段后间距应为{expected_after}行（当前：{space_after}行）")
    
    # 4. 检查行间距
    if 'line_spacing' in expected_format:
        line_spacing = paragraph.paragraph_format.line_spacing
        expected_spacing = expected_format.get('line_spacing', 1.5)
        
        if line_spacing:
            actual_spacing = float(line_spacing) if isinstance(line_spacing, (int, float)) else 1.0
            print(f"  行间距: {actual_spacing}倍")
            
            if abs(actual_spacing - expected_spacing) > 0.1:
                report['ok'] = False
                report['messages'].append(f"行间距应为{expected_spacing}倍（当前：{actual_spacing}倍）")
    
    # 5. 检查缩进
    if 'first_line_indent' in expected_format:
        first_line_indent = paragraph.paragraph_format.first_line_indent
        actual_indent = first_line_indent.pt if first_line_indent else 0
        expected_indent = expected_format.get('first_line_indent', 0)
        print(f"  首行缩进: {actual_indent}pt")
        
        if abs(actual_indent - expected_indent) > 1:
            report['ok'] = False
            report['messages'].append(f"首行缩进应为{expected_indent}pt（当前：{actual_indent}pt）")
    
    if 'left_indent' in expected_format:
        left_indent = paragraph.paragraph_format.left_indent
        actual_indent = left_indent.pt if left_indent else 0
        expected_indent = expected_format.get('left_indent', 0)
        
        if abs(actual_indent - expected_indent) > 1:
            report['ok'] = False
            report['messages'].append(f"左缩进应为{expected_indent}pt（当前：{actual_indent}pt）")
    
    if 'right_indent' in expected_format:
        right_indent = paragraph.paragraph_format.right_indent
        actual_indent = right_indent.pt if right_indent else 0
        expected_indent = expected_format.get('right_indent', 0)
        
        if abs(actual_indent - expected_indent) > 1:
            report['ok'] = False
            report['messages'].append(f"右缩进应为{expected_indent}pt（当前：{actual_indent}pt）")
    
    return report

def check_chinese_title(doc, chinese_section, tpl):
    """检测中文标题格式"""
    report = {'ok': True, 'messages': []}
    
    if chinese_section is None or chinese_section['title_index'] is None:
        report['ok'] = False
        report['messages'].append("未找到中文标题")
        return report
    
    title_para = doc.paragraphs[chinese_section['title_index']]
    title_text = title_para.text.strip()
    print(f"检测中文标题段落: '{title_text[:50]}...'")
    
    expected_format = tpl.get('format_rules', {}).get('title', {})
    return check_paragraph_full_format(title_para, "中文标题", expected_format, tpl)

def parse_chinese_authors(author_text):
    """
    解析中文作者的单位编号
    例如：刘影1，惠明明1，卜凡兴1*，罗维1*
    返回：[{'name': '刘影', 'affs': [1], 'corresponding': False}, ...]
    """
    authors = []
    
    # 用中英文逗号、分号分隔作者
    parts = re.split(r'[,，;；]+', author_text)
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        
        # 检测通讯作者标记（* 或 ☆）
        corresponding = bool(re.search(r'[*☆✉]', part))
        
        # 移除标记
        part_clean = re.sub(r'[*☆✉]+', '', part).strip()
        
        # 提取单位编号（数字在末尾，可能多个如"1,2"）
        match = re.search(r'(\d+(?:[,，]\d+)*)\s*$', part_clean)
        affs = []
        name = part_clean
        
        if match:
            # 提取编号
            aff_str = match.group(1)
            affs = [int(x) for x in re.findall(r'\d+', aff_str)]
            # 去掉编号，得到姓名
            name = part_clean[:match.start()].strip()
        
        authors.append({
            'name': name,
            'affs': affs,
            'corresponding': corresponding,
            'raw': part
        })
    
    return authors

def check_chinese_author(doc, chinese_section, tpl):
    """检测中文作者格式和内容"""
    report = {'ok': True, 'messages': []}
    
    if chinese_section is None or chinese_section['author_index'] is None:
        report['ok'] = False
        report['messages'].append("未找到中文作者")
        return report
    
    author_para = doc.paragraphs[chinese_section['author_index']]
    author_text = author_para.text.strip()
    print(f"检测中文作者段落: '{author_text[:50]}...'")
    
    # 1. 检测格式
    expected_format = tpl.get('format_rules', {}).get('author', {})
    format_result = check_paragraph_full_format(author_para, "中文作者", expected_format, tpl)
    report['ok'] = format_result['ok']
    report['messages'].extend(format_result['messages'])
    
    # 2. 解析作者和单位编号
    authors = parse_chinese_authors(author_text)
    print(f"  解析到 {len(authors)} 个作者")
    
    # 3. 检查作者单位编号规则
    if authors:
        # 收集所有使用的单位编号
        used_affs = set()
        for author in authors:
            used_affs.update(author['affs'])
        
        print(f"  使用的单位编号: {sorted(used_affs) if used_affs else '无'}")
        
        # 规则：如果所有作者来自同一单位，应该省略编号
        if len(used_affs) == 1 and len(used_affs) > 0:
            single_aff_id = list(used_affs)[0]
            report['ok'] = False
            report['messages'].append(
                f"所有作者来自同一单位（编号{single_aff_id}），应该省略作者姓名后的单位编号"
            )
            print(f"  作者单位检查: ✗ 所有作者都标注了单位{single_aff_id}，应省略编号")
        elif len(used_affs) == 0:
            # 没有单位编号（可能是正确的，如果只有一个单位）
            print(f"  作者单位检查: ✓ 作者未标注单位编号")
        else:
            print(f"  作者单位检查: ✓ 作者来自{len(used_affs)}个不同单位")
    
    return report

def check_chinese_affiliation(doc, chinese_section, authors_data, tpl):
    """检测中文单位格式和内容"""
    report = {'ok': True, 'messages': []}
    
    if chinese_section is None or chinese_section['affiliation_index'] is None:
        report['ok'] = False
        report['messages'].append("未找到中文单位")
        return report
    
    affiliation_start_idx = chinese_section['affiliation_index']
    
    # 1. 查找并解析所有单位段落
    print(f"  从段落 {affiliation_start_idx} 开始查找单位...")
    affiliation_paragraphs = []
    doc_affiliations = {}  # {编号: 单位文本}
    stop_keywords = ['摘要', '关键词', '关键字', 'Abstract', 'Keywords']
    
    for idx in range(affiliation_start_idx, min(affiliation_start_idx + 10, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        
        if not text:
            continue
        
        # 遇到摘要或关键词，停止
        if any(keyword in text for keyword in stop_keywords):
            print(f"  段落 {idx}: 遇到停止关键词 '{text[:30]}...' - 停止查找")
            break
        
        # 判断是否为单位段落
        # 支持半角和全角标点：. 。 、 ． ， 
        has_number_match = re.match(r'^\s*(\d+)[\.\。\、\．\，\s:-]+(.+)$', text)
        has_institution_keyword = bool(re.search(r'(大学|学院|研究院|研究所|中心|实验室|University|College|Institute)', text))
        
        # 方法1：以数字编号开头（如"1. 某某大学" 或 "1．某某大学"）
        if has_number_match:
            aff_number = int(has_number_match.group(1))
            aff_text = has_number_match.group(2).strip()
            doc_affiliations[aff_number] = aff_text
            affiliation_paragraphs.append({'para': para, 'text': text, 'has_number': True, 'number': aff_number, 'idx': idx})
            print(f"  段落 {idx}: 识别为单位（编号{aff_number}） - '{text[:40]}...'")
        # 方法2：第一个单位段落（可能没有编号）
        elif idx == affiliation_start_idx:
            affiliation_paragraphs.append({'para': para, 'text': text, 'has_number': False, 'number': None, 'idx': idx})
            print(f"  段落 {idx}: 识别为单位（首个段落，无编号） - '{text[:40]}...'")
        # 方法3：包含单位关键字
        elif has_institution_keyword:
            affiliation_paragraphs.append({'para': para, 'text': text, 'has_number': False, 'number': None, 'idx': idx})
            print(f"  段落 {idx}: 识别为单位（包含关键字，无编号） - '{text[:40]}...'")
        else:
            # 不是单位段落，停止查找
            print(f"  段落 {idx}: 不是单位段落 '{text[:30]}...' - 停止查找")
            break
    
    if not affiliation_paragraphs:
        report['ok'] = False
        report['messages'].append("未找到中文单位段落")
        return report
    
    print(f"  总共找到 {len(affiliation_paragraphs)} 个单位段落")
    print(f"  文档中的单位编号: {sorted(doc_affiliations.keys()) if doc_affiliations else '无'}")
    
    # 2. 检测第一个单位段落的格式
    first_affiliation = affiliation_paragraphs[0]
    print(f"检测中文单位段落: '{first_affiliation['text'][:50]}...'")
    
    expected_format = tpl.get('format_rules', {}).get('affiliation', {})
    format_result = check_paragraph_full_format(first_affiliation['para'], "中文单位", expected_format, tpl)
    report['ok'] = format_result['ok']
    report['messages'].extend(format_result['messages'])
    
    # 3. 检测单位引用完整性
    if authors_data:
        # 从作者数据中提取使用的单位编号
        used_affs = set()
        for author in authors_data:
            used_affs.update(author['affs'])
        
        # 检查作者引用的编号是否都在文档中定义
        if used_affs and doc_affiliations:
            missing_affs = [aff_id for aff_id in used_affs if aff_id not in doc_affiliations]
            if missing_affs:
                report['ok'] = False
                report['messages'].append(f"作者引用了不存在的单位编号: {missing_affs}")
                print(f"  单位引用检查: ✗ 编号{missing_affs}在单位列表中不存在")
            else:
                print(f"  单位引用检查: ✓ 所有引用的编号都存在")
    
    # 4. 检测单位编号规则
    affiliation_rules = tpl.get('check_rules', {}).get('affiliation_rules', {})
    
    # 规则：只有一个单位时，不应使用编号
    if affiliation_rules.get('check_single_affiliation', True):
        # 统计单位数量
        affiliation_count = len(affiliation_paragraphs)
        
        # 检查是否使用了编号
        has_any_number = any(aff['has_number'] for aff in affiliation_paragraphs)
        
        if affiliation_count == 1 and has_any_number:
            # 只有一个单位但使用了编号（支持全角和半角标点）
            match = re.match(r'^\s*(\d+)[\.\。\、\．\，\s:-]', first_affiliation['text'])
            if match:
                number = match.group(1)
                report['ok'] = False
                report['messages'].append(f"只有一个单位时，不应使用编号{number}（应直接写单位名称，去掉'{number}．'或'{number}. '）")
                print(f"  单位唯一性检查: ✗ 检测到编号{number}，但只有1个单位")
        elif affiliation_count == 1 and not has_any_number:
            print(f"  单位唯一性检查: ✓ 只有1个单位，未使用编号")
        elif affiliation_count > 1:
            print(f"  单位唯一性检查: ✓ 检测到{affiliation_count}个单位")
    
    return report

def check_chinese_abstract(doc, chinese_section, tpl):
    """检测中文摘要格式"""
    report = {'ok': True, 'messages': []}
    
    if chinese_section is None or chinese_section['abstract_index'] is None:
        report['ok'] = False
        report['messages'].append("未找到中文摘要")
        return report
    
    abstract_para = doc.paragraphs[chinese_section['abstract_index']]
    abstract_text = abstract_para.text.strip()
    print(f"检测中文摘要段落: '{abstract_text[:50]}...'")
    
    expected_format = tpl.get('format_rules', {}).get('abstract', {})
    
    # 1. 检查标题和间距
    label_match = re.search(r'^(摘(\s*)要)\s*[:：]?', abstract_text)
    if not label_match:
        report['ok'] = False
        report['messages'].append("未找到“摘要”标题，或标题格式不正确")
        # 如果连标题都找不到，后续检查无意义
        return report
    
    # 检查“摘要”两个字之间的空格
    space_between = label_match.group(2)
    if len(space_between) != 3:
        report['ok'] = False
        report['messages'].append(f"“摘要”二字之间应有3个空格（当前有{len(space_between)}个）")

    # 2. 逐个Run检查格式
    label_text_with_colon = label_match.group(0)
    label_end_pos = len(label_text_with_colon)
    current_text_pos = 0

    for run in abstract_para.runs:
        if not run.text.strip():
            current_text_pos += len(run.text)
            continue
            
        font_size, font_ascii, font_eastasia, is_bold, is_italic = detect_font_for_run(run, abstract_para)
        
        # 判断当前run是否属于标题部分
        is_label_part = current_text_pos < label_end_pos

        if is_label_part:
            # 标题部分应加粗
            if not is_bold:
                report['ok'] = False
                report['messages'].append(f"摘要标题“{run.text}”部分应加粗")
        else:
            # 正文部分不应加粗
            if is_bold:
                report['ok'] = False
                report['messages'].append(f"摘要正文“{run.text}”部分不应加粗")
            
            # 检查正文字体
            expected_chinese_font = expected_format.get('chinese_font', '宋体')
            if re.search(r'[\u4e00-\u9fff]', run.text) and font_eastasia != expected_chinese_font:
                msg = f"中文字体应为{expected_chinese_font}（当前：{font_eastasia}）"
                if msg not in report['messages']:
                    report['ok'] = False; report['messages'].append(msg)

            expected_english_font = expected_format.get('english_font', 'Times New Roman')
            if re.search(r'[a-zA-Z]', run.text) and font_ascii != expected_english_font:
                msg = f"英文字体应为{expected_english_font}（当前：{font_ascii}）"
                if msg not in report['messages']:
                    report['ok'] = False; report['messages'].append(msg)
            
            # 检查正文字号
            expected_size = expected_format.get('font_size_pt', 12)
            if abs(font_size - expected_size) > 0.5:
                msg = f"字号应为{get_font_size(expected_size)}({expected_size}pt)，当前为{get_font_size(font_size)}({font_size}pt)"
                if msg not in report['messages']:
                    report['ok'] = False; report['messages'].append(msg)

        current_text_pos += len(run.text)
        
    # 3. 检查段落级格式（如对齐、段间距等）
    # ... 此处可添加从 check_paragraph_full_format 移植过来的段落级检查 ...
    
    return report

def check_chinese_keywords(doc, chinese_section, tpl):
    """检测中文关键词格式和内容"""
    report = {'ok': True, 'messages': []}
    
    if chinese_section is None or chinese_section['keywords_index'] is None:
        report['ok'] = False
        report['messages'].append("未找到中文关键词")
        return report
    
    keywords_para = doc.paragraphs[chinese_section['keywords_index']]
    keywords_text = keywords_para.text.strip()
    print(f"检测中文关键词段落: '{keywords_text[:50]}...'")
    
    expected_format = tpl.get('format_rules', {}).get('keywords', {})
    
    # 1. 检查标题格式
    label_match = re.search(r'^(关键(?:词|字))\s*[:：]?', keywords_text)
    if not label_match:
        report['ok'] = False
        report['messages'].append("未找到“关键词”标题，或标题格式不正确")
        return report

    label_text_with_colon = label_match.group(0)
    label_end_pos = len(label_text_with_colon)
    current_text_pos = 0

    for run in keywords_para.runs:
        if not run.text.strip():
            current_text_pos += len(run.text)
            continue
            
        font_size, font_ascii, font_eastasia, is_bold, is_italic = detect_font_for_run(run, keywords_para)
        
        is_label_part = current_text_pos < label_end_pos

        if is_label_part:
            if not is_bold:
                report['ok'] = False
                report['messages'].append(f"关键词标题“{run.text}”部分应加粗")
        else:
            if is_bold:
                report['ok'] = False
                report['messages'].append(f"关键词内容“{run.text}”部分不应加粗")
            
            # 检查字体和字号
            expected_chinese_font = expected_format.get('chinese_font', '宋体')
            if re.search(r'[\u4e00-\u9fff]', run.text) and font_eastasia != expected_chinese_font:
                msg = f"中文字体应为{expected_chinese_font}（当前：{font_eastasia}）"
                if msg not in report['messages']:
                    report['ok'] = False; report['messages'].append(msg)

            expected_english_font = expected_format.get('english_font', 'Times New Roman')
            if re.search(r'[a-zA-Z]', run.text) and font_ascii != expected_english_font:
                msg = f"英文字体应为{expected_english_font}（当前：{font_ascii}）"
                if msg not in report['messages']:
                    report['ok'] = False; report['messages'].append(msg)

            expected_size = expected_format.get('font_size_pt', 12)
            if abs(font_size - expected_size) > 0.5:
                msg = f"字号应为{get_font_size(expected_size)}({expected_size}pt)，当前为{get_font_size(font_size)}({font_size}pt)"
                if msg not in report['messages']:
                    report['ok'] = False; report['messages'].append(msg)
                    
        current_text_pos += len(run.text)

    # 2. 检测分隔符 (保留原有逻辑)
    keywords_rules = tpl.get('check_rules', {}).get('keywords_rules', {})
    if keywords_rules.get('check_separator', True):
        match = re.search(r'关键(词|字)\s*[:：]\s*(.+)', keywords_text)
        if match:
            keywords_content = match.group(2)
            
            # 查找所有分隔符
            separators = re.findall(r'[，；,;]', keywords_content)
            
            # 期望的分隔符
            expected_separator = keywords_rules.get('separator', '；')
            
            print(f"  关键词分隔符: {separators}")
            
            # 检查是否有不符合要求的分隔符
            for sep in separators:
                if sep != expected_separator:
                    report['ok'] = False
                    report['messages'].append(f"关键词分隔符应为“{expected_separator}”（当前使用了“{sep}”）")
                    print(f"  分隔符检查: ✗ 发现错误的分隔符“{sep}”，应为“{expected_separator}”")
                    break  # 只报告第一个错误
    
    return report

def check_doc_with_template(doc_path, template_identifier):
    """
    使用指定的模板检测文档中的中文部分格式
    
    参数:
        doc_path: Word文档路径
        template_identifier: 模板标识符
    
    返回:
        检测报告字典
    """
    # 加载模板
    tpl = load_template(template_identifier)
    
    # 加载文档
    doc = Document(doc_path)
    
    # 初始化报告
    report = {
        'overall': {'ok': True, 'messages': []},
        'chinese_title': {'ok': True, 'messages': []},
        'chinese_author': {'ok': True, 'messages': []},
        'chinese_affiliation': {'ok': True, 'messages': []},
        'chinese_abstract': {'ok': True, 'messages': []},
        'chinese_keywords': {'ok': True, 'messages': []},
        'summary': {}
    }
    
    # 查找中文部分
    print("正在查找中文部分...")
    chinese_section = find_chinese_section(doc)
    
    if chinese_section is None:
        report['overall']['ok'] = False
        report['overall']['messages'].append("未找到参考文献（References）部分，无法定位中文部分")
        return report
    
    print(f"找到参考文献位置: 段落 {chinese_section['start_index']}")
    if chinese_section['title_index']:
        print(f"找到中文标题位置: 段落 {chinese_section['title_index']}")
    if chinese_section['abstract_index']:
        print(f"找到中文摘要位置: 段落 {chinese_section['abstract_index']}")
    if chinese_section['keywords_index']:
        print(f"找到中文关键词位置: 段落 {chinese_section['keywords_index']}")
    print()
    
    # 检查中文标题
    print("【检测中文标题】")
    title_result = check_chinese_title(doc, chinese_section, tpl)
    report['chinese_title'] = title_result
    if not title_result['ok']:
        report['overall']['ok'] = False
    print()
    
    # 检查中文作者
    print("【检测中文作者】")
    author_result = check_chinese_author(doc, chinese_section, tpl)
    report['chinese_author'] = author_result
    if not author_result['ok']:
        report['overall']['ok'] = False
    
    # 获取作者数据（用于单位引用检查）
    author_para = doc.paragraphs[chinese_section['author_index']] if chinese_section['author_index'] else None
    authors_data = parse_chinese_authors(author_para.text.strip()) if author_para else []
    print()
    
    # 检查中文单位
    print("【检测中文单位】")
    affiliation_result = check_chinese_affiliation(doc, chinese_section, authors_data, tpl)
    report['chinese_affiliation'] = affiliation_result
    if not affiliation_result['ok']:
        report['overall']['ok'] = False
    print()
    
    # 检查中文摘要
    print("【检测中文摘要】")
    abstract_result = check_chinese_abstract(doc, chinese_section, tpl)
    report['chinese_abstract'] = abstract_result
    if not abstract_result['ok']:
        report['overall']['ok'] = False
    print()
    
    # 检查中文关键词
    print("【检测中文关键词】")
    keywords_result = check_chinese_keywords(doc, chinese_section, tpl)
    report['chinese_keywords'] = keywords_result
    if not keywords_result['ok']:
        report['overall']['ok'] = False
    print()
    
    return report

# ---------- 报告生成 ----------
def print_report(report, tpl):
    """打印格式化的检测报告"""
    print("=" * 60)
    print("中文部分格式检测报告")
    print("=" * 60)
    print()
    
    # 总结
    print("【检查总结】")
    result_text = "通过" if report['overall']['ok'] else "发现问题"
    print(f"  中文部分格式检查结果: {result_text}")
    print()
    
    # 整体问题
    if report['overall']['messages']:
        print("【整体问题】")
        for msg in report['overall']['messages']:
            print(f"  ✗ {msg}")
        print()
    
    # 各部分检查结果
    sections = [
        ('chinese_title', '中文标题'),
        ('chinese_author', '中文作者'),
        ('chinese_affiliation', '中文单位'),
        ('chinese_abstract', '中文摘要'),
        ('chinese_keywords', '中文关键词')
    ]
    
    for section_key, section_name in sections:
        if section_key in report and report[section_key].get('messages'):
            print(f"【{section_name}检查】")
            if report[section_key]['ok']:
                print(f"  ✓ {section_name}格式正确")
            else:
                for msg in report[section_key]['messages']:
                    print(f"  ✗ {msg}")
            print()
    
    print("=" * 60)

# ---------- 命令行接口 ----------
def main():
    """命令行入口"""
    if len(sys.argv) < 3:
        print(__doc__)
        print("\n使用方法：")
        print("  python Chinese_section_detect.py check <docx文件路径> <模板标识符>")
        print("\n示例：")
        print("  python Chinese_section_detect.py check template/test.docx Chinese")
        print("  python Chinese_section_detect.py check template/test.docx templates/Chinese.json")
        sys.exit(1)
    
    command = sys.argv[1]
    
    if command == "check":
        if len(sys.argv) < 4:
            print("错误：check命令至少需要2个参数")
            sys.exit(1)
        
        doc_path = sys.argv[2]
        template_id = sys.argv[3]
        
        # 检查文件是否存在
        if not os.path.isfile(doc_path):
            print(f"错误：文件不存在: {doc_path}")
            sys.exit(1)
        
        try:
            # 执行检测
            report = check_doc_with_template(doc_path, template_id)
            
            # 加载模板用于打印
            tpl = load_template(template_id)
            
            # 打印报告
            print_report(report, tpl)
            
        except Exception as e:
            print(f"检测过程中发生错误: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)
    else:
        print(f"未知命令: {command}")
        print("支持的命令: check")
        sys.exit(1)

if __name__ == '__main__':
    main()

