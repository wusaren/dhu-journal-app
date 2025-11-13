#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

"""
=== 论文格式检测系统 - 正文内容检测器 ===

【正文内容检测 (Content Detection)】

1. 【标题层级结构检测】
   - 0 Introduction（必须有编号0，I大写，0后恰好1个空格）
   - 1 Materials and Methods（一级标题，实词首字母大写）
   - 1.1 二级标题（仅第一个单词首字母大写，其余小写）
   - 1.1.1 三级标题（格式同二级标题）
   - 验证编号连续性和层级逻辑
   - 注意：系统可以识别没有编号的 "Introduction"，但会标记为格式错误

2. 【标题大小写规则】
   - 一级标题：每个实词首字母大写（Title Case）
   - 二三级标题：仅第一个单词首字母大写（Sentence Case）
   - 虚词（and, or, of, in等）在中间位置小写

3. 【标题格式要求】
   - 一级标题：四号(14pt)，段前段后空一行，Times New Roman
   - 二三级标题：五号(10.5pt)，Times New Roman
   - 所有标题加粗，左对齐

4. 【正文格式要求】
   - 字体：Times New Roman，五号(10.5pt)
   - 行距：1.5倍行距
   - 对齐：两端对齐
   - 缩进：无特殊要求

5. 【技术特性】
   - 支持多层级标题识别和验证
   - 智能大小写规则检查
   - 完整的编号连续性验证
   - 段落样式和直接格式混合检测
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 字体检测函数 ----------
def detect_font_for_run(run, paragraph=None):
    """
    更健壮的字体检测函数，同时检查直接格式和段落样式。
    返回 (font_size_pt, font_name, is_bold, is_italic, line_spacing)
    """
    font_source = run.font if run else (paragraph.style.font if paragraph else None)
    if not font_source:
        return 12.0, "Unknown", False, False, 1.0

    font_name = font_source.name if font_source.name else "Times New Roman"

    font_size = None
    try:
        # 优先级1: 检查run的直接格式
        if run and run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        # 优先级2: 检查run的XML元素
        if font_size is None and run and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
        
        # 优先级3: 检查段落样式的直接格式
        if font_size is None and paragraph and paragraph.style and hasattr(paragraph.style.font, 'size') and hasattr(paragraph.style.font.size, 'pt'):
            font_size = float(paragraph.style.font.size.pt)
            
        # 优先级4: 检查段落样式的XML元素
        if font_size is None and paragraph and hasattr(paragraph.style, 'element'):
            sz_nodes = paragraph.style.element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass

    font_size = font_size if font_size is not None else 12.0

    is_bold = font_source.bold if font_source.bold is not None else False
    is_italic = font_source.italic if font_source.italic is not None else False
    if run and run.font:
        is_bold = run.font.bold if run.font.bold is not None else is_bold
        is_italic = run.font.italic if run.font.italic is not None else is_italic
    
    # 确保返回明确的布尔值，而不是None
    is_bold = bool(is_bold) if is_bold is not None else False
    is_italic = bool(is_italic) if is_italic is not None else False

    # 行间距检测
    line_spacing = 1.0  # 默认单倍行距
    try:
        if paragraph and paragraph.paragraph_format.line_spacing:
            line_spacing = float(paragraph.paragraph_format.line_spacing)
    except Exception:
        pass

    return font_size, font_name, is_bold, is_italic, line_spacing

def get_font_size(pt_size, tpl=None):
    """字体大小转换为中文字号"""
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        size_map = {
            9: "小五", 10.5: "五号", 12: "小四", 14: "四号",
            16: "三号", 18: "小二", 22: "二号", 24: "小一", 26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

def get_line_spacing_name(spacing, tpl=None):
    """行间距转换为中文描述"""
    if tpl and 'check_rules' in tpl and 'line_spacing_mapping' in tpl['check_rules']:
        spacing_config = tpl['check_rules']['line_spacing_mapping']
        spacing_map = {}
        for key, value in spacing_config.items():
            try:
                spacing_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        spacing_map = {
            1.0: "单倍行距", 1.15: "1.15倍行距", 
            1.5: "1.5倍行距", 2.0: "双倍行距"
        }
    
    if not spacing_map:
        return f"{spacing}倍行距"
    
    closest_spacing = min(spacing_map.keys(), key=lambda x: abs(x - spacing))
    return spacing_map[closest_spacing]

def get_alignment_name(alignment_value, tpl=None):
    """段落对齐方式转换为中文描述"""
    alignment_map = {
        0: "左对齐", 1: "居中对齐", 2: "右对齐", 3: "两端对齐"
    }
    return alignment_map.get(alignment_value, f"未知对齐({alignment_value})")

def detect_paragraph_alignment(paragraph):
    """段落对齐检测"""
    # 优先级1: 检查直接格式
    direct_alignment = paragraph.paragraph_format.alignment
    if direct_alignment is not None:
        return int(direct_alignment)
    
    # 优先级2: 如果直接格式为None，检查样式格式
    if paragraph.style:
        try:
            style_alignment = paragraph.style.paragraph_format.alignment
            if style_alignment is not None:
                return int(style_alignment)
        except Exception:
            pass
    
    # 默认值: 左对齐
    return WD_PARAGRAPH_ALIGNMENT.LEFT

def detect_paragraph_indent(paragraph):
    """检测段落缩进"""
    try:
        fmt = paragraph.paragraph_format
        first_line_indent = fmt.first_line_indent.pt if fmt.first_line_indent else 0.0
        left_indent = fmt.left_indent.pt if fmt.left_indent else 0.0
        right_indent = fmt.right_indent.pt if fmt.right_indent else 0.0
        return first_line_indent, left_indent, right_indent
    except Exception:
        return 0.0, 0.0, 0.0

# ---------- 正文检测逻辑 ----------

def get_paragraph_numbering_info(paragraph):
    """
    获取段落的 Word 自动编号信息
    返回: (是否有编号, 编号级别)
    """
    try:
        if paragraph._p.pPr is None:
            return False, None
        
        numPr = paragraph._p.pPr.find(qn('w:numPr'))
        if numPr is None:
            return False, None
        
        # 获取编号级别
        ilvl = numPr.find(qn('w:ilvl'))
        level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0
        
        # 获取编号ID
        numId = numPr.find(qn('w:numId'))
        num_id = numId.get(qn('w:val')) if numId is not None else None
        
        if num_id is None or num_id == '0':
            return False, None
        
        return True, level
    except:
        return False, None

def identify_title_hierarchy(doc, tpl):
    """
    识别文档中的标题层级结构
    返回 {'ok': bool, 'messages': [], 'titles': [{'level': int, 'number': str, 'text': str, 'paragraph': obj}]}
    """
    report = {'ok': True, 'messages': [], 'titles': []}
    
    # 获取结构规则
    structure_rules = tpl.get('structure_rules', {})
    level0_pattern = structure_rules.get('level0_pattern', r'^\\s*0\\s+Introduction\\s*$')
    level1_pattern = structure_rules.get('level1_pattern', r'^\\s*(\\d+)\\s+(.+)$')
    level2_pattern = structure_rules.get('level2_pattern', r'^\\s*(\\d+\\.\\d+)\\s+(.+)$')
    level3_pattern = structure_rules.get('level3_pattern', r'^\\s*(\\d+\\.\\d+\\.\\d+)\\s+(.+)$')
    
    found_introduction = False
    title_sequence = []
    
    # 扫描所有段落，识别标题
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text or not paragraph.text.strip():
            continue
            
        text = paragraph.text.strip()
        
        # 检查段落是否有 Word 自动编号
        has_auto_num, auto_num_level = get_paragraph_numbering_info(paragraph)
        
        # 检查0 Introduction（带编号的格式）
        if re.match(level0_pattern, text, re.IGNORECASE):
            found_introduction = True
            
            report['titles'].append({
                'level': 0,
                'number': '0',
                'text': 'Introduction',
                'paragraph': paragraph,
                'paragraph_index': para_idx,
                'full_text': text,
                'has_number': True,
                'has_auto_numbering': False
            })
            title_sequence.append('0')
            print(f"找到Introduction: {text} (索引: {para_idx})")
        
        # 检查单独的 Introduction（没有编号的格式，但可能有 Word 自动编号）
        elif re.match(r'^\s*Introduction\s*$', text, re.IGNORECASE) and not found_introduction:
            found_introduction = True
            
            # 检查是否使用了 Word 自动编号
            if has_auto_num and auto_num_level == 0:
                # 有 Word 自动编号，视为有编号（但格式仍不正确，应该是文本形式的 "0 Introduction"）
                report['titles'].append({
                    'level': 0,
                    'number': '[自动编号]',
                    'text': 'Introduction',
                    'paragraph': paragraph,
                    'paragraph_index': para_idx,
                    'full_text': text,
                    'has_number': False,  # 虽然有自动编号，但格式不对
                    'has_auto_numbering': True
                })
                print(f"找到Introduction: {text} (索引: {para_idx}, 使用Word自动编号)")
            else:
                # 完全没有编号
                report['titles'].append({
                    'level': 0,
                    'number': '',
                    'text': 'Introduction',
                    'paragraph': paragraph,
                    'paragraph_index': para_idx,
                    'full_text': text,
                    'has_number': False,
                    'has_auto_numbering': False
                })
                print(f"找到Introduction: {text} (索引: {para_idx}, 无编号格式)")
            title_sequence.append('0')
        
        # 检查是否是空格数量错误的Introduction
        elif re.match(r'^\s*0\s+Introduction\s*$', text, re.IGNORECASE):
            print(f"发现Introduction但空格数量不正确: '{text}'")
            # 分析空格数量
            space_match = re.search(r'0(\s+)Introduction', text)
            if space_match:
                spaces = space_match.group(1)
                space_count = len(spaces)
                print(f"0后有{space_count}个空格，应为1个空格")
        
        # 检查一级标题
        elif re.match(level1_pattern, text):
            match = re.match(level1_pattern, text)
            number = match.group(1)
            title_text = match.group(2).strip()
            report['titles'].append({
                'level': 1,
                'number': number,
                'text': title_text,
                'paragraph': paragraph,
                'paragraph_index': para_idx,
                'full_text': text
            })
            title_sequence.append(number)
            print(f"找到一级标题: {number} {title_text}")
        
        # 检查是否是空格数量错误的一级标题
        elif re.match(r'^\s*(\d+)\s+(.+)$', text) and not re.match(level1_pattern, text):
            space_match = re.search(r'(\d+)(\s+)(.+)', text)
            if space_match:
                number = space_match.group(1)
                spaces = space_match.group(2)
                space_count = len(spaces)
                title_text = space_match.group(3)
                print(f"发现一级标题但空格数量不正确: '{number}' 后有{space_count}个空格，应为1个空格")
        
        # 检查二级标题
        elif re.match(level2_pattern, text):
            match = re.match(level2_pattern, text)
            number = match.group(1)
            title_text = match.group(2).strip()
            report['titles'].append({
                'level': 2,
                'number': number,
                'text': title_text,
                'paragraph': paragraph,
                'full_text': text
            })
            title_sequence.append(number)
            print(f"找到二级标题: {number} {title_text}")
        
        # 检查三级标题
        elif re.match(level3_pattern, text):
            match = re.match(level3_pattern, text)
            number = match.group(1)
            title_text = match.group(2).strip()
            report['titles'].append({
                'level': 3,
                'number': number,
                'text': title_text,
                'paragraph': paragraph,
                'full_text': text
            })
            title_sequence.append(number)
            print(f"找到三级标题: {number} {title_text}")
    
    # 检测疑似标题但缺少编号的段落（基于格式特征）
    # 只检测 Introduction 之后的内容，避免误报 Title、Authors、Keywords 等
    missing_number_titles = []
    introduction_index = None
    
    # 先找到 Introduction 的位置
    for title in report['titles']:
        if title['level'] == 0:
            introduction_index = title.get('paragraph_index')
            break
    
    # 只在 Introduction 之后检测
    if introduction_index is not None:
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # 只检测 Introduction 之后的段落
            if para_idx <= introduction_index:
                continue
                
            if not paragraph.text or not paragraph.text.strip():
                continue
            
            text = paragraph.text.strip()
            
            # 跳过已识别的标题
            already_identified = False
            for title in report['titles']:
                if title.get('paragraph_index') == para_idx:
                    already_identified = True
                    break
            if already_identified:
                continue
            
            # 检查是否有 Word 自动编号
            has_auto_num, auto_num_level = get_paragraph_numbering_info(paragraph)
            
            # 如果有 Word 自动编号且级别为 0（一级编号）
            if has_auto_num and auto_num_level == 0:
                # 排除图表、参考文献等
                text_lower = text.lower()
                # 判断是否是参考文献：通常包含作者名、期刊名、年份等特征
                is_reference = (
                    text_lower.startswith('reference') or
                    # 常见作者姓氏开头
                    re.match(r'^[A-Z]{2,}\s+[A-Z]', text) or  # "ZHU R X, LI J" 格式
                    # 包含期刊/会议标记
                    '[J]' in text or '[D]' in text or '[C]' in text or '[M]' in text or
                    # 标准文献
                    text_lower.startswith('gb/t') or text_lower.startswith('iso')
                )
                
                # 判断是否是图表标题：通常包含 diagram, layout, chart 等词
                is_figure_caption = (
                    'diagram' in text_lower or
                    'layout' in text_lower or
                    'chart' in text_lower or
                    'schematic' in text_lower or
                    'physical' in text_lower or
                    text_lower.endswith('values') or
                    'measurement values' in text_lower or
                    'relation between' in text_lower
                )
                
                if not (text_lower.startswith('table ') or 
                       text_lower.startswith('figure ') or 
                       text_lower.startswith('fig.') or
                       text_lower.startswith('abstract') or
                       text_lower.startswith('keywords') or
                       text_lower.startswith('clc') or
                       is_reference or
                       is_figure_caption or
                       re.match(r'^\d+[\-\—]', text) or  # 排除类似 "1-laptop" 这种
                       re.match(r'^\(\d+\)', text)):  # 排除公式编号
                    missing_number_titles.append({
                        'paragraph_index': para_idx,
                        'text': text,
                        'font_size': None,
                        'has_word_numbering': True
                    })
                    print(f"警告: 发现使用Word自动编号的标题 (索引 {para_idx}): '{text[:60]}...'")
                    
                    # 也将此标题添加到 report['titles'] 中，以便进行格式检查
                    # 标记为一级标题（level=1），因为使用了Word一级编号
                    report['titles'].append({
                        'level': 1,
                        'number': '[Word自动编号]',
                        'text': text,
                        'paragraph': paragraph,
                        'paragraph_index': para_idx,
                        'full_text': text,
                        'has_number': False,
                        'has_auto_numbering': True  # 标记使用了Word自动编号
                    })
            # 否则，检测格式特征：加粗 + 较大字体 + 长度适中
            elif paragraph.runs and len(text) > 10 and len(text) < 150:
                first_run = paragraph.runs[0]
                is_bold = first_run.font.bold if first_run.font.bold is not None else False
                
                # 获取字体大小
                font_size = None
                try:
                    if first_run.font.size:
                        font_size = first_run.font.size.pt
                except:
                    pass
                
                # 如果是加粗且字体不是很小（可能是标题）
                if is_bold and (font_size is None or font_size >= 10.5):
                    # 排除一些明显不是标题的内容
                    text_lower = text.lower()
                    if not (text_lower.startswith('table ') or 
                           text_lower.startswith('figure ') or 
                           text_lower.startswith('fig.') or
                           text_lower.startswith('abstract') or
                           text_lower.startswith('keywords') or
                           text_lower.startswith('clc') or
                           re.match(r'^\d+[\-\—]', text) or  # 排除类似 "1-laptop" 这种
                           re.match(r'^\(\d+\)', text)):  # 排除公式编号
                        missing_number_titles.append({
                            'paragraph_index': para_idx,
                            'text': text,
                            'font_size': font_size,
                            'has_word_numbering': False
                        })
                        print(f"警告: 发现疑似标题但缺少编号的段落 (索引 {para_idx}): '{text[:60]}...'")
    
    # 如果发现缺少编号的标题，添加到错误消息
    if missing_number_titles:
        report['ok'] = False
        
        # 统计使用 Word 自动编号的标题数量
        word_numbered_count = sum(1 for item in missing_number_titles if item.get('has_word_numbering', False))
        no_number_count = len(missing_number_titles) - word_numbered_count
        
        if word_numbered_count > 0:
            report['messages'].append(f"发现 {word_numbered_count} 个使用Word自动编号的标题，应改为文本形式的编号（如 '1 Title'，而非使用Word编号库）")
            for item in [x for x in missing_number_titles if x.get('has_word_numbering', False)][:5]:
                report['messages'].append(f"  - 段落 {item['paragraph_index']}: '{item['text'][:60]}...'")
        
        if no_number_count > 0:
            report['messages'].append(f"发现 {no_number_count} 个疑似标题但完全缺少编号，标题应有编号格式（如 '1 Title', '1.1 Subtitle'）")
            for item in [x for x in missing_number_titles if not x.get('has_word_numbering', False)][:5]:
                report['messages'].append(f"  - 段落 {item['paragraph_index']}: '{item['text'][:60]}...'")
    
    # 验证Introduction
    if not found_introduction:
        report['ok'] = False
        error_msg = tpl.get('messages', {}).get('structure_introduction_error')
        if error_msg:
            report['messages'].append(error_msg)
    else:
        # 检查Introduction是否有编号
        intro_has_number = False
        intro_has_auto_numbering = False
        for title in report['titles']:
            if title['level'] == 0:
                intro_has_number = title.get('has_number', False)
                intro_has_auto_numbering = title.get('has_auto_numbering', False)
                break
        
        if intro_has_number:
            ok_msg = tpl.get('messages', {}).get('structure_introduction_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
        else:
            # Introduction没有正确的编号
            report['ok'] = False
            if intro_has_auto_numbering:
                error_msg = "Introduction标题使用了Word自动编号，应改为文本形式 '0 Introduction'（0后恰好1个空格）"
            else:
                error_msg = "Introduction标题缺少编号，应为 '0 Introduction'（0后恰好1个空格）"
            report['messages'].append(error_msg)
            print(f"警告: {error_msg}")
    
    # 验证编号连续性
    if len(report['titles']) > 1:
        numbering_ok = validate_title_numbering(report['titles'], tpl)
        if numbering_ok:
            ok_msg = tpl.get('messages', {}).get('structure_numbering_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
        else:
            report['ok'] = False
            error_msg = tpl.get('messages', {}).get('structure_numbering_error')
            if error_msg:
                report['messages'].append(error_msg)
    
    print(f"总共找到 {len(report['titles'])} 个标题")
    return report

def validate_title_numbering(titles, tpl):
    """
    验证标题编号的连续性和逻辑性
    """
    try:
        # 按层级分组
        level1_numbers = []
        level2_numbers = []
        level3_numbers = []
        
        for title in titles:
            if title['level'] == 1:
                level1_numbers.append(int(title['number']))
            elif title['level'] == 2:
                level2_numbers.append(title['number'])
            elif title['level'] == 3:
                level3_numbers.append(title['number'])
        
        # 检查一级标题连续性（应从1开始）
        if level1_numbers:
            level1_numbers.sort()
            expected = list(range(1, len(level1_numbers) + 1))
            if level1_numbers != expected:
                print(f"一级标题编号不连续: {level1_numbers} vs 期望: {expected}")
                return False
        
        # 检查二级标题格式（应为x.y格式）
        for level2_num in level2_numbers:
            if not re.match(r'^\d+\.\d+$', level2_num):
                print(f"二级标题编号格式错误: {level2_num}")
                return False
        
        # 检查三级标题格式（应为x.y.z格式）
        for level3_num in level3_numbers:
            if not re.match(r'^\d+\.\d+\.\d+$', level3_num):
                print(f"三级标题编号格式错误: {level3_num}")
                return False
        
        return True
    except Exception as e:
        print(f"编号验证出错: {e}")
        return False

def check_title_format(titles, tpl):
    """
    检查标题格式（字体、字号、加粗、行距等）
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not titles:
        report['ok'] = False
        report['messages'].append("没有标题可供格式检查")
        return report
    
    # 获取格式规则
    format_rules = tpl.get('format_rules', {})
    issues = []
    
    for title_info in titles:
        paragraph = title_info['paragraph']
        level = title_info['level']
        title_text = title_info['text']
        has_auto_numbering = title_info.get('has_auto_numbering', False)
        
        # 为使用Word自动编号的标题添加前缀
        title_prefix = "[使用Word自动编号] " if has_auto_numbering else ""
        
        if not paragraph.runs:
            continue
        
        # 根据标题级别选择格式规则
        if level == 0 or level == 1:
            rule_key = 'level1_title'
        elif level == 2:
            rule_key = 'level2_title'
        elif level == 3:
            rule_key = 'level3_title'
        else:
            continue
        
        title_rules = format_rules.get(rule_key, {})
        if not title_rules:
            continue
        
        # 检查第一个run的格式
        main_run = paragraph.runs[0]
        actual_size_pt, actual_font_name, actual_bold, actual_italic, actual_line_spacing = detect_font_for_run(main_run, paragraph)
        
        # 字体大小检查
        if 'font_size_pt' in title_rules:
            expected_size_pt = float(title_rules['font_size_pt'])
            actual_size_name = get_font_size(actual_size_pt, tpl)
            expected_size_name = get_font_size(expected_size_pt, tpl)
            print(f"{title_prefix}标题 '{title_text}' 字体大小: {actual_size_name}（{actual_size_pt}pt）(期望: {expected_size_name}（{expected_size_pt}pt）)")
            if abs(actual_size_pt - expected_size_pt) > 0.5:
                issues.append(f"{title_prefix}标题 '{title_text}' 字体大小应为{expected_size_name}（{expected_size_pt}pt），实际为{actual_size_name}（{actual_size_pt}pt）")
        
        # 字体名称检查
        if 'font_name' in title_rules:
            expected_font_name = str(title_rules['font_name'])
            print(f"{title_prefix}标题 '{title_text}' 字体名称: {actual_font_name} (期望: {expected_font_name})")
            if expected_font_name.lower() not in actual_font_name.lower():
                issues.append(f"{title_prefix}标题 '{title_text}' 字体应为{expected_font_name}，实际为{actual_font_name}")
        
        # 加粗检查
        if 'bold' in title_rules:
            expected_bold = bool(title_rules['bold'])
            print(f"{title_prefix}标题 '{title_text}' 加粗: {'是' if actual_bold else '否'} (期望: {'是' if expected_bold else '否'})")
            if actual_bold != expected_bold:
                bold_status = "加粗" if expected_bold else "不加粗"
                actual_status = "加粗" if actual_bold else "不加粗"
                issues.append(f"{title_prefix}标题 '{title_text}' 应为{bold_status}，实际为{actual_status}")
        
        # 斜体检查
        if 'italic' in title_rules:
            expected_italic = bool(title_rules['italic'])
            print(f"{title_prefix}标题 '{title_text}' 斜体: {'是' if actual_italic else '否'} (期望: {'是' if expected_italic else '否'})")
            if actual_italic != expected_italic:
                italic_status = "斜体" if expected_italic else "正体"
                actual_status = "斜体" if actual_italic else "正体"
                issues.append(f"{title_prefix}标题 '{title_text}' 应为{italic_status}，实际为{actual_status}")
        
        # 段前段后间距检查（所有标题级别）
        if 'space_before' in title_rules:
            expected_space_before = float(title_rules['space_before'])
            expected_lines = expected_space_before / 12.0
            actual_space_before = paragraph.paragraph_format.space_before
            actual_lines = actual_space_before.pt / 12.0 if actual_space_before and actual_space_before.pt else 0.0
            print(f"{title_prefix}标题 '{title_text}' 段前间距: {actual_lines:.1f}行 (期望: {expected_lines:.1f}行)")
            
            # 特殊处理：1行可能显示为1.3行
            if expected_lines == 1.0 and 1.0 <= actual_lines <= 1.35:
                pass  # 认为是正确的
            elif abs(actual_lines - expected_lines) > 0.2:
                issues.append(f"{title_prefix}标题 '{title_text}' 段前间距应为{expected_lines:.1f}行，实际为{actual_lines:.1f}行")
        
        if 'space_after' in title_rules:
            expected_space_after = float(title_rules['space_after'])
            expected_lines = expected_space_after / 12.0
            actual_space_after = paragraph.paragraph_format.space_after
            actual_lines = actual_space_after.pt / 12.0 if actual_space_after and actual_space_after.pt else 0.0
            print(f"{title_prefix}标题 '{title_text}' 段后间距: {actual_lines:.1f}行 (期望: {expected_lines:.1f}行)")
            
            # 特殊处理：1行可能显示为1.3行
            if expected_lines == 1.0 and 1.0 <= actual_lines <= 1.35:
                pass  # 认为是正确的
            elif abs(actual_lines - expected_lines) > 0.2:
                issues.append(f"{title_prefix}标题 '{title_text}' 段后间距应为{expected_lines:.1f}行，实际为{actual_lines:.1f}行")
    
    print(f"发现 {len(issues)} 个标题格式问题")
    
    if issues:
        report['ok'] = False
        report['messages'].append("标题格式问题：")
        report['messages'].extend([f"  - {i}" for i in issues])
    else:
        if level == 0 or level == 1:
            ok_msg = tpl.get('messages', {}).get('format_level1_ok')
        else:
            ok_msg = tpl.get('messages', {}).get('format_level23_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
    
    return report

def check_title_case(titles, tpl):
    """
    检查标题大小写规则
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not titles:
        return report
    
    # 获取大小写规则
    case_rules = tpl.get('title_case_rules', {})
    minor_words = case_rules.get('minor_words', ['and', 'or', 'of', 'in', 'on', 'at', 'to', 'for', 'with', 'by', 'from', 'the', 'a', 'an'])
    
    issues = []
    
    for title_info in titles:
        level = title_info['level']
        title_text = title_info['text']
        has_auto_numbering = title_info.get('has_auto_numbering', False)
        
        # 为使用Word自动编号的标题添加前缀
        title_prefix = "[使用Word自动编号] " if has_auto_numbering else ""
        
        if level == 0:  # Introduction特殊处理
            if title_text != 'Introduction':
                issues.append(f"{title_prefix}Introduction标题应为'Introduction'，实际为'{title_text}'")
        elif level == 1:  # 一级标题：实词首字母大写
            corrected = apply_title_case(title_text, minor_words)
            if title_text != corrected:
                issues.append(f"{title_prefix}一级标题 '{title_text}' 大小写不正确，应为 '{corrected}'")
        elif level in [2, 3]:  # 二三级标题：仅首词大写
            corrected = apply_sentence_case(title_text)
            if title_text != corrected:
                issues.append(f"{title_prefix}{'二' if level == 2 else '三'}级标题 '{title_text}' 大小写不正确，应为 '{corrected}'")
    
    if issues:
        report['ok'] = False
        if any('一级标题' in issue for issue in issues):
            error_msg = tpl.get('messages', {}).get('title_case_level1_error')
            if error_msg:
                report['messages'].append(error_msg)
        if any('二级标题' in issue or '三级标题' in issue for issue in issues):
            error_msg = tpl.get('messages', {}).get('title_case_level23_error')
            if error_msg:
                report['messages'].append(error_msg)
        report['messages'].extend([f"  - {i}" for i in issues])
    else:
        ok_msg = tpl.get('messages', {}).get('title_case_level1_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
        ok_msg = tpl.get('messages', {}).get('title_case_level23_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
    
    return report

def apply_title_case(title, minor_words):
    """应用Title Case规则（一级标题）"""
    words = title.split()
    result = []
    
    for i, word in enumerate(words):
        if i == 0 or i == len(words) - 1:  # 首词和末词总是大写
            result.append(word.capitalize())
        elif word.lower() in minor_words:  # 虚词小写
            result.append(word.lower())
        else:  # 实词大写
            result.append(word.capitalize())
    
    return ' '.join(result)

def apply_sentence_case(title):
    """应用Sentence Case规则（二三级标题）"""
    if not title:
        return title
    return title[0].upper() + title[1:].lower()

def check_content_text_format(doc, titles, tpl):
    """
    检查正文内容格式（非标题段落的格式）
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    # 获取格式规则
    format_rules = tpl.get('format_rules', {}).get('content_text', {})
    if not format_rules:
        report['messages'].append("没有正文格式规则可供检查")
        return report
    
    # 找到Introduction段落的索引
    introduction_index = None
    title_paragraph_indices = set()
    
    if titles:
        for title_info in titles:
            # 使用预存的段落索引
            if 'paragraph_index' in title_info:
                idx = title_info['paragraph_index']
                title_paragraph_indices.add(idx)
                # 找到Introduction段落索引
                if title_info['level'] == 0 and title_info['text'] == 'Introduction':
                    introduction_index = idx
    
    # 查找正文段落（只检测Introduction之后、References之前的内容）
    content_paragraphs = []
    consecutive_references = 0  # 连续的参考文献段落计数
    
    if introduction_index is not None:
        # 从Introduction段落的下一个段落开始检查
        for i in range(introduction_index + 1, len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            
            # 排除标题段落
            if i in title_paragraph_indices:
                continue
            
            # 检查是否为有效正文段落
            text = paragraph.text.strip()
            if not text or len(text) <= 20:
                consecutive_references = 0  # 重置计数
                continue
            
            # 检测是否为参考文献
            text_lower = text.lower()
            is_reference = (
                # 以作者姓名开头（全大写字母 + 空格）
                re.match(r'^[A-Z]{2,}\s+[A-Z]', text) or
                # 包含文献类型标记
                '[J]' in text or '[D]' in text or '[C]' in text or '[M]' in text or
                # 标准文献编号
                text_lower.startswith('gb/t') or text_lower.startswith('iso') or
                # References标题
                text_lower.startswith('reference')
            )
            
            if is_reference:
                consecutive_references += 1
                # 连续出现3个参考文献，认为已经进入参考文献部分，停止收集正文
                if consecutive_references >= 3:
                    print(f"检测到参考文献部分（从段落 {i-2} 开始），停止收集正文段落")
                    break
                continue
            else:
                consecutive_references = 0  # 重置计数
            
            # 检查是否有 Word 自动编号（可能是标题）
            has_auto_num, auto_num_level = get_paragraph_numbering_info(paragraph)
            if has_auto_num and auto_num_level == 0:
                # 有一级自动编号，很可能是标题，排除
                # 但要排除参考文献等非标题内容
                text_lower = text.lower()
                is_likely_title = not (
                    text_lower.startswith('reference') or
                    re.match(r'^[A-Z]{2,}\s+[A-Z]', text) or  # 参考文献格式
                    '[J]' in text or '[D]' in text or '[C]' in text or '[M]' in text or
                    text_lower.startswith('gb/t') or text_lower.startswith('iso')
                )
                if is_likely_title:
                    continue  # 排除这个疑似标题的段落
            
            # 排除表格标题、图片标题等
            # 表格标题通常以"Table"、"表"开头，或者包含居中对齐的数字标题
            text_lower = text.lower()
            if (text_lower.startswith('table ') or 
                text_lower.startswith('figure ') or 
                text_lower.startswith('fig.') or
                text_lower.startswith('图 ') or 
                text_lower.startswith('表 ') or
                re.match(r'^(表|图|Table|Figure)\s*\d+', text, re.IGNORECASE)):
                continue
            
            # 排除居中对齐的短段落（可能是标题或表格标题）
            alignment = paragraph.alignment
            if alignment is not None:
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and len(text) < 80:
                    continue
            
            # 排除加粗的短段落（很可能是标题但没有编号）
            if paragraph.runs and len(text) < 150:
                first_run = paragraph.runs[0]
                is_bold = first_run.font.bold if first_run.font.bold is not None else False
                if is_bold:
                    # 加粗的段落，检查字体大小
                    try:
                        font_size = first_run.font.size.pt if first_run.font.size else 10.5
                        # 如果字体大于正文（>10.5pt）且加粗，很可能是标题
                        if font_size > 10.5:
                            continue
                    except:
                        pass
            
            # 这是有效的正文段落
            content_paragraphs.append(paragraph)
    
    print(f"找到 {len(content_paragraphs)} 个正文段落")
    
    if not content_paragraphs:
        report['ok'] = False
        report['messages'].append("未找到正文内容段落")
        return report
    
    issues = []
    paragraphs_with_issues = []
    
    # 检查所有正文段落的格式
    for i, paragraph in enumerate(content_paragraphs):
        if not paragraph.runs:
            continue
        
        # 检查第一个run的格式
        main_run = paragraph.runs[0]
        actual_size_pt, actual_font_name, actual_bold, actual_italic, actual_line_spacing = detect_font_for_run(main_run, paragraph)
        
        paragraph_preview = paragraph.text[:40] + "..." if len(paragraph.text) > 40 else paragraph.text
        paragraph_issues = []
        
        # 字体大小检查
        if 'font_size_pt' in format_rules:
            expected_size_pt = float(format_rules['font_size_pt'])
            if abs(actual_size_pt - expected_size_pt) > 0.5:
                actual_size_name = get_font_size(actual_size_pt, tpl)
                expected_size_name = get_font_size(expected_size_pt, tpl)
                paragraph_issues.append(f"字体大小应为{expected_size_name}（{expected_size_pt}pt），实际为{actual_size_name}（{actual_size_pt}pt）")
        
        # 字体名称检查
        if 'font_name' in format_rules:
            expected_font_name = str(format_rules['font_name'])
            if expected_font_name.lower() not in actual_font_name.lower():
                paragraph_issues.append(f"字体应为{expected_font_name}，实际为{actual_font_name}")
        
        # 加粗检查
        if 'bold' in format_rules:
            expected_bold = bool(format_rules['bold'])
            if actual_bold != expected_bold:
                bold_status = "加粗" if expected_bold else "不加粗"
                actual_status = "加粗" if actual_bold else "不加粗"
                paragraph_issues.append(f"应为{bold_status}，实际为{actual_status}")
        
        # 斜体检查
        if 'italic' in format_rules:
            expected_italic = bool(format_rules['italic'])
            if actual_italic != expected_italic:
                italic_status = "斜体" if expected_italic else "正体"
                actual_status = "斜体" if actual_italic else "正体"
                paragraph_issues.append(f"应为{italic_status}，实际为{actual_status}")
        
        # 行间距检查
        if 'line_spacing' in format_rules:
            expected_line_spacing = float(format_rules['line_spacing'])
            if abs(actual_line_spacing - expected_line_spacing) > 0.1:
                actual_spacing_name = get_line_spacing_name(actual_line_spacing, tpl)
                expected_spacing_name = get_line_spacing_name(expected_line_spacing, tpl)
                paragraph_issues.append(f"行间距应为{expected_spacing_name}（{expected_line_spacing}倍），实际为{actual_spacing_name}（{actual_line_spacing}倍）")
        
        # 对齐方式检查
        if 'alignment' in format_rules:
            expected_alignment_str = str(format_rules['alignment'])
            alignment_map = {"left": 0, "center": 1, "right": 2, "justify": 3}
            expected_alignment = alignment_map.get(expected_alignment_str, 0)
            actual_alignment = detect_paragraph_alignment(paragraph)
            if actual_alignment != expected_alignment:
                actual_alignment_name = get_alignment_name(actual_alignment, tpl)
                expected_alignment_name = get_alignment_name(expected_alignment, tpl)
                paragraph_issues.append(f"对齐方式应为{expected_alignment_name}，实际为{actual_alignment_name}")
        
        # 首行缩进检查
        if 'first_line_indent' in format_rules:
            expected_first_indent = float(format_rules['first_line_indent'])
            first_line_indent, left_indent, right_indent = detect_paragraph_indent(paragraph)
            if abs(first_line_indent - expected_first_indent) > 2.0:  # 2pt容差
                paragraph_issues.append(f"首行缩进应为{expected_first_indent}pt（约2字符），实际为{first_line_indent:.1f}pt")
        
        # 如果这个段落有问题，记录下来
        if paragraph_issues:
            paragraphs_with_issues.append({
                'index': i + 1,
                'preview': paragraph_preview,
                'issues': paragraph_issues
            })
            issues.extend([f"正文段落 {i+1} {issue}" for issue in paragraph_issues])
    
    # 输出有问题的段落（简洁格式）
    print(f"\n=== 正文段落格式检查结果 ===")
    if paragraphs_with_issues:
        print(f"发现 {len(paragraphs_with_issues)} 个段落有格式问题（共检查 {len(content_paragraphs)} 个段落）")
        print("\n有问题的段落（简洁显示）：")
        for p in paragraphs_with_issues[:10]:  # 控制台只显示前10个
            print(f"  段落 {p['index']} '{p['preview']}'")
            print(f"    问题: {'; '.join(p['issues'])}")
        
        if len(paragraphs_with_issues) > 10:
            print(f"  ... 还有 {len(paragraphs_with_issues) - 10} 个段落有问题")
    else:
        print("✓ 所有正文段落格式正确")
    
    print(f"\n总计 {len(issues)} 个正文格式问题")
    
    if issues:
        report['ok'] = False
        header = tpl.get('messages', {}).get('format_content_error')
        if header:
            report['messages'].append(header + "：")
        report['messages'].extend([f"  - {i}" for i in issues])
    else:
        ok_msg = tpl.get('messages', {}).get('format_content_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
    
    # 保存详细结果到字典中供后续使用
    report['paragraphs_with_issues'] = paragraphs_with_issues
    report['total_paragraphs'] = len(content_paragraphs)
    
    return report

# ---------- 主检测函数 ----------
def check_content_with_template(doc_path, template_identifier):
    """
    主检查函数：检查正文内容格式
    """
    tpl = load_template(template_identifier)
    doc = Document(doc_path)
    
    # 执行标题层级检查
    hierarchy_report = identify_title_hierarchy(doc, tpl)
    
    # 执行标题格式检查
    if hierarchy_report['titles']:
        format_report = check_title_format(hierarchy_report['titles'], tpl)
        case_report = check_title_case(hierarchy_report['titles'], tpl)
    else:
        format_report = {'ok': False, 'messages': ['没有标题可供格式检查']}
        case_report = {'ok': False, 'messages': ['没有标题可供大小写检查']}
    
    # 执行正文内容格式检查
    content_format_report = check_content_text_format(doc, hierarchy_report['titles'], tpl)
    
    # 组装报告（添加titles信息用于批注定位）
    report = {
        'hierarchy': hierarchy_report,
        'format': format_report,
        'case': case_report,
        'content_format': content_format_report,
        'summary': [],
        'titles': hierarchy_report.get('titles', [])  # 传递标题信息用于批注定位
    }
    
    # 生成总结
    all_ok = (hierarchy_report['ok'] and format_report['ok'] and case_report['ok'] and content_format_report['ok'])
    summary_tpl = tpl.get('messages', {}).get('summary_overall')
    if summary_tpl:
        try:
            report['summary'].append(summary_tpl.format(ok=all_ok))
        except Exception:
            report['summary'].append(str(summary_tpl))
    
    # 保存详细的正文段落检查结果到文件
    if 'paragraphs_with_issues' in content_format_report and content_format_report['paragraphs_with_issues']:
        try:
            import os
            output_dir = os.path.dirname(doc_path)
            base_name = os.path.splitext(os.path.basename(doc_path))[0]
            detail_file = os.path.join(output_dir, f"{base_name}_content_details.txt")
            
            with open(detail_file, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("正文段落格式检查详细报告\n")
                f.write("=" * 80 + "\n\n")
                f.write(f"文档: {os.path.basename(doc_path)}\n")
                f.write(f"总段落数: {content_format_report.get('total_paragraphs', 0)}\n")
                f.write(f"有问题段落数: {len(content_format_report['paragraphs_with_issues'])}\n\n")
                f.write("-" * 80 + "\n\n")
                
                for p in content_format_report['paragraphs_with_issues']:
                    f.write(f"段落 {p['index']}\n")
                    f.write(f"内容: {p['preview']}\n")
                    f.write(f"问题:\n")
                    for issue in p['issues']:
                        f.write(f"  - {issue}\n")
                    f.write("\n" + "-" * 80 + "\n\n")
            
            print(f"\n详细报告已保存至: {detail_file}")
        except Exception as e:
            print(f"\n警告: 保存详细报告失败: {e}")
    
    return report

# ---------- 报告输出 ----------
def print_content_report(report):
    """打印正文检查报告"""
    print("=== Content Format Check Report ===")
    
    sections = [
        ('hierarchy', 'TITLE HIERARCHY'),
        ('format', 'TITLE FORMAT'),
        ('case', 'TITLE CASE'),
        ('content_format', 'CONTENT FORMAT')
    ]
    
    for sec_key, sec_name in sections:
        info = report.get(sec_key, {'ok': False, 'messages': ['检测项不存在']})
        print(f"--- {sec_name} ---")
        print(" OK:", info['ok'])
        for m in info['messages']:
            print("  -", m)
    
    print("--- SUMMARY ---")
    for s in report['summary']:
        print(" ", s)

def print_help():
    print("Usage:")
    print("  python Content_detect.py check <paper.docx> <template.json_or_name>")

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print_help()
        sys.exit(0)
    
    cmd = sys.argv[1]
    if cmd == 'check':
        paper_path = sys.argv[2]
        tpl_id = sys.argv[3]
        if not os.path.isfile(paper_path):
            print(f"论文文件不存在: {paper_path}")
            sys.exit(1)
        try:
            print("=== 开始正文格式检查 ===")
            report = check_content_with_template(paper_path, tpl_id)
        except Exception as e:
            print("检查时出错:", e)
            sys.exit(1)
        print_content_report(report)
    else:
        print_help()
        sys.exit(0)
'''
python paper_detect\Content_detect.py check template\test.docx templates\Content.json     
'''