#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
from pathlib import Path
import logging
logger = logging.getLogger(__name__)

# 添加项目根目录到路径，以便导入模块
if __name__ == '__main__':
    project_root = Path(__file__).parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

"""
=== 论文格式检测系统 - 图片检测器 ===

【图片检测 (Figure Detection)】

1. 【图片标题格式检测】
   - 图片标题格式：Fig. 编号 标题文字（如：Fig. 1 Material structure）
   - 注意：Fig. 后必须有空格
   - 标题居中对齐
   - 标题字体：Times New Roman，五号(10.5pt)
   - 图片编号应连续递增

2. 【图片位置关系检测】
   - 图片应居中对齐
   - 标题位于图片下方
   - 图片与标题间隔合理

3. 【图片编号检测】
   - 编号从1开始
   - 编号连续递增（Fig. 1, Fig. 2, Fig. 3...）

4. 【技术特性】
   - 支持图片标题的智能识别
   - 验证图片编号的连续性
   - 检测图片对象的存在性
   - 分析图片和标题的对齐方式
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    """解析模板路径，支持文件路径和模板名称"""
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    """加载JSON模板文件"""
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 字体检测函数 ----------
def detect_font_for_run(run, paragraph=None):
    """
    更健壮的字体检测函数，同时检查直接格式和段落样式。
    返回 (font_size_pt, font_name, is_bold, is_italic, candidates_dict)
    """
    font_source = run.font if run else (paragraph.style.font if paragraph else None)
    if not font_source:
        return 12.0, "Unknown", False, False, {}

    font_name = font_source.name if font_source.name else "Times New Roman"

    font_size = None
    try:
        if run and run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        if font_size is None and run and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
        
        if font_size is None and paragraph and paragraph.style and hasattr(paragraph.style.font, 'size') and hasattr(paragraph.style.font.size, 'pt'):
            font_size = float(paragraph.style.font.size.pt)
            
        if font_size is None and paragraph and hasattr(paragraph.style, 'element'):
            sz_nodes = paragraph.style.element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass

    font_size = font_size if font_size is not None else 12.0

    is_bold = font_source.bold if font_source.bold is not None else False
    is_italic = font_source.italic if font_source.italic is not None else False
    if run and run.font:
        is_bold = run.font.bold if run.font.bold is not None else is_bold
        is_italic = run.font.italic if run.font.italic is not None else is_italic
    
    # 确保返回明确的布尔值，而不是None
    is_bold = bool(is_bold) if is_bold is not None else False
    is_italic = bool(is_italic) if is_italic is not None else False

    return font_size, font_name, is_bold, is_italic, {}

def get_font_size(pt_size, tpl=None):
    """获取字体大小的中文名称"""
    # 从模板读取字体大小映射，提供默认值以保持向后兼容
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        # 将字符串键转换为浮点数
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        # 默认的字体大小映射（向后兼容）
        size_map = {
            9: "小五",
            10.5: "五号",
            12: "小四",
            14: "四号",
            16: "三号",
            18: "小二",
            22: "二号",
            24: "小一",
            26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

# ---------- 段落对齐检测函数 ----------
def get_style_alignment(doc, style_id):
    """从文档的styles.xml中查找并返回指定样式的对齐方式。"""
    try:
        styles = doc.styles.element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        style_path = f".//w:style[@w:styleId='{style_id}']//w:jc"
        jc_element = styles.find(style_path, ns)
        if jc_element is not None:
            val = jc_element.get(ns['w'] + 'val')
            mapping = {'left': 0, 'center': 1, 'right': 2, 'both': 3, 'justify': 3}
            return mapping.get(val)
    except Exception:
        return None
    return None

def detect_paragraph_alignment(paragraph):
    """
    改进的段落对齐检测，按照明确的优先级处理直接格式和样式格式
    
    检测逻辑：
    1. 如果直接格式不为None，直接使用直接格式结果
    2. 如果直接格式为None，检查样式格式
    3. 检查XML样式定义（处理复杂的样式继承）
    4. 默认为左对齐
    """
    
    # 优先级1: 检查直接格式
    direct_alignment = paragraph.paragraph_format.alignment
    if direct_alignment is not None:
        # 如果有直接格式设置，直接使用
        return int(direct_alignment)
    
    # 优先级2: 如果直接格式为None，检查样式格式
    if paragraph.style:
        try:
            style_alignment = paragraph.style.paragraph_format.alignment
            if style_alignment is not None:
                # 如果样式有对齐设置，使用样式对齐
                return int(style_alignment)
        except Exception:
            pass
    
    # 优先级3: 检查XML样式定义（处理复杂的样式继承）
    if paragraph.style and paragraph.style.style_id:
        try:
            xml_alignment = get_style_alignment(paragraph.part.document, paragraph.style.style_id)
            if xml_alignment is not None:
                return xml_alignment
        except Exception:
            pass
            
    # 默认值: 左对齐
    return 0  # WD_PARAGRAPH_ALIGNMENT.LEFT

# ---------- 图片检测核心函数 ----------

def has_picture_object(paragraph):
    """
    检测段落是否包含图片对象
    支持现代格式(w:drawing)和旧格式(w:pict)
    返回：True/False
    """
    try:
        para_xml = paragraph._element
        # 检查现代Word格式的图片（w:drawing）
        drawings = para_xml.xpath('.//w:drawing')
        if drawings:
            return True
        
        # 检查旧格式的图片（w:pict）
        pictures = para_xml.xpath('.//w:pict')
        if pictures:
            return True
        
        # 额外检查：通过v:imagedata（VML图片）
        imagedata = para_xml.xpath('.//v:imagedata')
        if imagedata:
            return True
            
    except Exception:
        pass
    
    return False

def find_picture_captions(doc, tpl):
    """
    识别文档中的图片标题
    返回：图片标题段落列表，每个元素包含段落对象、编号、标题文本
    """
    caption_pattern = tpl.get('figure_detection_rules', {}).get('caption_pattern', r'^\s*Fig\.\s+(\d+)\s+(.+)$')
    captions = []
    
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        match = re.match(caption_pattern, text, re.IGNORECASE)
        if match:
            figure_num = int(match.group(1))
            figure_title = match.group(2).strip()
            captions.append({
                'paragraph': paragraph,
                'paragraph_index': idx,
                'number': figure_num,
                'title': figure_title,
                'full_text': text
            })
    
    return captions

def find_picture_before_caption(doc, caption_info):
    """
    在标题上方查找对应的图片
    返回：包含图片的段落对象或None
    """
    caption_idx = caption_info['paragraph_index']
    max_distance = 2  # 最多向上查找2个段落
    
    # 在标题前的几个段落中查找图片
    for i in range(caption_idx - 1, max(0, caption_idx - max_distance - 1), -1):
        if i < 0 or i >= len(doc.paragraphs):
            continue
        
        paragraph = doc.paragraphs[i]
        if has_picture_object(paragraph):
            return paragraph
    
    return None

def check_picture_alignment(picture_paragraph):
    """
    检查图片段落的对齐方式
    返回：(is_centered, actual_alignment)
    """
    if not picture_paragraph:
        return False, None
    
    actual_alignment = detect_paragraph_alignment(picture_paragraph)
    # 居中对齐的值为1
    is_centered = (actual_alignment == 1)
    
    return is_centered, actual_alignment

def check_caption_format(caption_info, tpl):
    """
    检查图片标题的格式
    返回：{'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    paragraph = caption_info['paragraph']
    expected_format = tpl.get('format_rules', {}).get('caption', {})
    
    # 检查对齐方式
    if expected_format.get('alignment') == 'center':
        actual_alignment = detect_paragraph_alignment(paragraph)
        
        # 对齐方式名称映射
        alignment_names = {
            0: '左对齐',
            1: '居中对齐',
            2: '右对齐',
            3: '两端对齐'
        }
        actual_alignment_name = alignment_names.get(actual_alignment, f'未知对齐方式({actual_alignment})')
        
        # 居中对齐的值为1
        if actual_alignment != 1:
            report['ok'] = False
            msg = tpl.get('messages', {}).get('caption_alignment_error', '图片标题应居中对齐')
            report['messages'].append(f"{msg}（当前：{actual_alignment_name}）")
    
    # 检查字体大小
    expected_size = expected_format.get('font_size_pt', 10.5)
    main_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if main_run:
        actual_size, actual_font, _, _, _ = detect_font_for_run(main_run, paragraph)
        if abs(actual_size - expected_size) > 0.5:
            report['ok'] = False
            expected_size_name = get_font_size(expected_size, tpl)
            actual_size_name = get_font_size(actual_size, tpl)
            msg = tpl.get('messages', {}).get('caption_font_size_error', '图片标题字体大小不正确')
            report['messages'].append(f"{msg}（期望：{expected_size_name}，实际：{actual_size_name}）")
    
    # 检查字体名称
    expected_font = expected_format.get('font_name', 'Times New Roman')
    if main_run:
        actual_size, actual_font, _, _, _ = detect_font_for_run(main_run, paragraph)
        if actual_font and actual_font != expected_font:
            report['ok'] = False
            report['messages'].append(f"图片标题字体应为{expected_font}（当前：{actual_font}）")
    
    return report

def check_figure_numbering(captions, tpl):
    """
    检查图片编号的连续性
    返回：{'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not captions:
        return report
    
    numbers = sorted([c['number'] for c in captions])
    
    # 检查是否从1开始
    if tpl.get('check_rules', {}).get('caption_numbering_check', True):
        expected_start = tpl.get('check_rules', {}).get('numbering_start', 1)
        if numbers[0] != expected_start:
            report['ok'] = False
            msg_template = tpl.get('messages', {}).get('caption_numbering_start_error', '图片编号应从{start}开始，实际从{actual}开始')
            report['messages'].append(msg_template.format(start=expected_start, actual=numbers[0]))
    
    # 检查是否连续
    if tpl.get('check_rules', {}).get('caption_sequential_check', True):
        expected_sequence = list(range(numbers[0], numbers[0] + len(numbers)))
        if numbers != expected_sequence:
            report['ok'] = False
            msg_template = tpl.get('messages', {}).get('caption_numbering_error', '图片编号不连续')
            report['messages'].append(msg_template.format(numbers=', '.join(map(str, numbers))))
    
    return report

def check_doc_with_template(doc_path, template_identifier, enable_content_check=True):
    """
    使用指定的模板检测文档中的图片格式
    
    参数:
        doc_path: Word文档路径
        template_identifier: 模板标识符（文件路径或模板名称）
        enable_content_check: 是否启用图片内容智能检测（需要API密钥）
    返回:
        检测报告字典
    """
    # 加载模板
    tpl = load_template(template_identifier)
    
    # 加载文档
    doc = Document(doc_path)
    
    # 如果启用内容检测，导入相关模块
    content_detector = None
    if enable_content_check:
        try:
            # 尝试绝对导入
            try:
                from paper_detect.Figure_content_detect import FigureContentDetector
            except ImportError:
                # 如果失败，尝试相对导入
                from .Figure_content_detect import FigureContentDetector
            
            # 初始化检测器（会自动从配置文件读取）
            content_detector = FigureContentDetector()
            print(f"✓ 图片内容智能检测已启用")
            if content_detector.save_images:
                print(f"  图片将保存到: {content_detector.image_dir}/ 目录")
            else:
                print(f"  使用临时文件（分析后自动删除）")
        except ValueError as e:
            print(f"警告: {e}")
            print("  将跳过内容检测")
        except ImportError as e:
            print(f"警告: 无法加载图片内容检测模块: {e}")
            print("  将跳过内容检测")
    
    # 初始化报告
    report = {
        'overall': {'ok': True, 'messages': []},
        'figures': [],  # 改为按图片组织
        'numbering': {'ok': True, 'messages': []},
        'summary': {}
    }
    
    # 1. 找出所有包含图片的段落（按文档顺序）
    picture_paragraphs = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if has_picture_object(paragraph):
            picture_paragraphs.append({
                'paragraph': paragraph,
                'paragraph_index': idx
            })
    
    if not picture_paragraphs:
        report['overall']['ok'] = False
        report['overall']['messages'].append('文档中未找到任何图片')
        report['summary']['figure_count'] = 0
        return report
    
    report['summary']['figure_count'] = len(picture_paragraphs)
    
    # 2. 对每张图片进行检查
    caption_pattern = tpl.get('figure_detection_rules', {}).get('caption_pattern', r'^\s*Fig\.\s+(\d+)\s+(.+)$')
    figure_numbers = []
    
    for fig_idx, pic_info in enumerate(picture_paragraphs, start=1):
        picture_para = pic_info['paragraph']
        pic_index = pic_info['paragraph_index']
        
        figure_report = {
            'figure_index': fig_idx,
            'paragraph_index': pic_index,
            'has_caption': False,
            'caption_info': None,
            'format_check': {'ok': True, 'messages': []},
            'position_check': {'ok': True, 'messages': []},
            'picture_check': {'ok': True, 'messages': []}
        }
        
        # 2.1 检查是否有标题（向下查找）
        caption_found = None
        for i in range(pic_index + 1, min(pic_index + 3, len(doc.paragraphs))):
            caption_para = doc.paragraphs[i]
            match = re.match(caption_pattern, caption_para.text.strip(), re.IGNORECASE)
            if match:
                figure_num = int(match.group(1))
                figure_title = match.group(2).strip()
                caption_found = {
                    'paragraph': caption_para,
                    'number': figure_num,
                    'title': figure_title,
                    'full_text': caption_para.text.strip()
                }
                figure_numbers.append(figure_num)
                break
        
        if caption_found:
            figure_report['has_caption'] = True
            figure_report['caption_info'] = caption_found
            
            # 2.2 检查标题格式
            format_result = check_caption_format(caption_found, tpl)
            figure_report['format_check'] = format_result
            if not format_result['ok']:
                report['overall']['ok'] = False
        else:
            # 没有标题
            figure_report['has_caption'] = False
            figure_report['format_check']['ok'] = False
            figure_report['format_check']['messages'].append('❌ 图片缺少标题（应为：Fig. 编号 标题文字）')
            report['overall']['ok'] = False
        
        # 2.3 检查图片对齐
        if tpl.get('check_rules', {}).get('picture_alignment_check', True):
            is_centered, alignment = check_picture_alignment(picture_para)
            if not is_centered:
                alignment_names = {0: '左对齐', 1: '居中对齐', 2: '右对齐', 3: '两端对齐'}
                alignment_name = alignment_names.get(alignment, f'未知({alignment})')
                figure_report['picture_check']['ok'] = False
                figure_report['picture_check']['messages'].append(
                    f"图片应居中对齐（当前：{alignment_name}）"
                )
                report['overall']['ok'] = False
        
        # 2.4 检查图片内容（如果启用）- 不管有没有标题都检查
        if content_detector:
            print(f"  正在检测第 {fig_idx} 张图片的内容规范性...")
            fig_num = caption_found['number'] if caption_found else fig_idx
            content_result = content_detector.detect_figure_content(
                picture_para, 
                doc_path, 
                figure_number=fig_num
            )
            figure_report['content_check'] = content_result
            
            if not content_result['ok']:
                report['overall']['ok'] = False
        
        report['figures'].append(figure_report)
    
    # 3. 检查编号连续性（只对有标题的图片）
    if figure_numbers:
        numbers_sorted = sorted(figure_numbers)
        expected_start = tpl.get('check_rules', {}).get('numbering_start', 1)
        
        if numbers_sorted[0] != expected_start:
            report['numbering']['ok'] = False
            report['numbering']['messages'].append(
                f"图片编号应从{expected_start}开始，实际从{numbers_sorted[0]}开始"
            )
            report['overall']['ok'] = False
        
        expected_sequence = list(range(numbers_sorted[0], numbers_sorted[0] + len(numbers_sorted)))
        if numbers_sorted != expected_sequence:
            report['numbering']['ok'] = False
            report['numbering']['messages'].append(
                f"图片编号不连续，发现编号：{', '.join(map(str, numbers_sorted))}"
            )
            report['overall']['ok'] = False
    
    # 兼容处理：生成旧格式的captions列表
    report['captions'] = []
    for fig_report in report['figures']:
        if fig_report['has_caption']:
            report['captions'].append({
                'number': fig_report['caption_info']['number'],
                'title': fig_report['caption_info']['title'],
                'full_text': fig_report['caption_info']['full_text'],
                'format_check': fig_report['format_check'],
                'position_check': fig_report.get('position_check', {'ok': True, 'messages': []}),
                'content_check': fig_report.get('content_check')
            })
    
    return report



# ---------- 报告生成 ----------

def print_report(report, tpl):
    """打印格式化的检测报告"""
    print("=" * 60)
    print("图片格式检测报告")
    print("=" * 60)
    print()
    
    # 总结
    print("【检查总结】")
    result_text = "通过" if report['overall']['ok'] else "发现问题"
    msg_template = tpl.get('messages', {}).get('summary_overall', '图片格式检查结果: {ok}')
    print(f"  {msg_template.format(ok=result_text)}")
    
    count_msg = tpl.get('messages', {}).get('figure_count', '检测到 {count} 个图片')
    print(f"  {count_msg.format(count=report['summary'].get('figure_count', 0))}")
    
    print()
    
    # 编号检查（如果有标题的图片）
    has_captions = any(fig['has_caption'] for fig in report.get('figures', []))
    if has_captions and report['numbering']['messages']:
        print("【图片编号检查】")
        for msg in report['numbering']['messages']:
            print(f"  ✗ {msg}")
        print()
    
    # 按图片顺序显示详细检查
    if report.get('figures'):
        print("【各图片详细检查】")
        print()
        
        for fig_report in report['figures']:
            fig_idx = fig_report['figure_index']
            
            # 显示图片标题
            if fig_report['has_caption']:
                caption = fig_report['caption_info']
                print(f"  第 {fig_idx} 张图片: Fig.{caption['number']} {caption['title']}")
            else:
                print(f"  第 {fig_idx} 张图片: (无标题)")
            
            # 所有问题
            all_issues = []
            
            # 1. 标题问题
            if not fig_report['has_caption']:
                all_issues.append("❌ 缺少标题")
            elif not fig_report['format_check']['ok']:
                all_issues.append("标题格式问题：")
                for msg in fig_report['format_check']['messages']:
                    all_issues.append(f"  - {msg}")
            else:
                all_issues.append("✓ 标题格式正确")
            
            # 2. 图片本身的问题
            if not fig_report['picture_check']['ok']:
                all_issues.append("图片问题：")
                for msg in fig_report['picture_check']['messages']:
                    all_issues.append(f"  - {msg}")
            else:
                if fig_report['has_caption']:  # 只有有标题的才显示图片OK
                    all_issues.append("✓ 图片位置正确")
            
            # 3. 内容检查（如果有）
            if 'content_check' in fig_report:
                content_check = fig_report['content_check']
                if content_check.get('is_chart', False):
                    if content_check['ok']:
                        all_issues.append("✓ 图表内容符合规范")
                    else:
                        all_issues.append("图表内容问题：")
                        for msg in content_check['messages']:
                            all_issues.append(f"  - {msg}")
                else:
                    all_issues.append(f"ℹ 图片类型: {content_check.get('messages', ['非图表'])[0]}")
            
            # 显示所有问题
            for issue in all_issues:
                if issue.startswith("  -"):
                    print(f"    {issue}")
                else:
                    print(f"    {issue}")
            
            print()
    
    print("=" * 60)

# ---------- 命令行接口 ----------

def main():
    """命令行入口"""
    if len(sys.argv) < 3:
        print(__doc__)
        print("\n使用方法：")
        print("  python Figure_detect.py check <docx文件路径> <模板标识符> [--no-api]")
        print("\n参数说明：")
        print("  --no-api    不调用API检测图片内容（只检测格式）")
        print("\n示例：")
        print("  python Figure_detect.py check template/test.docx Figure")
        print("  python Figure_detect.py check template/test.docx Figure --no-api")
        print("\n注意：图片标题格式为 Fig. 编号 标题（Fig.后必须有空格）")
        sys.exit(1)
    
    command = sys.argv[1]
    
    if command == "check":
        if len(sys.argv) < 4:
            print("错误：check命令至少需要2个参数")
            print("使用方法：python Figure_detect.py check <docx文件路径> <模板标识符> [--no-api]")
            sys.exit(1)
        
        doc_path = sys.argv[2]
        template_id = sys.argv[3]
        
        # 检查是否有 --no-api 参数
        enable_api = True
        if len(sys.argv) >= 5 and sys.argv[4] == '--no-api':
            enable_api = False
            print("注意：已禁用API内容检测，只检测图片格式")
        elif len(sys.argv) >= 5:
            print(f"警告：未知参数 '{sys.argv[4]}'，将忽略")
        
        # 检查文件是否存在
        if not os.path.isfile(doc_path):
            print(f"错误：文件不存在: {doc_path}")
            sys.exit(1)
        
        try:
            # 执行检测（根据命令行参数决定是否启用内容检测）
            report = check_doc_with_template(
                doc_path, 
                template_id,
                enable_content_check=enable_api
            )
            
            # 加载模板用于打印
            tpl = load_template(template_id)
            
            # 打印报告
            print_report(report, tpl)
            
            # 保存完整的API响应到JSON文件（只在启用API检测时）
            if enable_api and report.get('figures'):
                import json
                from pathlib import Path
                
                output_dir = Path('api_results')
                output_dir.mkdir(exist_ok=True)
                
                doc_name = Path(doc_path).stem
                output_file = output_dir / f'{doc_name}_api_response.json'
                
                # 提取API响应
                api_responses = {}
                for fig_report in report['figures']:
                    if 'content_check' in fig_report and fig_report['content_check']:
                        fig_idx = fig_report['figure_index']
                        content_check = fig_report['content_check']
                        
                        # 确定图片标识
                        if fig_report['has_caption']:
                            fig_key = f"Fig.{fig_report['caption_info']['number']}"
                        else:
                            fig_key = f"图片{fig_idx}"
                        
                        api_responses[fig_key] = {
                            'image_path': content_check.get('image_path'),
                            'is_chart': content_check.get('is_chart'),
                            'ok': content_check.get('ok'),
                            'messages': content_check.get('messages'),
                            'api_details': content_check.get('details')
                        }
                
                if api_responses:
                    with open(output_file, 'w', encoding='utf-8') as f:
                        json.dump(api_responses, f, ensure_ascii=False, indent=2)
                    print(f"\n✓ API响应已保存到: {output_file}")
            
        except Exception as e:
            print(f"检测过程中发生错误: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)
    else:
        print(f"未知命令: {command}")
        print("支持的命令: check")
        sys.exit(1)

if __name__ == '__main__':
    main()

