#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

"""
=== 论文格式检测系统 - 公式检测器 ===

【公式检测 (Formula Detection)】

1. 【公式格式检测】
   - 公式字体为小四号(12pt)
   - 编号为Times New Roman字体
   - 公式内容为Cambria Math字体
   - 公式样式设置制表位：20字符居中对齐，40字符右对齐

2. 【制表位检测】
   - 第一个制表位：20字符位置，居中对齐
   - 第二个制表位：40字符位置，右对齐
   - 制表位设置通过段落样式或直接格式应用

3. 【公式对象检测】
   - 检测Word文档中的Office Math (oMath)对象
   - 验证公式内容的存在性和格式正确性
   - 区分公式内容和编号文本

4. 【技术特性】
   - 支持XML解析检测oMath对象
   - 智能制表位位置和对齐方式验证
   - 分别检测公式内容和编号的字体设置
   - 提供详细的中文检查报告
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    """解析模板路径，支持文件路径和模板名称"""
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    """加载JSON模板文件"""
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 字体检测函数 ----------
def detect_font_for_run(run, paragraph=None):
    """
    更健壮的字体检测函数，同时检查直接格式和段落样式。
    返回 (font_size_pt, font_name, is_bold, is_italic, candidates_dict)
    """
    font_source = run.font if run else (paragraph.style.font if paragraph else None)
    if not font_source:
        return 12.0, "Unknown", False, False, {}

    font_name = font_source.name if font_source.name else "Times New Roman"

    font_size = None
    try:
        if run and run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        if font_size is None and run and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
        
        if font_size is None and paragraph and paragraph.style and hasattr(paragraph.style.font, 'size') and hasattr(paragraph.style.font.size, 'pt'):
            font_size = float(paragraph.style.font.size.pt)
            
        if font_size is None and paragraph and hasattr(paragraph.style, 'element'):
            sz_nodes = paragraph.style.element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass

    font_size = font_size if font_size is not None else 12.0

    is_bold = font_source.bold if font_source.bold is not None else False
    is_italic = font_source.italic if font_source.italic is not None else False
    if run and run.font:
        is_bold = run.font.bold if run.font.bold is not None else is_bold
        is_italic = run.font.italic if run.font.italic is not None else is_italic
    
    # 确保返回明确的布尔值，而不是None
    is_bold = bool(is_bold) if is_bold is not None else False
    is_italic = bool(is_italic) if is_italic is not None else False

    return font_size, font_name, is_bold, is_italic, {}

def get_font_size(pt_size, tpl=None):
    """获取字体大小的中文名称"""
    # 从模板读取字体大小映射，提供默认值以保持向后兼容
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        # 将字符串键转换为浮点数
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        # 默认的字体大小映射（向后兼容）
        size_map = {
            9: "小五",
            10.5: "五号",
            12: "小四",
            14: "四号",
            16: "三号",
            18: "小二",
            22: "二号",
            24: "小一",
            26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

# ---------- 公式检测核心函数 ----------

def identify_formula_paragraphs(doc):
    """
    识别真正的Word公式段落
    优先检测Office Math对象和制表位设置，减少误识别
    """
    formula_paragraphs = []
    
    for paragraph in doc.paragraphs:
        # 跳过空段落
        if not paragraph.text.strip():
            continue
            
        # 1. 优先检查是否包含Office Math对象（最可靠的指标）
        has_math_object = False
        math_object_count = 0
        try:
            para_xml = paragraph._element
            # 检查多个可能的数学对象命名空间
            math_elements = (
                para_xml.xpath('.//w:oMath') + 
                para_xml.xpath('.//m:oMath') +
                para_xml.xpath('.//oMath') +
                para_xml.xpath('.//w:r/w:object[contains(@w:progId, "Equation")]') +  # 公式对象
                para_xml.xpath('.//w:r/w:object[contains(@w:progId, "MathType")]')    # MathType对象
            )
            if math_elements:
                has_math_object = True
                math_object_count = len(math_elements)
        except Exception:
            pass
        
        # 2. 检查是否包含公式样式的制表位设置
        has_formula_tab_stops = False
        center_tabs = []
        right_tabs = []
        
        try:
            # 首先检查段落格式中的制表位
            if hasattr(paragraph.paragraph_format, '_element'):
                pPr = paragraph.paragraph_format._element
                tabs_elements = pPr.xpath('.//w:tabs/w:tab')
                
                for tab in tabs_elements:
                    pos = tab.get(qn('w:pos'))
                    val = tab.get(qn('w:val'))
                    if pos and val:
                        pos_twips = int(pos)
                        pos_chars = round(pos_twips / 210.0)  # 1字符 ≈ 210 twips (根据实际测试修正)
                        
                        if val == 'center' and 10 <= pos_chars <= 50:
                            center_tabs.append(pos_chars)
                        elif val == 'right' and 20 <= pos_chars <= 100:
                            right_tabs.append(pos_chars)
            
            # 如果段落格式中没有找到制表位，检查样式中的制表位
            if not center_tabs and not right_tabs and paragraph.style and hasattr(paragraph.style, '_element'):
                style_elem = paragraph.style._element
                style_tabs = style_elem.xpath('.//w:tabs/w:tab')
                
                for tab in style_tabs:
                    pos = tab.get(qn('w:pos'))
                    val = tab.get(qn('w:val'))
                    if pos and val:
                        pos_twips = int(pos)
                        pos_chars = round(pos_twips / 210.0)  # 1字符 ≈ 210 twips (根据实际测试修正)
                        
                        if val == 'center' and 10 <= pos_chars <= 50:
                            center_tabs.append(pos_chars)
                        elif val == 'right' and 20 <= pos_chars <= 100:
                            right_tabs.append(pos_chars)
            
            # 如果同时有居中和右对齐制表位，很可能是公式样式
            if center_tabs and right_tabs:
                has_formula_tab_stops = True
                
        except Exception:
            pass
        
        # 3. 获取段落文本（用于后续处理）
        text = paragraph.text.strip()
        
        # 4. 检查段落样式名称（如果应用了公式样式）
        has_formula_style = False
        try:
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            formula_style_keywords = ['formula', 'equation', '公式', 'math']
            if any(keyword in style_name for keyword in formula_style_keywords):
                has_formula_style = True
        except Exception:
            pass
        
        # 4. 决策逻辑：只识别真正的公式
        is_formula_paragraph = False
        confidence_score = 0
        
        # 最高优先级：包含Office Math对象
        if has_math_object:
            is_formula_paragraph = True
            confidence_score = 10
        # 次高优先级：有公式制表位设置
        elif has_formula_tab_stops:
            is_formula_paragraph = True
            confidence_score = 8
        # 第三优先级：应用了公式样式
        elif has_formula_style:
            is_formula_paragraph = True
            confidence_score = 6
        
        if is_formula_paragraph:
            formula_paragraphs.append({
                'paragraph': paragraph,
                'has_math_object': has_math_object,
                'math_object_count': math_object_count,
                'has_formula_tab_stops': has_formula_tab_stops,
                'has_formula_style': has_formula_style,
                'confidence_score': confidence_score,
                'text': text
            })
    
    return formula_paragraphs

def check_tab_stops(paragraph, expected_tabs):
    """
    检测段落的制表位设置
    expected_tabs: [{"position_chars": 20, "alignment": "center"}, {"position_chars": 40, "alignment": "right"}]
    返回 (is_correct, issues, detected_tabs)
    """
    issues = []
    detected_tabs = []
    
    try:
        # 首先检查段落格式中的制表位
        pPr = paragraph.paragraph_format._element
        tabs_elements = pPr.xpath('.//w:tabs')
        
        # 如果段落格式中没有制表位，检查样式中的制表位
        if not tabs_elements and paragraph.style and hasattr(paragraph.style, '_element'):
            style_elem = paragraph.style._element
            tabs_elements = style_elem.xpath('.//w:tabs')
        
        if not tabs_elements:
            issues.append("未检测到制表位设置（段落格式和样式中都没有）")
            return False, issues, detected_tabs
        
        # 解析制表位
        for tabs_elem in tabs_elements:
            tab_elements = tabs_elem.xpath('.//w:tab')
            for tab_elem in tab_elements:
                pos_attr = tab_elem.get(qn('w:pos'))
                val_attr = tab_elem.get(qn('w:val'))
                
                if pos_attr and val_attr:
                    # 转换位置从twips到字符（1字符 ≈ 210 twips，根据实际测试修正）
                    pos_twips = int(pos_attr)
                    pos_chars = round(pos_twips / 210.0)
                    
                    # 映射对齐方式
                    alignment_map = {
                        'center': 'center',
                        'right': 'right',
                        'left': 'left',
                        'decimal': 'decimal'
                    }
                    alignment = alignment_map.get(val_attr, val_attr)
                    
                    detected_tabs.append({
                        'position_chars': pos_chars,
                        'position_twips': pos_twips,
                        'alignment': alignment
                    })
        
        # 验证检测到的制表位是否符合要求
        if len(detected_tabs) < len(expected_tabs):
            issues.append(f"制表位数量不足，期望{len(expected_tabs)}个，实际{len(detected_tabs)}个")
        
        for i, expected_tab in enumerate(expected_tabs):
            expected_pos = expected_tab['position_chars']
            expected_align = expected_tab['alignment']
            
            # 查找匹配的制表位（允许±2字符的误差）
            matching_tab = None
            for detected_tab in detected_tabs:
                if abs(detected_tab['position_chars'] - expected_pos) <= 2:
                    matching_tab = detected_tab
                    break
            
            if not matching_tab:
                issues.append(f"未找到位置为{expected_pos}字符的制表位")
            else:
                if matching_tab['alignment'] != expected_align:
                    issues.append(f"制表位{expected_pos}字符处对齐方式错误，期望{expected_align}，实际{matching_tab['alignment']}")
        
        # 检查制表符的实际使用情况
        para_text = paragraph.text
        tab_char_count = para_text.count('\t')
        expected_tab_chars = len(expected_tabs)  # 期望的制表符数量
        
        if tab_char_count == 0:
            issues.append("设置了制表位但没有使用制表符，公式不会按预期对齐")
            issues.append("建议：在公式前按Tab键跳到居中位置，在公式后按Tab键跳到右对齐位置")
        elif tab_char_count < expected_tab_chars:
            issues.append(f"制表符使用不足，期望{expected_tab_chars}个，实际{tab_char_count}个")
            if tab_char_count == 1:
                issues.append("建议：在公式前和公式后都要按Tab键，实现居中和右对齐")
        elif tab_char_count > expected_tab_chars:
            issues.append(f"制表符使用过多，期望{expected_tab_chars}个，实际{tab_char_count}个")
        
        is_correct = len(issues) == 0
        return is_correct, issues, detected_tabs
        
    except Exception as e:
        issues.append(f"制表位检测异常: {str(e)}")
        return False, issues, detected_tabs

def detect_math_objects(paragraph):
    """
    在段落中检测Office Math对象
    返回 (has_math, math_objects, math_info)
    """
    math_objects = []
    math_info = {
        'count': 0,
        'types': [],
        'content_preview': []
    }
    
    try:
        para_xml = paragraph._element
        
        # 搜索多种可能的数学对象命名空间
        math_elements = []
        
        # 常见的数学对象XML路径
        math_paths = [
            './/w:oMath',      # Word 2007+标准路径
            './/m:oMath',      # 数学命名空间
            './/oMath',        # 简化路径
            './/w:r/w:object', # 嵌入对象
            './/w:r[.//w:oMath]' # 包含数学对象的run
        ]
        
        for path in math_paths:
            try:
                elements = para_xml.xpath(path)
                math_elements.extend(elements)
            except Exception:
                continue
        
        # 去重
        unique_elements = []
        for elem in math_elements:
            if elem not in unique_elements:
                unique_elements.append(elem)
        
        math_info['count'] = len(unique_elements)
        
        for math_elem in unique_elements:
            try:
                # 提取数学对象信息
                math_obj = {
                    'element': math_elem,
                    'tag': math_elem.tag,
                    'text_content': '',
                    'type': 'unknown'
                }
                
                # 尝试提取文本内容
                if hasattr(math_elem, 'text') and math_elem.text:
                    math_obj['text_content'] = math_elem.text.strip()
                else:
                    # 递归提取所有文本节点
                    text_parts = []
                    for text_node in math_elem.iter():
                        if hasattr(text_node, 'text') and text_node.text:
                            text_parts.append(text_node.text.strip())
                    math_obj['text_content'] = ' '.join(filter(None, text_parts))
                
                # 判断数学对象类型
                if 'oMath' in math_elem.tag:
                    math_obj['type'] = 'office_math'
                elif 'object' in math_elem.tag:
                    math_obj['type'] = 'embedded_object'
                else:
                    math_obj['type'] = 'math_container'
                
                math_objects.append(math_obj)
                math_info['types'].append(math_obj['type'])
                
                # 保存内容预览（限制长度）
                preview = math_obj['text_content'][:50] + ('...' if len(math_obj['text_content']) > 50 else '')
                if preview:
                    math_info['content_preview'].append(preview)
                    
            except Exception as e:
                # 记录但不中断处理
                print(f"处理数学对象时出错: {e}")
                continue
        
        has_math = len(math_objects) > 0
        return has_math, math_objects, math_info
        
    except Exception as e:
        print(f"检测数学对象时出错: {e}")
        return False, [], math_info

def check_formula_fonts(paragraph, math_objects, template):
    """
    检测公式内容和编号的字体设置
    返回 (is_correct, issues, font_info)
    """
    issues = []
    font_info = {
        'formula_fonts': [],
        'number_fonts': [],
        'all_fonts': []
    }
    
    try:
        # 获取模板中的字体要求
        font_requirements = template.get('formula_detection_rules', {}).get('font_requirements', {})
        expected_formula_font = font_requirements.get('formula_content', 'Cambria Math')
        expected_number_font = font_requirements.get('formula_number', 'Times New Roman')
        expected_size_pt = float(font_requirements.get('font_size_pt', 12))
        
        # 检测段落中所有run的字体
        for run in paragraph.runs:
            if not run.text.strip():
                continue
                
            font_size, font_name, is_bold, is_italic, _ = detect_font_for_run(run, paragraph)
            
            font_detail = {
                'text': run.text.strip(),
                'font_name': font_name,
                'font_size_pt': font_size,
                'is_bold': is_bold,
                'is_italic': is_italic,
                'run': run
            }
            
            font_info['all_fonts'].append(font_detail)
            
            # 判断是否为公式内容或编号
            text = run.text.strip()
            
            # 启发式判断：包含数学符号或希腊字母的可能是公式内容
            is_likely_formula = False
            formula_indicators = [
                r'[=<>≤≥≠±∞∑∏∫]',  # 数学符号
                r'[α-ωΑ-Ω]',         # 希腊字母
                r'[₀-₉⁰-⁹]',         # 上下标
                r'[∂∇∆]',           # 微分符号
            ]
            
            for pattern in formula_indicators:
                if re.search(pattern, text):
                    is_likely_formula = True
                    break
            
            # 判断是否为编号（更宽松的匹配）
            is_likely_number = False
            
            # 1. 完整的编号格式：(1), [2], {3} 等
            if re.match(r'^\s*[\(\[\{]\s*\d+\s*[\)\]\}]\s*$', text):
                is_likely_number = True
            # 2. 单独的数字（在公式段落中很可能是编号）
            elif re.match(r'^\s*\d+\s*$', text) and len(text.strip()) <= 3:
                is_likely_number = True
            # 3. 单独的括号（可能是编号的一部分）
            elif re.match(r'^\s*[\(\)\[\]\{\}]\s*$', text):
                is_likely_number = True
            # 4. 引用格式：[1], [2-3] 等
            elif re.match(r'^\s*\[\s*\d+(-\d+)?\s*\]\s*$', text):
                is_likely_number = True
            
            if is_likely_formula:
                font_info['formula_fonts'].append(font_detail)
                # 检查公式内容字体
                if expected_formula_font.lower() not in font_name.lower():
                    issues.append(f"公式内容字体应为{expected_formula_font}，实际为{font_name}（内容：{text[:20]}...）")
            elif is_likely_number:
                font_info['number_fonts'].append(font_detail)
                # 检查编号字体
                if expected_number_font.lower() not in font_name.lower():
                    issues.append(f"公式编号字体应为{expected_number_font}，实际为{font_name}（编号：{text}）")
            
            # 检查字体大小（所有文本都应该是小四号）
            if abs(font_size - expected_size_pt) > 0.5:
                size_name = get_font_size(font_size, template)
                expected_size_name = get_font_size(expected_size_pt, template)
                issues.append(f"字体大小应为{expected_size_name}({expected_size_pt}pt)，实际为{size_name}({font_size}pt)（内容：{text[:20]}...）")
        
        # 如果检测到数学对象，验证是否有Cambria Math字体
        if math_objects:
            # 检查常规run中是否有Cambria Math字体
            has_cambria_math_in_runs = any('cambria math' in font_detail['font_name'].lower() 
                                         for font_detail in font_info['all_fonts'])
            
            # 检查Office Math对象内部是否有Cambria Math字体
            has_cambria_math_in_math = False
            try:
                for math_obj in math_objects:
                    math_elem = math_obj['element']
                    # 简单检查：如果XML中包含Cambria Math，就认为有
                    math_xml = ET.tostring(math_elem, encoding='unicode')
                    if 'cambria math' in math_xml.lower():
                        has_cambria_math_in_math = True
                        break
            except Exception:
                pass
            
            # 如果在常规run或Math对象中都没有找到Cambria Math，才报错
            if not has_cambria_math_in_runs and not has_cambria_math_in_math:
                issues.append("检测到Office Math对象但未找到Cambria Math字体")
            elif has_cambria_math_in_math:
                # 如果在Math对象中找到了Cambria Math，添加到公式字体信息中
                math_font_detail = {
                    'text': math_obj.get('text_content', '数学公式'),
                    'font_name': 'Cambria Math',
                    'font_size_pt': expected_size_pt,
                    'is_bold': False,
                    'is_italic': False,
                    'source': 'office_math'
                }
                font_info['formula_fonts'].append(math_font_detail)
                font_info['all_fonts'].append(math_font_detail)
        
        # 如果段落看起来像公式但没有找到编号，给出提示
        if (font_info['formula_fonts'] or math_objects) and not font_info['number_fonts']:
            # 检查是否有可能的编号文本（更宽松的匹配）
            possible_numbers = []
            for font_detail in font_info['all_fonts']:
                text = font_detail['text']
                if re.search(r'\d+', text) and len(text.strip()) <= 10:  # 短文本且包含数字
                    possible_numbers.append(text)
            
            if possible_numbers:
                issues.append(f"疑似公式编号但字体可能不正确：{', '.join(possible_numbers)}")
        
        is_correct = len(issues) == 0
        return is_correct, issues, font_info
        
    except Exception as e:
        issues.append(f"字体检测异常: {str(e)}")
        return False, issues, font_info

def validate_formula_format(paragraph, template):
    """
    综合验证公式格式，整合所有检测结果
    返回 {'ok': bool, 'messages': [], 'details': {}}
    """
    report = {'ok': True, 'messages': [], 'details': {}}
    
    try:
        # 获取模板中的制表位要求
        tab_requirements = template.get('formula_detection_rules', {}).get('tab_stops', [])
        
        # 1. 检测制表位
        tab_correct, tab_issues, detected_tabs = check_tab_stops(paragraph, tab_requirements)
        report['details']['tab_stops'] = {
            'correct': tab_correct,
            'issues': tab_issues,
            'detected': detected_tabs
        }
        
        if not tab_correct:
            report['ok'] = False
            report['messages'].extend([f"制表位问题: {issue}" for issue in tab_issues])
        
        # 2. 检测数学对象
        has_math, math_objects, math_info = detect_math_objects(paragraph)
        report['details']['math_objects'] = {
            'has_math': has_math,
            'count': math_info['count'],
            'types': math_info['types'],
            'preview': math_info['content_preview']
        }
        
        if not has_math:
            # 不一定是错误，可能是手动输入的公式
            report['messages'].append("未检测到Office Math对象，可能是手动输入的公式")
        
        # 3. 检测字体
        font_correct, font_issues, font_info = check_formula_fonts(paragraph, math_objects, template)
        report['details']['fonts'] = {
            'correct': font_correct,
            'issues': font_issues,
            'formula_fonts': len(font_info['formula_fonts']),
            'number_fonts': len(font_info['number_fonts']),
            'all_fonts': len(font_info['all_fonts'])
        }
        
        if not font_correct:
            report['ok'] = False
            report['messages'].extend([f"字体问题: {issue}" for issue in font_issues])
        
        # 4. 综合评估
        if report['ok']:
            if has_math and tab_correct and font_correct:
                report['messages'].append("公式格式检查通过：制表位设置正确，包含数学对象，字体符合要求")
            elif tab_correct and font_correct:
                report['messages'].append("公式格式基本正确：制表位和字体符合要求")
            else:
                report['messages'].append("公式格式检查通过")
        
        # 5. 添加详细信息
        if detected_tabs:
            tab_desc = []
            for tab in detected_tabs:
                tab_desc.append(f"{tab['position_chars']}字符({tab['alignment']})")
            report['messages'].append(f"检测到制表位: {', '.join(tab_desc)}")
        
        if math_info['content_preview']:
            preview_text = ', '.join(math_info['content_preview'][:2])  # 只显示前两个
            report['messages'].append(f"检测到数学内容: {preview_text}")
        
        return report
        
    except Exception as e:
        report['ok'] = False
        report['messages'].append(f"公式格式验证异常: {str(e)}")
        return report

def check_doc_with_template(doc_path, template_identifier):
    """
    主检查函数：使用模板检查文档中的公式格式
    返回完整的检查报告
    """
    try:
        # 加载模板
        template = load_template(template_identifier)
        
        # 打开文档
        doc = Document(doc_path)
        
        # 识别公式段落
        formula_paragraphs = identify_formula_paragraphs(doc)
        
        # 初始化报告
        report = {
            'formula_detection': {'ok': True, 'messages': []},
            'summary': [],
            'details': {
                'total_paragraphs': len(doc.paragraphs),
                'formula_paragraphs_count': len(formula_paragraphs),
                'formula_paragraphs': []
            }
        }
        
        if not formula_paragraphs:
            report['formula_detection']['ok'] = False
            report['formula_detection']['messages'].append("未检测到公式段落")
            report['summary'].append("文档中未找到符合公式格式的段落")
            return report
        
        # 检查每个公式段落
        all_formulas_ok = True
        
        for i, formula_para in enumerate(formula_paragraphs):
            paragraph = formula_para['paragraph']
            
            # 验证公式格式
            para_report = validate_formula_format(paragraph, template)
            
            # 获取更好的文本预览（数学内容在前，编号在后）
            math_content = ""
            formula_number = ""
            
            # 如果段落包含Office Math对象，提取数学内容
            if formula_para['has_math_object']:
                try:
                    para_xml = paragraph._element
                    math_elements = para_xml.xpath('.//m:oMath | .//w:oMath | .//oMath')
                    
                    math_texts = []
                    for math_elem in math_elements:
                        # 提取Math对象中的文本内容
                        t_elements = math_elem.xpath('.//*[local-name()="t"]')
                        math_text = ''.join(t.text for t in t_elements if t.text)
                        if math_text:
                            math_texts.append(math_text)
                    
                    if math_texts:
                        math_content = ' | '.join(math_texts)
                except Exception:
                    pass
            
            # 提取编号（从段落文本中）
            para_text = paragraph.text.strip()
            number_matches = re.findall(r'\((\d+)\)', para_text)
            if number_matches:
                formula_number = f"({number_matches[-1]})"
            
            # 构建预览文本：数学内容 + 编号
            if math_content and formula_number:
                text_preview = f"{math_content} {formula_number}"
            elif math_content:
                text_preview = math_content
            elif formula_number:
                text_preview = f"[公式] {formula_number}"
            else:
                text_preview = para_text
            
            # 添加段落信息
            para_info = {
                'index': i + 1,
                'text_preview': text_preview[:150] + ('...' if len(text_preview) > 150 else ''),
                'has_math_object': formula_para['has_math_object'],
                'math_object_count': formula_para['math_object_count'],
                'has_formula_tab_stops': formula_para['has_formula_tab_stops'],
                'has_formula_style': formula_para['has_formula_style'],
                'confidence_score': formula_para['confidence_score'],
                'format_check': para_report
            }
            
            report['details']['formula_paragraphs'].append(para_info)
            
            if not para_report['ok']:
                all_formulas_ok = False
                report['formula_detection']['ok'] = False
                
                # 添加具体的错误信息
                header = f"公式段落 {i + 1} 格式问题："
                report['formula_detection']['messages'].append(header)
                for msg in para_report['messages']:
                    report['formula_detection']['messages'].append(f"  - {msg}")
            else:
                # 添加成功信息
                success_msg = f"公式段落 {i + 1} 格式正确"
                report['formula_detection']['messages'].append(success_msg)
                for msg in para_report['messages']:
                    if "检查通过" in msg or "正确" in msg:
                        report['formula_detection']['messages'].append(f"  - {msg}")
        
        # 生成总结
        if all_formulas_ok:
            summary_msg = template.get('messages', {}).get('formula_complete_ok', 
                                                         f"所有 {len(formula_paragraphs)} 个公式段落格式检查通过")
            report['summary'].append(summary_msg)
        else:
            failed_count = sum(1 for para in report['details']['formula_paragraphs'] 
                             if not para['format_check']['ok'])
            summary_msg = f"发现 {failed_count} 个公式段落格式问题，共检查 {len(formula_paragraphs)} 个公式段落"
            report['summary'].append(summary_msg)
        
        # 添加检测统计
        stats_msg = f"检测统计：共 {report['details']['total_paragraphs']} 个段落，识别出 {len(formula_paragraphs)} 个可能的公式段落"
        report['summary'].append(stats_msg)
        
        # 检查公式编号连续性
        formula_numbers = []
        for para_info in report['details']['formula_paragraphs']:
            # 从文本预览中提取编号
            text_preview = para_info['text_preview']
            
            # 查找编号模式
            number_matches = re.findall(r'\((\d+)\)', text_preview)
            if number_matches:
                try:
                    formula_numbers.append(int(number_matches[-1]))  # 取最后一个匹配的编号
                except ValueError:
                    pass
        
        # 检查编号连续性
        if len(formula_numbers) > 1:
            formula_numbers.sort()
            expected_numbers = list(range(1, len(formula_numbers) + 1))
            
            if formula_numbers != expected_numbers:
                report['formula_detection']['ok'] = False
                report['formula_detection']['messages'].append(f"公式编号不连续或不从1开始：检测到编号 {formula_numbers}，期望 {expected_numbers}")
                
                # 添加到总结
                report['summary'].append(f"公式编号问题：应从(1)开始连续编号，当前为 {formula_numbers}")
            else:
                report['formula_detection']['messages'].append(f"公式编号连续性正确：{formula_numbers}")
        elif len(formula_numbers) == 1:
            if formula_numbers[0] != 1:
                report['formula_detection']['ok'] = False
                report['formula_detection']['messages'].append(f"单个公式编号应为(1)，实际为({formula_numbers[0]})")
                report['summary'].append(f"公式编号问题：单个公式应编号为(1)，当前为({formula_numbers[0]})")
            else:
                report['formula_detection']['messages'].append("公式编号正确：(1)")
        
        # 检查是否有可能的公式但没有使用Word公式功能
        potential_formula_suggestions = []
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            
            # 跳过已识别的公式段落
            is_identified_formula = any(fp['paragraph'] == paragraph for fp in formula_paragraphs)
            if is_identified_formula:
                continue
            
            text = paragraph.text.strip()
            
            # 检测可能是公式的模式
            if len(text) < 200:  # 限制段落长度
                potential_patterns = [
                    r'[=<>≤≥≠±∞∑∏∫]',  # 包含数学符号
                    r'[α-ωΑ-Ω]',         # 包含希腊字母
                    r'[₀-₉⁰-⁹]',         # 包含上下标
                    r'[∂∇∆]',           # 微分符号
                ]
                
                has_math_pattern = any(re.search(pattern, text) for pattern in potential_patterns)
                
                # 如果包含数学模式但不是真正的公式段落，给出建议
                if has_math_pattern:
                    potential_formula_suggestions.append({
                        'text': text[:100] + ('...' if len(text) > 100 else ''),
                        'paragraph_index': doc.paragraphs.index(paragraph) + 1
                    })
        
        # 添加建议到报告中
        if potential_formula_suggestions:
            report['formula_detection']['messages'].append("\n发现可能的公式内容，建议使用Word的插入公式功能：")
            for suggestion in potential_formula_suggestions[:3]:  # 最多显示3个
                report['formula_detection']['messages'].append(f"  - 段落{suggestion['paragraph_index']}: {suggestion['text']}")
            
            if len(potential_formula_suggestions) > 3:
                report['formula_detection']['messages'].append(f"  - 还有 {len(potential_formula_suggestions) - 3} 个类似段落...")
            
            report['formula_detection']['messages'].append("  建议：在Word中选中这些内容，使用 插入→公式 功能重新创建")
        
        return report
        
    except Exception as e:
        return {
            'formula_detection': {
                'ok': False, 
                'messages': [f"公式检测过程中发生异常: {str(e)}"]
            },
            'summary': [f"检查失败: {str(e)}"],
            'details': {}
        }

def print_report(report):
    """
    打印公式检测报告
    """
    print("=== 公式格式检查报告 ===")
    
    # 显示总体状态
    formula_ok = report.get('formula_detection', {}).get('ok', False)
    print(f"--- 公式检测 ---")
    print(f" OK: {formula_ok}")
    
    # 显示详细消息
    messages = report.get('formula_detection', {}).get('messages', [])
    for msg in messages:
        print(f"  - {msg}")
    
    # 显示检测详情
    details = report.get('details', {})
    if details:
        print(f"\n--- 检测详情 ---")
        print(f"总段落数: {details.get('total_paragraphs', 0)}")
        print(f"公式段落数: {details.get('formula_paragraphs_count', 0)}")
        
        # 显示每个公式段落的详细信息
        formula_paras = details.get('formula_paragraphs', [])
        for para in formula_paras:
            print(f"\n公式段落 {para['index']} (置信度: {para['confidence_score']}/10):")
            print(f"  内容预览: {para['text_preview']}")
            print(f"  Office Math对象: {'是' if para['has_math_object'] else '否'} ({para['math_object_count']}个)")
            print(f"  公式制表位: {'是' if para['has_formula_tab_stops'] else '否'}")
            print(f"  公式样式: {'是' if para['has_formula_style'] else '否'}")
            
            format_check = para.get('format_check', {})
            print(f"  格式检查: {'通过' if format_check.get('ok', False) else '失败'}")
            
            # 显示格式检查的详细信息
            if format_check.get('messages'):
                for msg in format_check['messages']:
                    print(f"    - {msg}")
    
    # 显示总结
    print(f"\n--- 总结 ---")
    summary = report.get('summary', [])
    for s in summary:
        print(f" {s}")

def print_help():
    """
    显示帮助信息
    """
    print("用法:")
    print("  python Formula_detect.py check <paper.docx> <template.json_or_name>")
    print("")
    print("示例:")
    print("  python Formula_detect.py check template/test.docx Formula")
    print("  python Formula_detect.py check template/test.docx templates/Formula.json")
    print("")
    print("说明:")
    print("  检查Word文档中的公式格式，包括：")
    print("  - 制表位设置（20字符居中，40字符右对齐）")
    print("  - 公式内容字体（Cambria Math）")
    print("  - 编号字体（Times New Roman）")
    print("  - 字体大小（小四号12pt）")

# ---------- CLI接口 ----------
if __name__ == '__main__':
    if len(sys.argv) != 4:
        print_help()
        sys.exit(0)
    
    cmd = sys.argv[1]
    if cmd == 'check':
        paper_path = sys.argv[2]
        tpl_id = sys.argv[3]
        
        # 检查文件是否存在
        if not os.path.isfile(paper_path):
            print(f"论文文件不存在: {paper_path}")
            sys.exit(1)
        
        try:
            # 执行检查
            report = check_doc_with_template(paper_path, tpl_id)
            
            # 打印报告
            print_report(report)
            
        except Exception as e:
            print("检查时出错:", e)
            sys.exit(1)
    else:
        print_help()
        sys.exit(0)

'''
使用示例:
python paper_detect\Formula_detect.py check template\test.docx Formula
python paper_detect\Formula_detect.py check template\test.docx templates\Formula.json
'''
