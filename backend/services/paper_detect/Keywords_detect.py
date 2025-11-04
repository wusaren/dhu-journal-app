#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# 导入Title_detect的作者和标题提取功能
try:
    from paper_detect.Title_detect import extract_from_docx, parse_authors_by_regex, split_authors_block, get_nonempty_paragraphs
    TITLE_DETECT_AVAILABLE = True
except ImportError:
    TITLE_DETECT_AVAILABLE = False
    print("警告: 无法导入Title_detect模块，脚注中的作者和标题比对功能将被禁用")

"""
=== 论文格式检测系统 - Keywords、CLC、脚注检测器 ===

本检测器包含以下检测功能：
【一、关键词检测 (Keywords Detection)】
1. Keywords字段后紧接冒号，冒号后紧接关键词内容
2. 关键词用分号(;)分割，而非逗号或冒号
3. 字体为Times New Roman，五号正体(10.5pt)，1.5倍行间距
4. Keywords标题加粗，关键词内容正体
5. 双端对齐
6. 支持检测缺少冒号的情况并给出错误提示

【二、CLC和Document Code检测】
1. 关键词后一行为CLC number和Document code
2. 格式：CLC number: xxx    Document code: A（同一行，中间8-15个空格）
3. 字体为Times New Roman，五号(10.5pt)
4. CLC number和Document code标签加粗
5. 支持检测多行情况并报错（应在同一行）
6. 支持检测空格数量不正确的情况

【三、Word脚注检测 (Real Footnotes Detection)】
1. 检测通过Word"插入->脚注"功能创建的真正脚注
2. 脚注内容包含四个必需项目：
   - Received date: （投稿日期，格式：YYYY-MM-DD）
   - Foundation item: xxxxxx (No. xxxxxx)
   - * Correspondence should be addressed to NAME, email: xxx (注意：不是Corresponding)
   - Citation: AUTHOR1, AUTHOR2, AUTHOR3, et al. Title [J]. Journal of Donghua University (English Edition), YEAR, VOL(ISSUE): PAGE-PAGE
3. 格式要求：
   - 字体：Times New Roman，小五号(9pt)，单倍行距，左对齐
   - Journal of Donghua University (English Edition)必须为斜体
   - 其他内容为正体
4. Citation格式要求：
   - 作者姓名缩写格式（三个人名后跟et al），支持SURNAME I或SURNAME I N格式
   - 作者列表必须与论文标题下的作者一致（顺序和姓名）
   - 标题必须与论文标题一致，但使用sentence case（仅首字母大写）
   - Journal名称固定且斜体
5. 新增检测：
   - 检测Correspondence拼写错误（如写成Corresponding）
   - 自动提取论文作者并与脚注作者比对
   - 自动提取论文标题并与脚注标题比对（考虑大小写转换）

【检测技术特性】
- 支持段落样式和直接格式的混合检测
- 支持XML级别的脚注格式检测
- 智能错误提示和格式建议
- 完整的7维度检测报告
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 字体检测函数 ----------
def detect_font_for_run(run, paragraph=None):
    """
    更健壮的字体检测函数，同时检查直接格式和段落样式。
    返回 (font_size_pt, font_name, is_bold, is_italic, line_spacing)
    """
    font_source = run.font if run else (paragraph.style.font if paragraph else None)
    if not font_source:
        return 12.0, "Unknown", False, False, 1.0

    font_name = font_source.name if font_source.name else "Times New Roman"

    font_size = None
    try:
        # 优先级1: 检查run的直接格式
        if run and run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        # 优先级2: 检查run的XML元素
        if font_size is None and run and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
        
        # 优先级3: 检查段落样式的直接格式
        if font_size is None and paragraph and paragraph.style and hasattr(paragraph.style.font, 'size') and hasattr(paragraph.style.font.size, 'pt'):
            font_size = float(paragraph.style.font.size.pt)
            
        # 优先级4: 检查段落样式的XML元素
        if font_size is None and paragraph and hasattr(paragraph.style, 'element'):
            sz_nodes = paragraph.style.element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass

    font_size = font_size if font_size is not None else 12.0

    is_bold = font_source.bold if font_source.bold is not None else False
    is_italic = font_source.italic if font_source.italic is not None else False
    if run and run.font:
        is_bold = run.font.bold if run.font.bold is not None else is_bold
        is_italic = run.font.italic if run.font.italic is not None else is_italic
    
    # 确保返回明确的布尔值，而不是None
    is_bold = bool(is_bold) if is_bold is not None else False
    is_italic = bool(is_italic) if is_italic is not None else False

    # 行间距检测
    line_spacing = 1.0  # 默认单倍行距
    try:
        if paragraph and paragraph.paragraph_format.line_spacing:
            line_spacing = float(paragraph.paragraph_format.line_spacing)
    except Exception:
        pass

    return font_size, font_name, is_bold, is_italic, line_spacing

def get_font_size(pt_size, tpl=None):
    """字体大小转换为中文字号"""
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        size_map = {
            9: "小五", 10.5: "五号", 12: "小四", 14: "四号",
            16: "三号", 18: "小二", 22: "二号", 24: "小一", 26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

def get_line_spacing_name(spacing, tpl=None):
    """行间距转换为中文描述"""
    if tpl and 'check_rules' in tpl and 'line_spacing_mapping' in tpl['check_rules']:
        spacing_config = tpl['check_rules']['line_spacing_mapping']
        spacing_map = {}
        for key, value in spacing_config.items():
            try:
                spacing_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        spacing_map = {
            1.0: "单倍行距", 1.15: "1.15倍行距", 
            1.5: "1.5倍行距", 2.0: "双倍行距"
        }
    
    if not spacing_map:
        return f"{spacing}倍行距"
    
    closest_spacing = min(spacing_map.keys(), key=lambda x: abs(x - spacing))
    return spacing_map[closest_spacing]

def get_alignment_name(alignment_value, tpl=None):
    """段落对齐方式转换为中文描述"""
    alignment_map = {
        0: "左对齐", 1: "居中对齐", 2: "右对齐", 3: "两端对齐"
    }
    return alignment_map.get(alignment_value, f"未知对齐({alignment_value})")

def detect_paragraph_alignment(paragraph):
    """
    段落对齐检测（基于Abstract_detect.py）
    """
    # 优先级1: 检查直接格式
    direct_alignment = paragraph.paragraph_format.alignment
    if direct_alignment is not None:
        return int(direct_alignment)
    
    # 优先级2: 如果直接格式为None，检查样式格式
    if paragraph.style:
        try:
            style_alignment = paragraph.style.paragraph_format.alignment
            if style_alignment is not None:
                return int(style_alignment)
        except Exception:
            pass
    
    # 默认值: 左对齐
    return WD_PARAGRAPH_ALIGNMENT.LEFT

def detect_paragraph_indent(paragraph):
    """检测段落缩进"""
    try:
        fmt = paragraph.paragraph_format
        first_line_indent = fmt.first_line_indent.pt if fmt.first_line_indent else 0.0
        left_indent = fmt.left_indent.pt if fmt.left_indent else 0.0
        right_indent = fmt.right_indent.pt if fmt.right_indent else 0.0
        return first_line_indent, left_indent, right_indent
    except Exception:
        return 0.0, 0.0, 0.0

# ---------- 关键词检测逻辑 ----------

def check_keywords_structure(text, tpl):
    """
    检查关键词结构（Keywords:冒号格式和分隔符）
    返回 {'ok': bool, 'messages': [], 'content': str, 'keywords_list': []}
    """
    report = {'ok': True, 'messages': [], 'content': '', 'keywords_list': []}
    
    # 检查Keywords:格式
    structure_rules = tpl.get('structure_rules', {})
    header_pattern = structure_rules.get('header_pattern', r'^\\s*Keywords\\s*:\\s*(.+)$')
    
    try:
        match = re.match(header_pattern, text.strip(), re.DOTALL | re.IGNORECASE)
        if match:
            report['content'] = match.group(1).strip()
            ok_msg = tpl.get('messages', {}).get('structure_header_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
        else:
            # 检查是否是缺少冒号的情况
            keywords_without_colon = re.match(r'^\s*Keywords\s+(.+)$', text.strip(), re.DOTALL | re.IGNORECASE)
            if keywords_without_colon:
                # 找到Keywords但缺少冒号
                report['ok'] = False
                error_msg = tpl.get('messages', {}).get('structure_missing_colon')
                if error_msg:
                    report['messages'].append(error_msg)
                # 仍然提取内容以便后续检测
                report['content'] = keywords_without_colon.group(1).strip()
            else:
                # 完全不匹配Keywords格式
                report['ok'] = False
                error_msg = tpl.get('messages', {}).get('structure_header_error')
                if error_msg:
                    report['messages'].append(error_msg)
                return report
    except Exception:
        report['ok'] = False
        report['messages'].append("关键词格式检查出错")
        return report
    
    # 检查关键词分隔符（冒号分割）
    if report['content']:
        separator_pattern = structure_rules.get('separator_pattern', r'\\s*:\\s*')
        
        # 分割关键词
        try:
            keywords_raw = re.split(separator_pattern, report['content'])
            keywords_list = [kw.strip() for kw in keywords_raw if kw.strip()]
            report['keywords_list'] = keywords_list
            
            # 检查是否正确使用分号分割
            has_semicolon = ';' in report['content'] or '；' in report['content']
            has_comma = ',' in report['content']
            has_colon_in_content = ':' in report['content'].replace('Keywords:', '', 1)  # 排除Keywords:后的冒号
            
            if has_semicolon:
                # 如果使用了分号，检查是否混用了逗号
                if has_comma:
                    report['ok'] = False
                    error_msg = tpl.get('messages', {}).get('structure_separator_mixed_error', 
                                                            '关键词分隔符混用：应统一使用分号分割，不应混用逗号')
                    report['messages'].append(error_msg)
                else:
                    # 只使用分号，正确
                    ok_msg = tpl.get('messages', {}).get('structure_separator_ok')
                    if ok_msg:
                        report['messages'].append(ok_msg)
            else:
                # 没有使用分号，检查是否错误使用了其他分隔符
                if has_comma or has_colon_in_content:
                    report['ok'] = False
                    error_msg = tpl.get('messages', {}).get('structure_separator_error')
                    if error_msg:
                        report['messages'].append(error_msg)
        except Exception:
            report['ok'] = False
            report['messages'].append("关键词分隔符检查出错")
            return report
    
    # 检查关键词数量
    if report['keywords_list']:
        keywords_count = len(report['keywords_list'])
        min_count = structure_rules.get('min_keywords_count', 3)
        max_count = structure_rules.get('max_keywords_count', 8)
        
        if keywords_count < min_count:
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('structure_count_few')
            if msg_tpl:
                try:
                    report['messages'].append(msg_tpl.format(count=keywords_count, min=min_count))
                except:
                    report['messages'].append(msg_tpl)
        elif keywords_count > max_count:
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('structure_count_many')
            if msg_tpl:
                try:
                    report['messages'].append(msg_tpl.format(count=keywords_count, max=max_count))
                except:
                    report['messages'].append(msg_tpl)
        else:
            ok_msg = tpl.get('messages', {}).get('structure_count_ok')
            if ok_msg:
                try:
                    report['messages'].append(ok_msg.format(count=keywords_count))
                except:
                    report['messages'].append(ok_msg)
    
    # 检查内容长度
    if report['content']:
        content_length = len(report['content'])
        min_length = structure_rules.get('min_content_length', 10)
        max_length = structure_rules.get('max_content_length', 500)
        
        if content_length < min_length:
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('structure_length_short')
            if msg_tpl:
                try:
                    report['messages'].append(msg_tpl.format(min=min_length))
                except:
                    report['messages'].append(msg_tpl)
        elif content_length > max_length:
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('structure_length_long')
            if msg_tpl:
                try:
                    report['messages'].append(msg_tpl.format(max=max_length))
                except:
                    report['messages'].append(msg_tpl)
        else:
            ok_msg = tpl.get('messages', {}).get('structure_length_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
    
    return report

def check_keywords_paragraphs(doc, tpl):
    """
    检查关键词段落查找
    返回 {'ok': bool, 'messages': [], 'keywords_paragraph': paragraph}
    """
    report = {'ok': True, 'messages': [], 'keywords_paragraph': None}
    
    # 查找包含Keywords的段落（宽松匹配，不要求冒号）
    keywords_paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text and re.search(r'\bKeywords\b', paragraph.text, re.IGNORECASE):
            keywords_paragraphs.append(paragraph)
    
    if len(keywords_paragraphs) == 1:
        report['keywords_paragraph'] = keywords_paragraphs[0]
        ok_msg = tpl.get('messages', {}).get('structure_header_ok')
        if ok_msg:
            report['messages'].append("找到关键词段落")
    elif len(keywords_paragraphs) > 1:
        report['ok'] = False
        report['messages'].append("检测到多个Keywords段落，应该只有一个")
        # 选择第一个作为主要段落
        report['keywords_paragraph'] = keywords_paragraphs[0]
    else:
        report['ok'] = False
        report['messages'].append("未找到Keywords段落")
    
    return report

def check_clc_document_structure(doc, keywords_paragraph_index, tpl):
    """
    检查CLC number和Document code结构（应在关键词后一行）
    返回 {'ok': bool, 'messages': [], 'clc_content': str, 'document_content': str, 'clc_paragraph': paragraph}
    """
    report = {'ok': True, 'messages': [], 'clc_content': '', 'document_content': '', 'clc_paragraph': None}
    
    # 查找关键词段落后的下一段落
    paragraphs = [p for p in doc.paragraphs if p.text and p.text.strip()]
    
    if keywords_paragraph_index is None or keywords_paragraph_index >= len(paragraphs) - 1:
        report['ok'] = False
        error_msg = tpl.get('messages', {}).get('clc_document_missing')
        if error_msg:
            report['messages'].append(error_msg)
        return report
    
    # 检查CLC number和Document code格式（支持单行和多行）
    structure_rules = tpl.get('structure_rules', {})
    clc_pattern = structure_rules.get('clc_document_pattern', r'^\\s*CLC number\\s*:\\s*(.+?)\\s+Document code\\s*:\\s*(.+)$')
    clc_multiline_pattern = structure_rules.get('clc_document_multiline_pattern', r'^\\s*CLC number\\s*:\\s*(.+?)$')
    document_pattern = structure_rules.get('document_code_pattern', r'^\\s*Document code\\s*:\\s*(.+)$')
    
    try:
        found_clc = False
        found_document = False
        single_line_success = False  # 新增：标记单行匹配是否成功
        clc_content = ""
        document_content = ""
        clc_paragraph = None
        
        # 检查后续最多3个段落
        for i in range(keywords_paragraph_index + 1, min(keywords_paragraph_index + 4, len(paragraphs))):
            current_paragraph = paragraphs[i]
            current_text = current_paragraph.text.strip()
            
            if not current_text:
                continue
            
            # 首先尝试单行匹配（CLC和Document在同一行）
            single_line_match = re.match(clc_pattern, current_text, re.IGNORECASE)
            if single_line_match:
                report['clc_content'] = single_line_match.group(1).strip()
                report['document_content'] = single_line_match.group(2).strip()
                report['clc_paragraph'] = current_paragraph
                single_line_success = True  # 标记单行匹配成功
                
                # 检查空格数量
                spacing_match = re.search(r'CLC number\s*:\s*(.+?)(\s+)Document code\s*:\s*(.+)', current_text, re.IGNORECASE)
                if spacing_match:
                    spacing = spacing_match.group(2)
                    spacing_count = len(spacing)
                    min_spacing = structure_rules.get('clc_spacing_min', 8)
                    max_spacing = structure_rules.get('clc_spacing_max', 15)
                    
                    if spacing_count < min_spacing or spacing_count > max_spacing:
                        report['ok'] = False
                        error_msg = tpl.get('messages', {}).get('clc_document_spacing_error')
                        if error_msg:
                            report['messages'].append(error_msg)
                        ok_msg = tpl.get('messages', {}).get('clc_document_structure_ok')
                        if ok_msg:
                            report['messages'].append(ok_msg + "（但空格数量不符合要求）")
                    else:
                        ok_msg = tpl.get('messages', {}).get('clc_document_structure_ok')
                        if ok_msg:
                            report['messages'].append(ok_msg)
                break
            
            # 多行匹配：先找CLC number
            if not found_clc:
                clc_match = re.match(clc_multiline_pattern, current_text, re.IGNORECASE)
                if clc_match:
                    clc_content = clc_match.group(1).strip()
                    clc_paragraph = current_paragraph
                    found_clc = True
                    continue
            
            # 多行匹配：再找Document code
            if found_clc and not found_document:
                # 检查同一段落的其他行（通过换行分割）
                lines = current_text.split('\n')
                for line in lines:
                    doc_match = re.match(document_pattern, line.strip(), re.IGNORECASE)
                    if doc_match:
                        document_content = doc_match.group(1).strip()
                        found_document = True
                        break
                
                # 如果同一段落没找到，检查下一段落
                if not found_document:
                    doc_match = re.match(document_pattern, current_text, re.IGNORECASE)
                    if doc_match:
                        document_content = doc_match.group(1).strip()
                        found_document = True
        
        # 处理匹配结果
        if single_line_success:
            # 单行匹配已经在上面处理了，这里不需要额外处理
            pass
        elif found_clc and found_document:
            # 多行格式应该报错
            report['ok'] = False
            report['clc_content'] = clc_content
            report['document_content'] = document_content
            report['clc_paragraph'] = clc_paragraph
            error_msg = tpl.get('messages', {}).get('clc_document_multiline_error')
            if error_msg:
                report['messages'].append(error_msg)
        elif found_clc or found_document:
            # 只找到一部分
            report['ok'] = False
            if found_clc and not found_document:
                report['messages'].append("找到CLC number但未找到Document code")
            elif found_document and not found_clc:
                report['messages'].append("找到Document code但未找到CLC number")
        else:
            # 检查是否有相关关键词但格式不正确
            found_keywords = False
            for i in range(keywords_paragraph_index + 1, min(keywords_paragraph_index + 4, len(paragraphs))):
                if i < len(paragraphs):
                    text = paragraphs[i].text
                    if 'CLC number' in text or 'Document code' in text:
                        found_keywords = True
                        break
            
            if found_keywords:
                report['ok'] = False
                error_msg = tpl.get('messages', {}).get('clc_document_structure_error')
                if error_msg:
                    report['messages'].append(error_msg)
            else:
                report['ok'] = False
                error_msg = tpl.get('messages', {}).get('clc_document_missing')
                if error_msg:
                    report['messages'].append(error_msg)
    
    except Exception:
        report['ok'] = False
        report['messages'].append("CLC number和Document code格式检查出错")
    
    return report

def check_clc_document_format(paragraph, tpl):
    """
    检查CLC number和Document code格式（字体、加粗等）
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not paragraph or not paragraph.runs:
        report['ok'] = False
        report['messages'].append("CLC number和Document code段落没有文本内容")
        return report
    
    # 获取格式规则
    format_rules = tpl.get('format_rules', {}).get('clc_document', {})
    
    issues = []
    
    # 分析段落中的所有runs，寻找CLC number和Document code部分
    clc_runs = []  # CLC number部分的runs
    document_runs = []  # Document code部分的runs
    
    # 查找CLC number和Document code位置
    full_text = paragraph.text
    clc_match = re.search(r'CLC number\s*:', full_text, re.IGNORECASE)
    document_match = re.search(r'Document code\s*:', full_text, re.IGNORECASE)
    
    if clc_match and document_match:
        clc_end_pos = clc_match.end()
        document_start_pos = document_match.start()
        
        # 遍历runs，确定哪些属于CLC，哪些属于Document
        current_pos = 0
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if run_start < clc_end_pos:
                # 这个run包含CLC number标签部分
                clc_runs.append(run)
            elif run_start >= document_start_pos:
                # 这个run是Document code部分
                document_runs.append(run)
            
            current_pos = run_end
    else:
        # 没有找到标准格式，将所有runs作为检查对象
        clc_runs = paragraph.runs[:len(paragraph.runs)//2] if paragraph.runs else []
        document_runs = paragraph.runs[len(paragraph.runs)//2:] if paragraph.runs else []
    
    # 检查CLC number标签格式
    if clc_runs:
        clc_run = clc_runs[0]  # 取第一个作为代表
        clc_size_pt, clc_font_name, clc_bold, clc_italic, clc_line_spacing = detect_font_for_run(clc_run, paragraph)
        
        # 检查CLC number标签是否加粗
        expected_bold_labels = format_rules.get('bold_labels', True)
        if clc_bold != expected_bold_labels:
            bold_status = "加粗" if expected_bold_labels else "不加粗"
            actual_status = "加粗" if clc_bold else "不加粗"
            issues.append(f"CLC number标签应为{bold_status}，实际为{actual_status}")
    
    # 检查Document code标签格式
    if document_runs:
        document_run = document_runs[0]  # 取第一个作为代表
        doc_size_pt, doc_font_name, doc_bold, doc_italic, doc_line_spacing = detect_font_for_run(document_run, paragraph)
        
        # 检查Document code标签是否加粗
        expected_bold_labels = format_rules.get('bold_labels', True)
        if doc_bold != expected_bold_labels:
            bold_status = "加粗" if expected_bold_labels else "不加粗"
            actual_status = "加粗" if doc_bold else "不加粗"
            issues.append(f"Document code标签应为{bold_status}，实际为{actual_status}")
    
    # 统一格式检查（使用第一个run作为基准）
    if clc_runs or document_runs:
        check_run = clc_runs[0] if clc_runs else document_runs[0]
        actual_size_pt, actual_font_name, actual_bold, actual_italic, actual_line_spacing = detect_font_for_run(check_run, paragraph)
        
        # 字体大小检查
        if 'font_size_pt' in format_rules:
            expected_size_pt = float(format_rules['font_size_pt'])
            actual_size_name = get_font_size(actual_size_pt, tpl)
            expected_size_name = get_font_size(expected_size_pt, tpl)
            print(f"CLC/Document字体大小: {actual_size_name}（{actual_size_pt}pt）(期望: {expected_size_name}（{expected_size_pt}pt）)")
            if abs(actual_size_pt - expected_size_pt) > 0.5:
                issues.append(f"字体大小应为{expected_size_name}（{expected_size_pt}pt），实际为{actual_size_name}（{actual_size_pt}pt）")
        
        # 字体名称检查
        if 'font_name' in format_rules:
            expected_font_name = str(format_rules['font_name'])
            print(f"CLC/Document字体名称: {actual_font_name} (期望: {expected_font_name})")
            if expected_font_name.lower() not in actual_font_name.lower():
                issues.append(f"字体应为{expected_font_name}，实际为{actual_font_name}")
        
        # 斜体检查
        if 'italic' in format_rules:
            expected_italic = bool(format_rules['italic'])
            print(f"CLC/Document斜体: {'是' if actual_italic else '否'} (期望: {'是' if expected_italic else '否'})")
            if actual_italic != expected_italic:
                italic_status = "斜体" if expected_italic else "正体"
                actual_status = "斜体" if actual_italic else "正体"
                issues.append(f"字体应为{italic_status}，实际为{actual_status}")
    
    print(f"发现 {len(issues)} 个CLC/Document格式问题")
    print("---")
    
    if issues:
        report['ok'] = False
        header = tpl.get('messages', {}).get('clc_document_format_error')
        if header:
            report['messages'].append(header + "：")
        report['messages'].extend([f"  - {i}" for i in issues])
    else:
        ok_msg = tpl.get('messages', {}).get('clc_document_format_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
    
    return report

def check_keywords_format(paragraph, tpl):
    """
    检查关键词格式（字体、加粗、行间距等，包括混合格式）
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not paragraph or not paragraph.runs:
        report['ok'] = False
        report['messages'].append("关键词段落没有文本内容")
        return report
    
    # 获取格式规则
    format_rules = tpl.get('format_rules', {}).get('keywords', {})
    
    issues = []
    
    # 分析段落中的所有runs，寻找Keywords标题和内容部分
    keywords_title_runs = []  # Keywords标题部分的runs
    content_runs = []  # 关键词内容部分的runs
    
    # 查找Keywords标题位置
    full_text = paragraph.text
    keywords_match = re.search(r'\bKeywords\s*:', full_text, re.IGNORECASE)
    
    if keywords_match:
        keywords_end_pos = keywords_match.end()
        
        # 遍历runs，确定哪些属于标题，哪些属于内容
        current_pos = 0
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if run_start < keywords_end_pos:
                # 这个run包含Keywords标题部分
                keywords_title_runs.append(run)
            else:
                # 这个run是内容部分
                content_runs.append(run)
            
            current_pos = run_end
    else:
        # 没有找到Keywords标题，将第一个run作为标题，其余作为内容
        if paragraph.runs:
            keywords_title_runs = [paragraph.runs[0]]
            content_runs = paragraph.runs[1:] if len(paragraph.runs) > 1 else []
    
    # 检查Keywords标题部分格式
    if keywords_title_runs:
        title_run = keywords_title_runs[0]  # 取第一个作为代表
        title_size_pt, title_font_name, title_bold, title_italic, title_line_spacing = detect_font_for_run(title_run, paragraph)
        
        # 检查标题是否加粗
        expected_bold_title = format_rules.get('bold_title', True)
        if title_bold != expected_bold_title:
            bold_status = "加粗" if expected_bold_title else "不加粗"
            actual_status = "加粗" if title_bold else "不加粗"
            issues.append(f"Keywords标题应为{bold_status}，实际为{actual_status}")
    
    # 检查关键词内容部分格式
    if content_runs:
        content_run = content_runs[0]  # 取第一个作为代表
        content_size_pt, content_font_name, content_bold, content_italic, content_line_spacing = detect_font_for_run(content_run, paragraph)
        
        # 检查内容是否为正体（不加粗）
        expected_bold_content = format_rules.get('bold_content', False)
        if content_bold != expected_bold_content:
            bold_status = "加粗" if expected_bold_content else "正体"
            actual_status = "加粗" if content_bold else "正体"
            issues.append(f"关键词内容应为{bold_status}，实际为{actual_status}")
    else:
        # 如果没有内容runs，使用标题run作为整体检查
        content_run = keywords_title_runs[0] if keywords_title_runs else None
        if content_run:
            content_size_pt, content_font_name, content_bold, content_italic, content_line_spacing = detect_font_for_run(content_run, paragraph)
    
    # 统一格式检查（使用内容run的格式作为基准）
    if content_runs or keywords_title_runs:
        # 选择用于检查的run
        check_run = content_runs[0] if content_runs else keywords_title_runs[0]
        actual_size_pt, actual_font_name, actual_bold, actual_italic, actual_line_spacing = detect_font_for_run(check_run, paragraph)
        
        # 字体大小检查
        if 'font_size_pt' in format_rules:
            expected_size_pt = float(format_rules['font_size_pt'])
            actual_size_name = get_font_size(actual_size_pt, tpl)
            expected_size_name = get_font_size(expected_size_pt, tpl)
            print(f"字体大小: {actual_size_name}（{actual_size_pt}pt）(期望: {expected_size_name}（{expected_size_pt}pt）)")
            if abs(actual_size_pt - expected_size_pt) > 0.5:
                issues.append(f"字体大小应为{expected_size_name}（{expected_size_pt}pt），实际为{actual_size_name}（{actual_size_pt}pt）")
        
        # 字体名称检查
        if 'font_name' in format_rules:
            expected_font_name = str(format_rules['font_name'])
            print(f"字体名称: {actual_font_name} (期望: {expected_font_name})")
            if expected_font_name.lower() not in actual_font_name.lower():
                issues.append(f"字体应为{expected_font_name}，实际为{actual_font_name}")
        
        # 斜体检查
        if 'italic' in format_rules:
            expected_italic = bool(format_rules['italic'])
            print(f"斜体: {'是' if actual_italic else '否'} (期望: {'是' if expected_italic else '否'})")
            if actual_italic != expected_italic:
                italic_status = "斜体" if expected_italic else "正体"
                actual_status = "斜体" if actual_italic else "正体"
                issues.append(f"字体应为{italic_status}，实际为{actual_status}")
        
        # 行间距检查
        if 'line_spacing' in format_rules:
            expected_line_spacing = float(format_rules['line_spacing'])
            actual_spacing_name = get_line_spacing_name(actual_line_spacing, tpl)
            expected_spacing_name = get_line_spacing_name(expected_line_spacing, tpl)
            print(f"行间距: {actual_spacing_name}（{actual_line_spacing}倍）(期望: {expected_spacing_name}（{expected_line_spacing}倍）)")
            if abs(actual_line_spacing - expected_line_spacing) > 0.1:
                issues.append(f"行间距应为{expected_spacing_name}（{expected_line_spacing}倍），实际为{actual_spacing_name}（{actual_line_spacing}倍）")
        
        # 段落对齐检查
        if 'alignment' in format_rules:
            expected_alignment_str = str(format_rules['alignment'])
            alignment_map = {"left": 0, "center": 1, "right": 2, "justify": 3}
            expected_alignment = alignment_map.get(expected_alignment_str, 0)
            
            actual_alignment = detect_paragraph_alignment(paragraph)
            actual_alignment_name = get_alignment_name(actual_alignment, tpl)
            expected_alignment_name = get_alignment_name(expected_alignment, tpl)
            print(f"段落对齐: {actual_alignment_name} (期望: {expected_alignment_name})")
            if actual_alignment != expected_alignment:
                issues.append(f"段落应为{expected_alignment_name}，实际为{actual_alignment_name}")
        
        # 段落缩进检查
        if 'first_line_indent' in format_rules or 'left_indent' in format_rules or 'right_indent' in format_rules:
            first_line_indent, left_indent, right_indent = detect_paragraph_indent(paragraph)
            
            if 'first_line_indent' in format_rules:
                expected_first_indent = float(format_rules['first_line_indent'])
                print(f"首行缩进: {first_line_indent:.1f}pt (期望: {expected_first_indent}pt)")
                if abs(first_line_indent - expected_first_indent) > 1.0:
                    issues.append(f"首行缩进应为{expected_first_indent}pt，实际为{first_line_indent:.1f}pt")
            
            if 'left_indent' in format_rules:
                expected_left_indent = float(format_rules['left_indent'])
                print(f"左缩进: {left_indent:.1f}pt (期望: {expected_left_indent}pt)")
                if abs(left_indent - expected_left_indent) > 1.0:
                    issues.append(f"左缩进应为{expected_left_indent}pt，实际为{left_indent:.1f}pt")
            
            if 'right_indent' in format_rules:
                expected_right_indent = float(format_rules['right_indent'])
                print(f"右缩进: {right_indent:.1f}pt (期望: {expected_right_indent}pt)")
                if abs(right_indent - expected_right_indent) > 1.0:
                    issues.append(f"右缩进应为{expected_right_indent}pt，实际为{right_indent:.1f}pt")
    
    print(f"发现 {len(issues)} 个格式问题")
    print("---")
    
    if issues:
        report['ok'] = False
        header = tpl.get('messages', {}).get('format_keywords_issue_header')
        if header:
            report['messages'].append(header)
        report['messages'].extend([f"  - {i}" for i in issues])
    else:
        ok_msg = tpl.get('messages', {}).get('format_keywords_ok')
        if ok_msg:
            report['messages'].append(ok_msg)
        
        # 检查混合格式是否正确
        if keywords_title_runs and content_runs:
            mixed_ok_msg = tpl.get('messages', {}).get('format_mixed_ok')
            if mixed_ok_msg:
                report['messages'].append(mixed_ok_msg)
    
    return report

def check_keywords_with_template(doc_path, template_identifier):
    """
    主检查函数：检查关键词格式、CLC/Document code和脚注
    """
    tpl = load_template(template_identifier)
    doc = Document(doc_path)
    
    # 查找关键词段落（宽松匹配，不要求冒号）
    keywords_text = ""
    keywords_paragraph_index = None
    paragraphs = [p for p in doc.paragraphs if p.text and p.text.strip()]
    
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text and re.search(r'\bKeywords\b', paragraph.text, re.IGNORECASE):
            keywords_text = paragraph.text.strip()
            keywords_paragraph_index = i
            break
    
    print(f"keywords_text: {keywords_text}")
    if not keywords_text:
        return {
            'structure': {'ok': False, 'messages': ['未找到Keywords段落']},
            'paragraphs': {'ok': False, 'messages': ['未找到Keywords段落']},
            'format': {'ok': False, 'messages': ['未找到Keywords段落']},
            'clc_structure': {'ok': False, 'messages': ['未找到Keywords段落']},
            'clc_format': {'ok': False, 'messages': ['未找到Keywords段落']},
            'footnote_structure': {'ok': False, 'messages': ['未找到Keywords段落']},
            'footnote_format': {'ok': False, 'messages': ['未找到Keywords段落']},
            'summary': ['关键词检查失败：未找到Keywords段落']
        }
    
    # 执行关键词检查
    structure_report = check_keywords_structure(keywords_text, tpl)
    paragraphs_report = check_keywords_paragraphs(doc, tpl)
    
    if paragraphs_report['keywords_paragraph']:
        format_report = check_keywords_format(paragraphs_report['keywords_paragraph'], tpl)
    else:
        format_report = {'ok': False, 'messages': ['无法检测关键词格式']}
    
    # 执行CLC和Document code检查
    clc_structure_report = check_clc_document_structure(doc, keywords_paragraph_index, tpl)
    
    if clc_structure_report['clc_paragraph']:
        clc_format_report = check_clc_document_format(clc_structure_report['clc_paragraph'], tpl)
    else:
        clc_format_report = {'ok': False, 'messages': ['无法检测CLC和Document code格式']}
    
    # 执行脚注检查
    footnote_structure_report = check_footnote_structure(doc, tpl)
    footnote_format_report = check_footnote_format(doc, tpl)
    
    # 组装报告
    report = {
        'structure': structure_report,
        'paragraphs': paragraphs_report,
        'format': format_report,
        'clc_structure': clc_structure_report,
        'clc_format': clc_format_report,
        'footnote_structure': footnote_structure_report,
        'footnote_format': footnote_format_report,
        'summary': []
    }
    
    # 生成总结
    all_ok = (structure_report['ok'] and paragraphs_report['ok'] and format_report['ok'] and 
              clc_structure_report['ok'] and clc_format_report['ok'] and
              footnote_structure_report['ok'] and footnote_format_report['ok'])
    summary_tpl = tpl.get('messages', {}).get('summary_overall')
    if summary_tpl:
        try:
            report['summary'].append(summary_tpl.format(ok=all_ok))
        except Exception:
            report['summary'].append(str(summary_tpl))
    
    return report

def title_to_sentence_case(title):
    """
    将Title Case转换为Sentence Case（仅首字母大写）
    保留专有名词和化学式的大小写
    """
    if not title:
        return ""
    
    # 分割成单词
    words = title.split()
    if not words:
        return title
    
    result = []
    for i, word in enumerate(words):
        # 提取核心词和前后标点
        match = re.match(r'^([^\w]*)(\S+)([^\w]*)$', word)
        if not match:
            result.append(word)
            continue
        
        prefix, core_word, suffix = match.groups()
        
        # 规则1：如果单词包含多个大写字母（如FePc, SQL），保留原样
        if sum(1 for c in core_word if c.isupper()) > 1:
            result.append(word)
            continue
        
        # 规则2：首词首字母大写
        if i == 0:
            result.append(prefix + core_word.capitalize() + suffix)
        else:
            # 其他词全小写
            result.append(prefix + core_word.lower() + suffix)
    
    return ' '.join(result)

def normalize_author_name(author_str):
    """
    标准化作者姓名格式，用于比对
    将 'SURNAME I N' 或 'SURNAME I' 格式标准化为可比较的格式
    """
    # 移除多余空格和标点
    author_str = re.sub(r'[\s,]+', ' ', author_str.strip())
    # 转换为大写以便比较
    return author_str.upper()

def extract_authors_from_citation(citation_text):
    """
    从Citation文本中提取作者列表
    返回: [author1, author2, author3] 或 None
    """
    # 匹配格式：AUTHOR1, AUTHOR2, AUTHOR3, et al. Title
    match = re.search(r'^(.+?),\s*et al\.', citation_text, re.IGNORECASE)
    if not match:
        # 尝试不带et al.的格式
        match = re.search(r'^(.+?)\.[^.]*\[J\]', citation_text, re.IGNORECASE)
    
    if not match:
        return None
    
    authors_part = match.group(1).strip()
    # 分割作者（通过逗号）
    authors = [a.strip() for a in authors_part.split(',')]
    return authors

def extract_title_from_citation(citation_text):
    """
    从Citation文本中提取标题
    返回: title string 或 None
    """
    # 匹配格式：...et al. Title [J]. Journal...
    match = re.search(r'et al\.\s*(.+?)\s*\[J\]', citation_text, re.IGNORECASE)
    if not match:
        # 尝试不带et al.的格式：...AUTHOR. Title [J].
        match = re.search(r'\.[^.]*?\s+(.+?)\s*\[J\]', citation_text, re.IGNORECASE)
    
    if match:
        return match.group(1).strip()
    return None

def extract_author_from_correspondence(correspondence_text):
    """
    从Correspondence文本中提取作者名字
    格式: * Correspondence should be addressed to NAME, email: xxx
    返回: author_name 或 None
    """
    # 匹配格式：* Correspondence should be addressed to NAME, email:
    match = re.search(r'Correspondence\s+should\s+be\s+addressed\s+to\s+([^,]+)', correspondence_text, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    
    # 尝试简化格式：* Correspondence: NAME, email:
    match = re.search(r'Correspondence\s*:\s*([^,]+)', correspondence_text, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    
    return None

def match_author_name(correspondence_name, paper_authors):
    """
    匹配Correspondence中的作者名字与论文作者
    返回: (matched_author, is_corresponding) 或 (None, False)
    """
    if not correspondence_name or not paper_authors:
        return None, False
    
    # 标准化correspondence中的名字（转为大写，移除多余空格）
    corr_name_normalized = re.sub(r'\s+', ' ', correspondence_name.upper().strip())
    
    for author in paper_authors:
        # 获取作者的姓和名
        surname = author.get('surname', '').upper()
        given_en = author.get('given_en', '').upper()
        is_corresponding = author.get('corresponding', False)
        
        # 构建可能的匹配模式
        # 1. 完整名字: SURNAME GIVENNAME
        full_name = f"{surname} {given_en}".strip()
        # 2. 姓 + 名首字母: SURNAME G
        given_initial = given_en[0] if given_en else ''
        name_with_initial = f"{surname} {given_initial}".strip()
        
        # 检查是否匹配
        if (corr_name_normalized == full_name.upper() or 
            corr_name_normalized == name_with_initial.upper() or
            surname in corr_name_normalized):
            return author, is_corresponding
    
    return None, False

def extract_paper_title_authors(doc, tpl):
    """
    从论文文档中提取标题和作者信息
    返回: {'title': str, 'authors': [author_dict]} 或 None
    """
    if not TITLE_DETECT_AVAILABLE:
        return None
    
    try:
        # 使用Title_detect的提取功能
        # 首先需要获取非空段落
        nonempty = get_nonempty_paragraphs(doc)
        if not nonempty:
            return None
        
        # 猜测标题位置（跳过章节编号）
        section_pattern = re.compile(r'^\s*\d+(\s+\d+)*\s+[A-Z]')
        title_idx = 0
        for idx, text in enumerate([p.text.strip() for p in doc.paragraphs if p.text]):
            if not section_pattern.match(text):
                title = text
                title_idx = idx
                break
        else:
            title = nonempty[0] if nonempty else ""
        
        # 提取作者信息
        # 查找标题在非空段落中的索引
        try:
            ne_idx = nonempty.index(title)
        except ValueError:
            ne_idx = 0
        
        # 提取作者文本
        authors_text = ""
        i = ne_idx + 1
        stop_markers = re.compile(r'^\s*(Abstract|Keywords|CLC|Introduction|摘要|关键词)[\s:：]', re.IGNORECASE)
        
        while i < len(nonempty):
            line = nonempty[i]
            # 遇到停止标记或单位信息时停止
            if stop_markers.match(line):
                break
            if re.match(r'^\s*\d+[\.\、\s:-]', line):
                break
            if re.search(r'\b(College|University|Institute|School|Department)\b', line, flags=re.I):
                break
            authors_text += " " + line
            i += 1
        
        # 解析作者（使用模板中的author_regex，或使用默认正则）
        author_regex = tpl.get('author_regex', r'^([A-Z]+)\s+([A-Z\s]+)$')
        authors = parse_authors_by_regex(authors_text.strip(), author_regex)
        
        return {
            'title': title,
            'authors': authors
        }
    except Exception as e:
        print(f"提取论文标题和作者时出错: {str(e)}")
        return None

def check_footnote_structure(doc, tpl):
    """
    检查Word文档中的真正脚注结构
    返回 {'ok': bool, 'messages': [], 'footnote_paragraphs': []}
    """
    report = {'ok': True, 'messages': [], 'footnote_paragraphs': []}
    
    try:
        # 访问脚注XML
        rels = doc.part.rels
        footnote_rels = [rel for rel in rels if 'footnotes.xml' in rels[rel].target_ref]
        
        if not footnote_rels:
            report['ok'] = False
            error_msg = tpl.get('messages', {}).get('footnote_missing')
            if error_msg:
                report['messages'].append(error_msg)
            return report
        
        footnote_part = rels[footnote_rels[0]].target_part
        xml_content = footnote_part.blob.decode('utf-8')
        
        # 解析脚注XML
        import xml.etree.ElementTree as ET
        root = ET.fromstring(xml_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        footnotes = root.findall('.//w:footnote', ns)
        print(f"找到 {len(footnotes)} 个脚注")
        
        # 获取结构规则
        structure_rules = tpl.get('structure_rules', {})
        received_pattern = structure_rules.get('footnote_received_pattern', r'Received date\\s*:')
        foundation_pattern = structure_rules.get('footnote_foundation_pattern', r'Foundation item\\s*:')
        correspondence_pattern = structure_rules.get('footnote_correspondence_pattern', r'\\*\\s*Correspondence should be addressed to')
        citation_pattern = structure_rules.get('footnote_citation_pattern', r'Citation\\s*:')
        
        found_items = {
            'received': False,
            'foundation': False, 
            'correspondence': False,
            'citation': False
        }
        
        footnote_content = ""
        
        # 检查每个脚注的内容
        for footnote in footnotes:
            footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if footnote_id in ['-1', '0']:  # 跳过分隔符脚注
                continue
                
            # 获取脚注文本
            text_elements = footnote.findall('.//w:t', ns)
            footnote_text = ''.join([t.text for t in text_elements if t.text])
            footnote_content += footnote_text + " "
            
            print(f"脚注 {footnote_id} 内容: {footnote_text}")
            
            # 检查各个项目
            # 1. 检查Received date（正确拼写）
            if re.search(received_pattern, footnote_text, re.IGNORECASE):
                found_items['received'] = True
                ok_msg = tpl.get('messages', {}).get('footnote_received_ok')
                if ok_msg:
                    report['messages'].append(ok_msg)
            # 检查常见错误：Receive date
            elif re.search(r'Receive\s+date\s*:', footnote_text, re.IGNORECASE):
                found_items['received'] = False
                report['ok'] = False
                report['messages'].append("❌ 字段拼写错误：正确写法为'Received date'，当前为'Receive date'")
            
            # 2. 检查Foundation item（正确拼写）
            if re.search(foundation_pattern, footnote_text, re.IGNORECASE):
                found_items['foundation'] = True
                ok_msg = tpl.get('messages', {}).get('footnote_foundation_ok')
                if ok_msg:
                    report['messages'].append(ok_msg)
            # 检查常见错误：Foundation items
            elif re.search(r'Foundation\s+items\s*:', footnote_text, re.IGNORECASE):
                found_items['foundation'] = False
                report['ok'] = False
                report['messages'].append("❌ 字段拼写错误：正确写法为'Foundation item'，当前为'Foundation items'")
            
            # 检查Correspondence（正确拼写）
            if re.search(correspondence_pattern, footnote_text, re.IGNORECASE):
                found_items['correspondence'] = True
                
                # === 新增：检查Correspondence中的作者是否正确 ===
                if TITLE_DETECT_AVAILABLE:
                    try:
                        # 提取Correspondence中的作者名字
                        corr_author_name = extract_author_from_correspondence(footnote_text)
                        
                        if corr_author_name:
                            # 提取论文作者信息
                            doc_info = extract_paper_title_authors(doc, tpl)
                            
                            if doc_info and doc_info['authors']:
                                # 匹配Correspondence作者与论文作者
                                matched_author, is_corresponding = match_author_name(corr_author_name, doc_info['authors'])
                                
                                if matched_author:
                                    # 找到匹配的作者
                                    author_full_name = f"{matched_author.get('surname', '')} {matched_author.get('given_en', '')}".strip()
                                    
                                    if is_corresponding:
                                        # 作者被标记为通讯作者
                                        report['messages'].append(f"✓ Correspondence作者'{corr_author_name}'与通讯作者'{author_full_name}'匹配")
                                    else:
                                        # 作者存在但未标记为通讯作者
                                        report['ok'] = False
                                        report['messages'].append(f"⚠️ Correspondence作者'{corr_author_name}'在作者列表中但未标记为通讯作者（缺少*或†标记）")
                                else:
                                    # 未找到匹配的作者
                                    report['ok'] = False
                                    report['messages'].append(f"❌ Correspondence作者'{corr_author_name}'不在论文作者列表中")
                                    # 列出论文中的所有作者供参考
                                    author_names = [f"{a.get('surname', '')} {a.get('given_en', '')}".strip() for a in doc_info['authors']]
                                    report['messages'].append(f"   论文作者列表: {', '.join(author_names)}")
                        else:
                            report['messages'].append("⚠️ 无法从Correspondence中提取作者名字")
                    except Exception as e:
                        # 如果检查失败，记录但不影响主检测
                        print(f"Correspondence作者验证时出错: {str(e)}")
                else:
                    # Title_detect不可用，跳过作者验证
                    pass
                
                # 基本格式检查通过的消息
                ok_msg = tpl.get('messages', {}).get('footnote_correspondence_ok')
                if ok_msg and not any('Correspondence作者' in m for m in report['messages'][-3:]):
                    report['messages'].append(ok_msg)
                    
            # 检查是否错误拼写为Corresponding
            elif re.search(r'\*\s*Corresponding\s+', footnote_text, re.IGNORECASE):
                found_items['correspondence'] = False
                report['ok'] = False
                report['messages'].append("❌ 字段拼写错误：正确写法为'Correspondence'，当前为'Corresponding'")
            
            if re.search(citation_pattern, footnote_text, re.IGNORECASE):
                found_items['citation'] = True
                
                # 详细检查Citation格式
                citation_issues = []
                
                # 检查Journal名称
                if 'Journal of Donghua University (English Edition)' not in footnote_text:
                    citation_issues.append("缺少Journal of Donghua University (English Edition)")
                
                # 提取Citation内容进行详细分析
                citation_match = re.search(r'Citation\s*:\s*(.+)', footnote_text, re.IGNORECASE)
                if citation_match:
                    citation_content = citation_match.group(1).strip()
                    print(f"Citation内容: {citation_content}")
                    
                    # 检查作者格式和数量
                    # 支持两种匹配模式：
                    # 1. 带et al.：AUTHOR1, AUTHOR2, AUTHOR3, et al. Title [J]
                    # 2. 不带et al.：AUTHOR1, AUTHOR2, AUTHOR3. Title [J]
                    author_title_match = re.search(r'^(.+?),\s*et al\.\s*(.+?)\s*\[J\]', citation_content, re.IGNORECASE)
                    if not author_title_match:
                        # 尝试不带et al.的格式
                        author_title_match = re.search(r'^(.+?)\.\s*(.+?)\s*\[J\]', citation_content, re.IGNORECASE)
                    
                    if author_title_match:
                        authors_part = author_title_match.group(1).strip()
                        title_part = author_title_match.group(2).strip()
                        
                        print(f"作者部分: {authors_part}")
                        print(f"标题部分: {title_part}")
                        
                        # 计算作者数量（通过逗号分割）
                        authors_list = [author.strip() for author in authors_part.split(',')]
                        author_count = len(authors_list)
                        
                        print(f"检测到 {author_count} 个作者: {authors_list}")
                        
                        if author_count != 3:
                            citation_issues.append(f"Citation中作者数量应为3个，实际为{author_count}个")
                            report['ok'] = False
                        
                        # 检查作者姓名格式（支持灵活缩写格式）
                        for i, author in enumerate(authors_list):
                            # 检查是否为缩写格式：SURNAME I 或 SURNAME I N
                            author_clean = author.strip()
                            # 支持两种格式：SURNAME I（两字母）或 SURNAME I N（三字母）
                            if not re.match(r'^[A-Z]+\s+[A-Z](?:\s+[A-Z])?$', author_clean):
                                citation_issues.append(f"作者 {i+1} '{author}' 格式不正确，应为 'SURNAME I' 或 'SURNAME I N' 格式")
                                report['ok'] = False
                    else:
                        citation_issues.append("Citation格式不正确，无法解析作者和标题部分")
                        report['ok'] = False
                
                # 报告Citation检测结果
                if citation_issues:
                    for issue in citation_issues:
                        report['messages'].append(issue)
                else:
                    ok_msg = tpl.get('messages', {}).get('footnote_citation_ok')
                    if ok_msg:
                        report['messages'].append(ok_msg)
                
                # === 新增：比对Citation中的作者和标题与论文实际内容 ===
                if TITLE_DETECT_AVAILABLE:
                    try:
                        # 提取论文的标题和作者
                        doc_info = extract_paper_title_authors(doc, tpl)
                        
                        if doc_info:
                            # 1. 比对作者
                            citation_authors = extract_authors_from_citation(citation_content)
                            if citation_authors and doc_info['authors']:
                                # 比对前三位作者
                                paper_authors = doc_info['authors'][:3]
                                
                                if len(citation_authors) != len(paper_authors):
                                    report['ok'] = False
                                    report['messages'].append(f"❌ Citation中作者数量({len(citation_authors)}个)与论文作者数量({len(paper_authors)}个)不匹配")
                                else:
                                    # 逐个比对作者姓名
                                    for i, (cit_author, paper_author) in enumerate(zip(citation_authors, paper_authors)):
                                        # 标准化姓名进行比对
                                        cit_normalized = normalize_author_name(cit_author)
                                        # 从论文作者信息中提取姓和名首字母
                                        surname = paper_author.get('surname', '').upper()
                                        given_en = paper_author.get('given_en', '').upper()
                                        # 提取名的首字母
                                        given_initials = ' '.join([word[0] for word in given_en.split() if word])
                                        paper_normalized = f"{surname} {given_initials}".strip()
                                        
                                        if surname not in cit_normalized:
                                            report['ok'] = False
                                            report['messages'].append(f"❌ Citation中第{i+1}位作者'{cit_author}'与论文作者'{surname} {given_en}'不匹配")
                            
                            # 2. 比对标题
                            citation_title = extract_title_from_citation(citation_content)
                            if citation_title and doc_info['title']:
                                paper_title = doc_info['title']
                                # 将论文标题（Title Case）转换为Sentence Case
                                expected_citation_title = title_to_sentence_case(paper_title)
                                
                                # 比较标题（忽略多余空格和标点差异）
                                cit_title_normalized = re.sub(r'[\s\.,;:!?]+', ' ', citation_title.lower().strip())
                                expected_normalized = re.sub(r'[\s\.,;:!?]+', ' ', expected_citation_title.lower().strip())
                                
                                if cit_title_normalized != expected_normalized:
                                    report['ok'] = False
                                    report['messages'].append(f"❌ Citation中的标题与论文标题不一致")
                                    report['messages'].append(f"   论文标题: {paper_title}")
                                    report['messages'].append(f"   期望格式(sentence case): {expected_citation_title}")
                                    report['messages'].append(f"   实际Citation: {citation_title}")
                                else:
                                    report['messages'].append("✓ Citation标题与论文标题一致（sentence case格式）")
                    except Exception as e:
                        # 如果比对失败，记录但不影响主检测
                        print(f"作者/标题比对时出错: {str(e)}")
        
        # 检查是否找到所有必需项目
        missing_items = [key for key, value in found_items.items() if not value]
        if missing_items:
            report['ok'] = False
            for item in missing_items:
                report['messages'].append(f"脚注中未找到{item}项目")
        
        if report['ok'] and all(found_items.values()):
            ok_msg = tpl.get('messages', {}).get('footnote_structure_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
    
    except Exception as e:
        report['ok'] = False
        report['messages'].append(f"脚注检测出错: {str(e)}")
    
    return report

def check_footnote_format(doc, tpl):
    """
    检查Word脚注的格式（字体、大小、行距等）
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    try:
        # 访问脚注XML
        rels = doc.part.rels
        footnote_rels = [rel for rel in rels if 'footnotes.xml' in rels[rel].target_ref]
        
        if not footnote_rels:
            report['ok'] = False
            report['messages'].append("无法访问脚注进行格式检查")
            return report
        
        footnote_part = rels[footnote_rels[0]].target_part
        xml_content = footnote_part.blob.decode('utf-8')
        
        # 解析脚注XML
        import xml.etree.ElementTree as ET
        root = ET.fromstring(xml_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        footnotes = root.findall('.//w:footnote', ns)
        
        # 获取格式规则
        format_rules = tpl.get('format_rules', {}).get('footnote', {})
        issues = []
        
        # 检查实际脚注内容（跳过分隔符）
        for footnote in footnotes:
            footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if footnote_id in ['-1', '0']:  # 跳过分隔符脚注
                continue
            
            # 检查字体大小
            sz_elements = footnote.findall('.//w:sz', ns)
            if sz_elements:
                for sz in sz_elements:
                    sz_val = sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if sz_val:
                        actual_size_pt = float(sz_val) / 2.0
                        expected_size_pt = float(format_rules.get('font_size_pt', 9))
                        actual_size_name = get_font_size(actual_size_pt, tpl)
                        expected_size_name = get_font_size(expected_size_pt, tpl)
                        
                        print(f"脚注字体大小: {actual_size_name}（{actual_size_pt}pt）(期望: {expected_size_name}（{expected_size_pt}pt）)")
                        if abs(actual_size_pt - expected_size_pt) > 0.5:
                            issues.append(f"脚注字体大小应为{expected_size_name}（{expected_size_pt}pt），实际为{actual_size_name}（{actual_size_pt}pt）")
                        break
            
            # 检查字体名称
            font_elements = footnote.findall('.//w:rFonts', ns)
            if font_elements:
                for font_elem in font_elements:
                    ascii_font = font_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                    if ascii_font:
                        expected_font_name = str(format_rules.get('font_name', 'Times New Roman'))
                        print(f"脚注字体名称: {ascii_font} (期望: {expected_font_name})")
                        if expected_font_name.lower() not in ascii_font.lower():
                            issues.append(f"脚注字体应为{expected_font_name}，实际为{ascii_font}")
                        break
            
            # 检查是否有斜体（Journal名称应该斜体）
            text_elements = footnote.findall('.//w:t', ns)
            footnote_text = ''.join([t.text for t in text_elements if t.text])
            
            if 'Journal of Donghua University (English Edition)' in footnote_text:
                # 检查Journal部分是否为斜体
                journal_italic = format_rules.get('journal_italic', True)
                if journal_italic:
                    # 查找包含Journal文本的run，检查是否有斜体标记
                    journal_found_italic = False
                    for run in footnote.findall('.//w:r', ns):
                        run_text_elements = run.findall('.//w:t', ns)
                        run_text = ''.join([t.text for t in run_text_elements if t.text])
                        if 'Journal of Donghua University' in run_text:
                            italic_elements = run.findall('.//w:i', ns)
                            if italic_elements:
                                journal_found_italic = True
                                break
                    
                    if not journal_found_italic:
                        issues.append("Journal of Donghua University (English Edition)应为斜体")
        
        if issues:
            report['ok'] = False
            header = tpl.get('messages', {}).get('footnote_format_error')
            if header:
                report['messages'].append(header + "：")
            report['messages'].extend([f"  - {i}" for i in issues])
        else:
            ok_msg = tpl.get('messages', {}).get('footnote_format_ok')
            if ok_msg:
                report['messages'].append(ok_msg)
    
    except Exception as e:
        report['ok'] = False
        report['messages'].append(f"脚注格式检测出错: {str(e)}")
    
    return report

# ---------- 报告输出 ----------
def print_keywords_report(report):
    """打印关键词、CLC和脚注检查报告"""
    print("=== Paper Format Check Report ===")
    
    sections = [
        ('structure', 'KEYWORDS STRUCTURE'),
        ('paragraphs', 'KEYWORDS PARAGRAPHS'), 
        ('format', 'KEYWORDS FORMAT'),
        ('clc_structure', 'CLC STRUCTURE'),
        ('clc_format', 'CLC FORMAT'),
        ('footnote_structure', 'FOOTNOTE STRUCTURE'),
        ('footnote_format', 'FOOTNOTE FORMAT')
    ]
    
    for sec_key, sec_name in sections:
        info = report.get(sec_key, {'ok': False, 'messages': ['检测项不存在']})
        print(f"--- {sec_name} ---")
        print(" OK:", info['ok'])
        for m in info['messages']:
            print("  -", m)
    
    print("--- SUMMARY ---")
    for s in report['summary']:
        print(" ", s)

def print_help():
    print("Usage:")
    print("  python Keywords_detect.py check <paper.docx> <template.json_or_name>")

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print_help()
        sys.exit(0)
    
    cmd = sys.argv[1]
    if cmd == 'check':
        paper_path = sys.argv[2]
        tpl_id = sys.argv[3]
        if not os.path.isfile(paper_path):
            print(f"论文文件不存在: {paper_path}")
            sys.exit(1)
        try:
            print("=== 开始关键词格式检查 ===")
            report = check_keywords_with_template(paper_path, tpl_id)
        except Exception as e:
            print("检查时出错:", e)
            sys.exit(1)
        print_keywords_report(report)
    else:
        print_help()
        sys.exit(0)
'''
python paper_detect\Keywords_detect.py check template\test.docx templates\Keywords.json
'''