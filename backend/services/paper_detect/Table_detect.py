#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import Table

"""
=== 论文格式检测系统 - 表格检测器 ===

【表格检测 (Table Detection)】

1. 【表格标题格式检测】
   - 表格标题格式：Table X 表名（X为数字编号）
   - 标题加粗，居中对齐
   - 表名首字母大写
   - 字体大小与正文一致（小四号12pt）
   - 表格编号应连续递增

2. 【三线表格式检测】
   - 表格采用三线表格式
   - 仅有顶线、底线和表头底线
   - 无竖线和其他横线

3. 【表格内容对齐检测】
   - 单行内容居中对齐
   - 多行内容左对齐
   - 字体格式与正文一致

4. 【技术特性】
   - 支持表格标题的智能识别
   - 验证表格编号的连续性
   - 检测三线表格式的正确性
   - 分析表格内容的对齐方式
"""

# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    """解析模板路径，支持文件路径和模板名称"""
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    """加载JSON模板文件"""
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    return tpl

# ---------- 字体检测函数 ----------
def detect_font_for_run(run, paragraph=None):
    """
    更健壮的字体检测函数，同时检查直接格式和段落样式。
    返回 (font_size_pt, font_name, is_bold, is_italic, candidates_dict)
    """
    font_source = run.font if run else (paragraph.style.font if paragraph else None)
    if not font_source:
        return 12.0, "Unknown", False, False, {}

    font_name = font_source.name if font_source.name else "Times New Roman"

    font_size = None
    try:
        if run and run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        if font_size is None and run and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
        
        if font_size is None and paragraph and paragraph.style and hasattr(paragraph.style.font, 'size') and hasattr(paragraph.style.font.size, 'pt'):
            font_size = float(paragraph.style.font.size.pt)
            
        if font_size is None and paragraph and hasattr(paragraph.style, 'element'):
            sz_nodes = paragraph.style.element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass

    font_size = font_size if font_size is not None else 12.0

    is_bold = font_source.bold if font_source.bold is not None else False
    is_italic = font_source.italic if font_source.italic is not None else False
    if run and run.font:
        is_bold = run.font.bold if run.font.bold is not None else is_bold
        is_italic = run.font.italic if run.font.italic is not None else is_italic
    
    # 确保返回明确的布尔值，而不是None
    is_bold = bool(is_bold) if is_bold is not None else False
    is_italic = bool(is_italic) if is_italic is not None else False

    return font_size, font_name, is_bold, is_italic, {}

def get_font_size(pt_size, tpl=None):
    """获取字体大小的中文名称"""
    # 从模板读取字体大小映射，提供默认值以保持向后兼容
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        # 将字符串键转换为浮点数
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        # 默认的字体大小映射（向后兼容）
        size_map = {
            9: "小五",
            10.5: "五号",
            12: "小四",
            14: "四号",
            16: "三号",
            18: "小二",
            22: "二号",
            24: "小一",
            26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]

# ---------- 段落对齐检测函数 ----------
def get_style_alignment(doc, style_id):
    """从文档的styles.xml中查找并返回指定样式的对齐方式。"""
    try:
        styles = doc.styles.element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        style_path = f".//w:style[@w:styleId='{style_id}']//w:jc"
        jc_element = styles.find(style_path, ns)
        if jc_element is not None:
            val = jc_element.get(ns['w'] + 'val')
            mapping = {'left': 0, 'center': 1, 'right': 2, 'both': 3, 'justify': 3}
            return mapping.get(val)
    except Exception:
        return None
    return None

def detect_paragraph_alignment(paragraph):
    """
    改进的段落对齐检测，按照明确的优先级处理直接格式和样式格式
    
    检测逻辑：
    1. 如果直接格式不为None，直接使用直接格式结果
    2. 如果直接格式为None，检查样式格式
    3. 检查XML样式定义（处理复杂的样式继承）
    4. 默认为左对齐
    """
    
    # 优先级1: 检查直接格式
    direct_alignment = paragraph.paragraph_format.alignment
    if direct_alignment is not None:
        # 如果有直接格式设置，直接使用
        return int(direct_alignment)
    
    # 优先级2: 如果直接格式为None，检查样式格式
    if paragraph.style:
        try:
            style_alignment = paragraph.style.paragraph_format.alignment
            if style_alignment is not None:
                # 如果样式有对齐设置，使用样式对齐
                return int(style_alignment)
        except Exception:
            pass
    
    # 优先级3: 检查XML样式定义（处理复杂的样式继承）
    if paragraph.style and paragraph.style.style_id:
        try:
            xml_alignment = get_style_alignment(paragraph.part.document, paragraph.style.style_id)
            if xml_alignment is not None:
                return xml_alignment
        except Exception:
            pass
            
    # 默认值: 左对齐
    return 0  # WD_PARAGRAPH_ALIGNMENT.LEFT

# ---------- 表格检测核心函数 ----------

def identify_table_captions(doc, tpl):
    """
    识别文档中的表格标题
    返回：表格标题段落列表，每个元素包含段落对象、编号、标题文本
    """
    caption_pattern = tpl.get('table_detection_rules', {}).get('caption_pattern', r'^\s*Table\s+(\d+)\s+(.+)$')
    captions = []
    
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        match = re.match(caption_pattern, text, re.IGNORECASE)
        if match:
            table_num = int(match.group(1))
            table_title = match.group(2).strip()
            captions.append({
                'paragraph': paragraph,
                'paragraph_index': idx,
                'number': table_num,
                'title': table_title,
                'full_text': text
            })
    
    return captions

def find_table_after_caption(doc, caption_info):
    """
    在标题后查找对应的表格
    返回：Table对象或None
    """
    caption_idx = caption_info['paragraph_index']
    
    # 在标题后的几个元素中查找表格
    for i in range(caption_idx + 1, min(caption_idx + 5, len(doc.element.body))):
        element = doc.element.body[i]
        if element.tag.endswith('tbl'):  # 表格元素
            # 通过索引找到对应的Table对象
            table_idx = 0
            for j, elem in enumerate(doc.element.body):
                if elem.tag.endswith('tbl'):
                    if j == i:
                        return doc.tables[table_idx] if table_idx < len(doc.tables) else None
                    table_idx += 1
    
    return None

def check_table_style(table, tpl):
    """
    检查表格是否为三线表格式
    返回：(is_three_line, issues)
    """
    issues = []
    border_width_issues = []
    
    if not table:
        issues.append("未找到表格对象")
        return False, issues
    
    try:
        # 检查表格的边框设置
        # 三线表特征：只有顶部、底部和表头底部有线
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        
        if tblPr is None:
            issues.append("表格缺少格式属性")
            return False, issues
        
        # 检查表格边框
        tblBorders = tblPr.find(qn('w:tblBorders'))
        
        # 三线表应该：
        # 1. 有顶部边框（top）
        # 2. 有底部边框（bottom）
        # 3. 无左右边框或边框为none
        # 4. 表头行底部有边框
        
        has_issues = False
        
        if tblBorders is not None:
            # 检查是否有左右边框（三线表不应有）
            left_border = tblBorders.find(qn('w:left'))
            right_border = tblBorders.find(qn('w:right'))
            inside_v = tblBorders.find(qn('w:insideV'))
            
            if left_border is not None and left_border.get(qn('w:val')) not in [None, 'none', 'nil']:
                issues.append("表格不应有左边框（三线表格式）")
                has_issues = True
                
            if right_border is not None and right_border.get(qn('w:val')) not in [None, 'none', 'nil']:
                issues.append("表格不应有右边框（三线表格式）")
                has_issues = True
                
            if inside_v is not None and inside_v.get(qn('w:val')) not in [None, 'none', 'nil']:
                issues.append("表格不应有内部竖线（三线表格式）")
                has_issues = True
        
        # 检查是否只有三条横线（简化检查：检查表格的行边框设置）
        # 这里我们检查是否大部分内部横线都是none
        inside_h = None
        if tblBorders is not None:
            inside_h = tblBorders.find(qn('w:insideH'))
            
        # 如果内部横线不是none，说明可能不是标准三线表
        if inside_h is not None and inside_h.get(qn('w:val')) not in [None, 'none', 'nil']:
            # 需要进一步检查：可能第一行（表头）下方有线是正常的
            # 这里简化处理：如果有内部横线，我们检查是否只在表头下方
            pass  # 暂时允许，因为表头下方应该有线
        
        # 检查边框粗细（如果模板中启用了此检查）
        if tpl.get('check_rules', {}).get('border_width_check', False):
            border_config = tpl.get('table_detection_rules', {}).get('table_style', {}).get('border_width', {})
            tolerance = tpl.get('check_rules', {}).get('border_width_tolerance', 0.1)
            
            expected_top = border_config.get('top_line', 1.5)
            expected_bottom = border_config.get('bottom_line', 1.5)
            expected_header = border_config.get('header_line', 0.75)
            
            # 三线表的边框通常设置在单元格级别，而不是表格级别
            # 1. 检查顶线：第一行单元格的顶边框
            if len(table.rows) > 0:
                first_row = table.rows[0]
                if first_row.cells:
                    first_cell = first_row.cells[0]
                    tc = first_cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is not None:
                            # 检查顶线
                            top_border = tcBorders.find(qn('w:top'))
                            if top_border is not None:
                                sz = top_border.get(qn('w:sz'))
                                if sz:
                                    actual_width = float(sz) / 8.0
                                    if abs(actual_width - expected_top) > tolerance:
                                        msg_template = tpl.get('messages', {}).get('top_border_width_error', '顶线宽度应为{expected}磅，实际为{actual}磅')
                                        border_width_issues.append(msg_template.format(expected=expected_top, actual=round(actual_width, 2)))
                                        has_issues = True
                            
                            # 检查表头底线（第一行底边框）
                            bottom_border = tcBorders.find(qn('w:bottom'))
                            if bottom_border is not None:
                                sz = bottom_border.get(qn('w:sz'))
                                if sz:
                                    actual_width = float(sz) / 8.0
                                    if abs(actual_width - expected_header) > tolerance:
                                        msg_template = tpl.get('messages', {}).get('header_border_width_error', '表头底线宽度应为{expected}磅，实际为{actual}磅')
                                        border_width_issues.append(msg_template.format(expected=expected_header, actual=round(actual_width, 2)))
                                        has_issues = True
            
            # 2. 检查底线：最后一行单元格的底边框
            if len(table.rows) > 0:
                last_row = table.rows[-1]
                if last_row.cells:
                    first_cell = last_row.cells[0]
                    tc = first_cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is not None:
                            bottom_border = tcBorders.find(qn('w:bottom'))
                            if bottom_border is not None:
                                sz = bottom_border.get(qn('w:sz'))
                                if sz:
                                    actual_width = float(sz) / 8.0
                                    if abs(actual_width - expected_bottom) > tolerance:
                                        msg_template = tpl.get('messages', {}).get('bottom_border_width_error', '底线宽度应为{expected}磅，实际为{actual}磅')
                                        border_width_issues.append(msg_template.format(expected=expected_bottom, actual=round(actual_width, 2)))
                                        has_issues = True
        
        issues.extend(border_width_issues)
        
        if has_issues:
            return False, issues
        else:
            return True, []
            
    except Exception as e:
        issues.append(f"表格格式检测异常: {str(e)}")
        return False, issues

def check_table_content_alignment(table, tpl):
    """
    检查表格内容的对齐方式
    规则：
    - 表头行（第1行）：所有单元格居中对齐
    - 内容行（其他行）：较长内容左对齐，较短内容居中对齐
    返回：(is_correct, issues)
    """
    issues = []
    
    if not table:
        issues.append("未找到表格对象")
        return False, issues
    
    # 对齐方式名称映射
    alignment_names = {
        0: '左对齐',
        1: '居中对齐',
        2: '右对齐',
        3: '两端对齐'
    }
    
    # 获取长度阈值（超过此长度视为"较长的话"，应该左对齐）
    # 可以从模板配置中读取，默认20个字符
    length_threshold = tpl.get('check_rules', {}).get('alignment_length_threshold', 20)
    
    try:
        # 用于跟踪已检查的单元格（避免合并单元格重复检查）
        checked_cells = set()
        
        for row_idx, row in enumerate(table.rows):
            # 判断是否是表头行（第1行，索引为0）
            is_header_row = (row_idx == 0)
            
            # 计算逻辑列号（排除合并单元格的重复）
            logical_col = 0
            
            for cell_idx, cell in enumerate(row.cells):
                # 获取单元格的唯一标识（用于处理合并单元格）
                cell_id = id(cell._element)
                is_merged = cell_id in checked_cells
                
                if is_merged:
                    continue
                    
                checked_cells.add(cell_id)
                logical_col += 1
                
                # 获取单元格中的段落
                paragraphs = cell.paragraphs
                
                if not paragraphs:
                    continue
                
                # 获取非空段落
                non_empty_paras = [p for p in paragraphs if p.text.strip()]
                if not non_empty_paras:
                    continue
                
                # 获取单元格文本用于显示
                cell_text = cell.text.strip()
                text_preview = cell_text[:20] + '...' if len(cell_text) > 20 else cell_text
                text_length = len(cell_text)
                
                # 只检查第一个非空段落的对齐方式（避免重复报告）
                para = non_empty_paras[0]
                actual_alignment = detect_paragraph_alignment(para)
                actual_alignment_name = alignment_names.get(actual_alignment, f'未知({actual_alignment})')
                
                # 根据位置判断期望的对齐方式
                if is_header_row:
                    # 表头行：所有单元格都应该居中对齐
                    expected_alignment = 1  # 居中
                    if actual_alignment != expected_alignment:
                        issues.append(
                            f"单元格[第{row_idx+1}行(表头),第{logical_col}列]应居中对齐"
                            f"（当前：{actual_alignment_name}，内容：{text_preview}）"
                        )
                else:
                    # 内容行：根据文本长度判断
                    if text_length > length_threshold:
                        # 较长内容：应该左对齐
                        expected_alignment = 0  # 左对齐
                        if actual_alignment != expected_alignment:
                            issues.append(
                                f"单元格[第{row_idx+1}行,第{logical_col}列]较长内容应左对齐"
                                f"（当前：{actual_alignment_name}，长度：{text_length}字符，内容：{text_preview}）"
                            )
                    else:
                        # 较短内容：应该居中对齐
                        expected_alignment = 1  # 居中对齐
                        if actual_alignment != expected_alignment:
                            issues.append(
                                f"单元格[第{row_idx+1}行,第{logical_col}列]较短内容应居中对齐"
                                f"（当前：{actual_alignment_name}，长度：{text_length}字符，内容：{text_preview}）"
                            )
        
        return len(issues) == 0, issues
        
    except Exception as e:
        issues.append(f"表格对齐方式检测异常: {str(e)}")
        return False, issues

def check_caption_format(caption_info, tpl):
    """
    检查表格标题的格式
    返回：{'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    paragraph = caption_info['paragraph']
    expected_format = tpl.get('format_rules', {}).get('caption', {})
    
    # 检查加粗
    if expected_format.get('bold', True):
        main_run = next((r for r in paragraph.runs if r.text.strip()), None)
        if main_run:
            _, _, is_bold, _, _ = detect_font_for_run(main_run, paragraph)
            if not is_bold:
                report['ok'] = False
                report['messages'].append(tpl.get('messages', {}).get('caption_bold_error', '表格标题应加粗'))
    
    # 检查对齐方式
    if expected_format.get('alignment') == 'center':
        # 使用改进的对齐检测方法（直接格式 -> 样式格式 -> XML样式）
        actual_alignment = detect_paragraph_alignment(paragraph)
        
        # 获取对齐方式的可读名称
        alignment_names = {
            0: '左对齐',
            1: '居中对齐',
            2: '右对齐',
            3: '两端对齐'
        }
        actual_alignment_name = alignment_names.get(actual_alignment, f'未知对齐方式({actual_alignment})')
        
        # 居中对齐的值为1
        if actual_alignment != 1:
            report['ok'] = False
            msg = tpl.get('messages', {}).get('caption_alignment_error', '表格标题应居中对齐')
            report['messages'].append(f"{msg}（当前：{actual_alignment_name}）")
    
    # 检查字体大小
    expected_size = expected_format.get('font_size_pt', 12)
    main_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if main_run:
        actual_size, actual_font, _, _, _ = detect_font_for_run(main_run, paragraph)
        if abs(actual_size - expected_size) > 0.5:
            report['ok'] = False
            expected_size_name = get_font_size(expected_size, tpl)
            actual_size_name = get_font_size(actual_size, tpl)
            msg = tpl.get('messages', {}).get('caption_font_size_error', '表格标题字体大小不正确')
            report['messages'].append(f"{msg}（期望：{expected_size_name}，实际：{actual_size_name}）")
    
    # 检查标题大小写（首字母应大写）
    title = caption_info['title']
    if title and not title[0].isupper():
        report['ok'] = False
        report['messages'].append(tpl.get('messages', {}).get('caption_title_case_error', '表格名称首字母应大写'))
    
    return report

def check_table_numbering(captions, tpl):
    """
    检查表格编号的连续性
    返回：{'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if not captions:
        return report
    
    numbers = sorted([c['number'] for c in captions])
    
    # 检查是否从1开始
    if numbers[0] != 1:
        report['ok'] = False
        report['messages'].append(f"表格编号应从1开始，当前从{numbers[0]}开始")
    
    # 检查是否连续
    for i in range(len(numbers) - 1):
        if numbers[i + 1] - numbers[i] != 1:
            report['ok'] = False
            msg_template = tpl.get('messages', {}).get('caption_numbering_error', '表格编号不连续')
            report['messages'].append(msg_template.format(numbers=numbers) if '{numbers}' in msg_template else f"表格编号不连续：{numbers}")
            break
    
    return report

def check_doc_with_template(doc_path, template_identifier):
    """
    主检查函数：使用模板检查文档中的表格格式
    返回完整的检查报告
    """
    tpl = load_template(template_identifier)
    doc = Document(doc_path)
    
    # 识别所有表格标题
    captions = identify_table_captions(doc, tpl)
    
    report = {
        'captions': [],
        'tables': [],
        'numbering': {},
        'summary': [],
        'details': {
            'total_captions': len(captions),
            'caption_info': captions
        }
    }
    
    # 检查表格编号
    numbering_report = check_table_numbering(captions, tpl)
    report['numbering'] = numbering_report
    
    # 检查每个表格
    for caption_info in captions:
        table_report = {
            'caption': caption_info,
            'caption_format': {},
            'table_style': {},
            'table_alignment': {},
            'table_object': None
        }
        
        # 检查标题格式
        caption_format_report = check_caption_format(caption_info, tpl)
        table_report['caption_format'] = caption_format_report
        
        # 查找对应的表格
        table = find_table_after_caption(doc, caption_info)
        table_report['table_object'] = table is not None
        
        if table:
            # 检查三线表格式
            is_three_line, style_issues = check_table_style(table, tpl)
            table_report['table_style'] = {
                'ok': is_three_line,
                'messages': style_issues if not is_three_line else [tpl.get('messages', {}).get('table_style_ok', '表格为三线表格式')]
            }
            
            # 检查内容对齐
            is_aligned, alignment_issues = check_table_content_alignment(table, tpl)
            table_report['table_alignment'] = {
                'ok': is_aligned,
                'messages': alignment_issues if not is_aligned else [tpl.get('messages', {}).get('table_content_alignment_ok', '表格内容对齐方式正确')]
            }
        else:
            table_report['table_style'] = {
                'ok': False,
                'messages': [tpl.get('messages', {}).get('table_not_found', '未在标题下方找到表格')]
            }
            table_report['table_alignment'] = {
                'ok': False,
                'messages': []
            }
        
        report['tables'].append(table_report)
    
    # 生成总结
    all_ok = True
    if not numbering_report['ok']:
        all_ok = False
    
    for table_report in report['tables']:
        if not table_report['caption_format']['ok']:
            all_ok = False
        if not table_report['table_style']['ok']:
            all_ok = False
        if not table_report['table_alignment']['ok']:
            all_ok = False
    
    summary_template = tpl.get('messages', {}).get('summary_overall', '表格格式检查结果: {ok}')
    report['summary'].append(summary_template.format(ok='通过' if all_ok else '发现问题'))
    
    count_msg = tpl.get('messages', {}).get('table_count', '检测到 {count} 个表格')
    report['summary'].append(count_msg.format(count=len(captions)))
    
    report['overall_ok'] = all_ok
    
    return report

# ---------- 报表输出 ----------
def print_report(report):
    """打印检查报告"""
    print("=" * 60)
    print("表格格式检测报告")
    print("=" * 60)
    
    # 打印总结
    print("\n【检查总结】")
    for summary in report.get('summary', []):
        print(f"  {summary}")
    
    # 打印编号检查结果
    print("\n【表格编号检查】")
    numbering = report.get('numbering', {})
    if numbering.get('ok', True):
        print("  ✓ 表格编号正确且连续")
    else:
        print("  ✗ 表格编号存在问题：")
        for msg in numbering.get('messages', []):
            print(f"    - {msg}")
    
    # 打印每个表格的检查结果
    print("\n【各表格详细检查】")
    for idx, table_report in enumerate(report.get('tables', []), 1):
        caption_info = table_report['caption']
        print(f"\n  表格 {caption_info['number']}: {caption_info['full_text']}")
        
        # 标题格式检查
        caption_format = table_report['caption_format']
        if caption_format.get('ok', True):
            print("    ✓ 标题格式正确")
        else:
            print("    ✗ 标题格式问题：")
            for msg in caption_format.get('messages', []):
                print(f"      - {msg}")
        
        # 表格样式检查
        table_style = table_report['table_style']
        if table_style.get('ok', True):
            print("    ✓ 三线表格式正确")
        else:
            print("    ✗ 表格样式问题：")
            for msg in table_style.get('messages', []):
                print(f"      - {msg}")
        
        # 对齐方式检查
        table_alignment = table_report['table_alignment']
        if table_alignment.get('ok', True):
            print("    ✓ 内容对齐方式正确")
        else:
            if table_alignment.get('messages'):
                print("    ✗ 对齐方式问题：")
                for msg in table_alignment.get('messages', []):
                    print(f"      - {msg}")
    
    print("\n" + "=" * 60)

def print_help():
    """显示帮助信息"""
    print("用法:")
    print("  python Table_detect.py check <paper.docx> <template.json_or_name>")
    print("")
    print("示例:")
    print("  python Table_detect.py check template/test.docx Table")
    print("  python Table_detect.py check template/test.docx templates/Table.json")
    print("")
    print("说明:")
    print("  检查Word文档中的表格格式，包括：")
    print("  - 表格标题格式（Table X 表名，加粗，居中）")
    print("  - 表格编号连续性（Table 1, Table 2, ...）")
    print("  - 三线表格式（仅顶线、底线和表头底线）")
    print("  - 内容对齐方式（单行居中，多行左对齐）")

# ---------- CLI接口 ----------
if __name__ == '__main__':
    if len(sys.argv) != 4:
        print_help()
        sys.exit(0)
    
    cmd = sys.argv[1]
    if cmd == 'check':
        paper_path = sys.argv[2]
        tpl_id = sys.argv[3]
        
        # 检查文件是否存在
        if not os.path.isfile(paper_path):
            print(f"论文文件不存在: {paper_path}")
            sys.exit(1)
        
        try:
            # 执行检查
            report = check_doc_with_template(paper_path, tpl_id)
            
            # 打印报告
            print_report(report)
            
            # 返回状态码
            sys.exit(0 if report.get('overall_ok', False) else 1)
            
        except Exception as e:
            print("检查时出错:", e)
            import traceback
            traceback.print_exc()
            sys.exit(1)
    else:
        print_help()
        sys.exit(0)

'''
使用示例:
python paper_detect\Table_detect.py check template\test.docx Table
python paper_detect\Table_detect.py check template\test.docx templates\Table.json
'''
