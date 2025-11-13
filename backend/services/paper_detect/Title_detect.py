import os
import sys
import json
import re
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# spaCy支持（可选）
try:
    import spacy
    # 尝试加载英文模型
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except OSError:
        try:
            nlp = spacy.load("en_core_web_lg")
            SPACY_AVAILABLE = True
        except OSError:
            nlp = None
            SPACY_AVAILABLE = False
            print("警告: 未找到spaCy英文模型，将使用简化的Title Case检查")
except ImportError:
    nlp = None
    SPACY_AVAILABLE = False
    print("警告: 未安装spaCy，将使用简化的Title Case检查")


"""
=== 论文格式检测系统 - 标题、作者、单位检测器 ===

【标题、作者、单位检测 (Title, Authors, Affiliations Detection)】

1. 【标题格式检测】
   - 标题每个单词首字母大写，除了冠词、介词、并列连词
   - 使用芝加哥格式规则（基于spaCy词性分析+显式虚词列表）
   - 字号为四号(14pt)，加粗，字体为Times New Roman
   - 段前一行，段后一行
   - 支持化学式和缩写的特殊大小写保留

2. 【作者格式检测】
   - 作者姓氏全大写，名首字母大写
   - 作者英文名后可以添加括号写入中文名
   - 字号为小四(12pt)，字体为Times New Roman
   - 段前0行，段后0行
   - 支持通讯作者标记(*, †, ✉等)

3. 【单位格式检测】
   - 数字下标引用的单位编号必须存在于模板或本文中
   - 如果所有作者来自同一单位，则应该省略小标
   - 该唯一单位在单位列表中不应出现编号(例如"1.")
   - 字号为五号(10.5pt)，字体为Times New Roman
   - 段前0行，段后1行

4. 【技术特性】
   - 支持段落样式和直接格式的混合检测
   - 智能作者解析（正则表达式+启发式）
   - 单位引用完整性验证
   - spaCy智能标题大小写检查
"""
   
# ---------- 模板加载 ----------
def resolve_template_path(identifier):
    if os.path.isfile(identifier):
        return identifier
    candidate = os.path.join("templates", identifier + ".json")
    if os.path.isfile(candidate):
        return candidate
    raise FileNotFoundError(f"Template not found: '{identifier}' (tried file path and {candidate})")

def load_template(identifier):
    tpl_path = resolve_template_path(identifier)
    with open(tpl_path, 'r', encoding='utf-8') as f:
        tpl = json.load(f)
    # 严格仅读取 JSON，不注入任何默认规则
    if 'author_regex' not in tpl:
        raise ValueError("Template JSON 缺少 'author_regex' 字段。")
    return tpl


# ---------- docx 解析辅助 ----------
def get_nonempty_paragraphs(doc):
    """获取非空段落文本列表"""
    return [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

def guess_title_index(doc):
    """
    猜测标题在文档中的索引位置
    跳过章节编号（如"0 Introduction", "0 0 Introduction", "1 Introduction"等）
    """
    # 章节编号模式：数字开头，后面可能跟更多数字和空格，然后是章节名
    section_pattern = re.compile(r'^\s*\d+(\s+\d+)*\s+[A-Z]')
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # 跳过看起来像章节标题的段落
        if section_pattern.match(text):
            continue
            
        # 找到第一个不是章节编号的非空段落
        return idx
    
    # 如果所有段落都被跳过，返回第一个非空段落
    for idx, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip():
            return idx
    return 0

def split_authors_block(nonempty, start_idx):
    """
    从指定位置开始分割作者块，直到遇到单位信息或停止标记
    """
    authors_lines = []
    i = start_idx
    n = len(nonempty)
    # 停止标记
    stop_markers = re.compile(r'^\s*(Abstract|Keywords|CLC|Introduction|摘要|关键词)[\s:：]', re.IGNORECASE)
    
    while i < n:
        line = nonempty[i]
        
        # 遇到停止标记时，停止提取作者
        if stop_markers.match(line):
            break
            
        # 遇到单位编号时停止
        if re.match(r'^\s*\d+[\.\、\s:-]', line):
            break
            
        # 遇到单位关键词时停止
        if re.search(r'\b(College|University|Institute|School|Department|学院|大学|研究院)\b', line, flags=re.I):
            break
            
        authors_lines.append(line)
        i += 1
    return " ".join(authors_lines).strip(), i

def parse_authors_by_regex(authors_text, author_regex):
    """使用正则表达式解析作者信息"""
    if not authors_text:
        return []
    parts = re.split(r'[,\uFF0C;；]+', authors_text)
    authors = []
    try:
        re_compiled = re.compile(author_regex)
        use_regex = True
    except Exception:
        re_compiled = None
        use_regex = False
    paren_cn_re = re.compile(r'[\(\（]([^()\（\）]*[\u4e00-\u9fff]+[^()\（\）]*)[\)\）]')
    for p in parts:
        s = p.strip()
        if not s:
            continue
        # 检测通讯作者标记
        corresponding = any(ch in s for ch in ['', '*', '†', '✉'])
        
        # 移除中文名（括号内容）
        given_cn = None
        cn_match = paren_cn_re.search(s)
        if cn_match:
            given_cn = cn_match.group(1).strip()
            s = s[:cn_match.start()] + s[cn_match.end():]
            s = s.strip()
        
        # 移除通讯作者标记，以便正确提取单位编号
        if corresponding:
            s = re.sub(r'[☆*†✉]+', '', s).strip()
        
        # 提取单位编号（现在标记已被移除）
        m_num = re.search(r'(\d+(?:,\d+)*)\s*$', s)
        affs = []
        if m_num:
            affs = [int(x) for x in re.findall(r'\d+', m_num.group(1))]
            s = s[:m_num.start()].strip()
        surname = ""
        given_en = ""
        if use_regex and re_compiled is not None:
            m = re_compiled.search(s)
            if m:
                groups = m.groups()
                if groups:
                    if len(groups) >= 1 and groups[0]:
                        surname = groups[0].strip()
                    if len(groups) >= 2 and groups[1]:
                        given_en = groups[1].strip()
        if not surname:
            toks = s.split()
            surname = toks[0] if toks else s
            given_en = " ".join(toks[1:]) if len(toks) > 1 else ""
        authors.append({
            'raw': p.strip(),
            'surname': surname,
            'given_en': given_en,
            'given_cn': given_cn,
            'affs': affs,
            'corresponding': corresponding
        })
    return authors

def parse_affiliations_from(nonempty, start_idx):
    """
    从指定位置开始解析单位信息
    遇到Abstract、Keywords等标记时停止
    """
    affs = {}
    i = start_idx
    n = len(nonempty)
    # 停止标记
    stop_markers = re.compile(r'^\s*(Abstract|Keywords|CLC|Introduction|摘要|关键词)[\s:：]', re.IGNORECASE)
    
    while i < n:
        line = nonempty[i]
        
        # 遇到停止标记时，停止解析
        if stop_markers.match(line):
            break
            
        m = re.match(r'^\s*([0-9]+)[\.\、\s:-]+(.+)$', line)
        if not m:
            break
        affs[int(m.group(1))] = m.group(2).strip()
        i += 1
    return affs, i

def extract_from_docx(path, template):
    """从DOCX文件中提取标题、作者、单位信息"""
    doc = Document(path)
    nonempty = get_nonempty_paragraphs(doc)
    if not nonempty:
        raise ValueError("文档没有可用段落")
    t_idx = guess_title_index(doc)
    title = doc.paragraphs[t_idx].text.strip()
    try:
        ne_idx = nonempty.index(title)
    except ValueError:
        ne_idx = 0
    authors_text, next_idx = split_authors_block(nonempty, ne_idx + 1)
    authors = parse_authors_by_regex(authors_text, template.get('author_regex',''))
    affiliations, _ = parse_affiliations_from(nonempty, next_idx)
    return {
        'title': title,
        'authors_text': authors_text,
        'authors_struct': authors,
        'affiliations': affiliations
    }

def detect_font_for_run(run, paragraph=None):
    """
    更健壮的字体检测函数，同时检查直接格式和段落样式。
    返回 (font_size_pt, font_name, is_bold, is_italic, candidates_dict)
    注意：font_name是英文字体（ASCII字体）
    """
    font_size = None
    font_name_ascii = None
    is_bold = False
    is_italic = False
    
    if not run:
        return 12.0, "Times New Roman", False, False, {}
    
    # 1. 检测字号
    try:
        if run.font and run.font.size and hasattr(run.font.size, 'pt'):
            font_size = float(run.font.size.pt)
        
        if font_size is None and hasattr(run._element, 'rPr'):
            sz_nodes = run._element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
        
        if font_size is None and paragraph and paragraph.style and hasattr(paragraph.style.font, 'size') and hasattr(paragraph.style.font.size, 'pt'):
            font_size = float(paragraph.style.font.size.pt)
            
        if font_size is None and paragraph and hasattr(paragraph.style, 'element'):
            sz_nodes = paragraph.style.element.xpath('.//w:sz')
            if sz_nodes and sz_nodes[0].get(qn('w:val')):
                font_size = float(sz_nodes[0].get(qn('w:val'))) / 2.0
    except Exception:
        pass
    
    font_size = font_size if font_size is not None else 12.0
    
    # 2. 检测英文字体（ASCII）和中文字体（EastAsia）
    font_name_eastasia = None
    
    try:
        # 优先从run.font.name读取
        if run.font and run.font.name:
            font_name_ascii = run.font.name
        
        # 从XML中读取w:rFonts
        if hasattr(run._element, 'rPr'):
            rpr = run._element.rPr
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    # 读取各种字体属性
                    xml_ascii = rfonts.get(qn('w:ascii'))
                    xml_hansi = rfonts.get(qn('w:hAnsi'))
                    xml_eastasia = rfonts.get(qn('w:eastAsia'))
                    
                    # ASCII字体（英文）
                    if xml_ascii:
                        font_name_ascii = xml_ascii
                    elif xml_hansi and font_name_ascii is None:
                        font_name_ascii = xml_hansi
                    
                    # EastAsia字体（中文）
                    if xml_eastasia:
                        font_name_eastasia = xml_eastasia
                    elif xml_hansi and font_name_eastasia is None:
                        font_name_eastasia = xml_hansi
    except Exception:
        pass
    
    font_name_ascii = font_name_ascii if font_name_ascii else "Times New Roman"
    font_name_eastasia = font_name_eastasia if font_name_eastasia else "宋体"
    
    # 3. 检测加粗和斜体
    try:
        if run.font:
            is_bold = run.font.bold if run.font.bold is not None else False
            is_italic = run.font.italic if run.font.italic is not None else False
        
        is_bold = bool(is_bold)
        is_italic = bool(is_italic)
    except Exception:
        pass

    return font_size, font_name_ascii, is_bold, is_italic, {'font_eastasia': font_name_eastasia}

# ---------- 其余工具（你现有的作者/单位解析） ----------
def get_font_size(pt_size, tpl=None):
    # 从模板读取字体大小映射，提供默认值以保持向后兼容
    if tpl and 'check_rules' in tpl and 'font_size_mapping' in tpl['check_rules']:
        size_config = tpl['check_rules']['font_size_mapping']
        # 将字符串键转换为浮点数
        size_map = {}
        for key, value in size_config.items():
            try:
                size_map[float(key)] = value
            except (ValueError, TypeError):
                continue
    else:
        # 默认的字体大小映射（向后兼容）
        size_map = {
            9: "小五",
            10.5: "五号",
            12: "小四",
            14: "四号",
            16: "三号",
            18: "小二",
            22: "二号",
            24: "小一",
            26: "一号"
        }
    
    if not size_map:
        return f"{pt_size}pt"
    
    closest_size = min(size_map.keys(), key=lambda x: abs(x - pt_size))
    return size_map[closest_size]


def apply_chicago_title_case(title, nlp_model):
    """
    应用芝加哥格式的标题大小写规则，并特别处理化学式和缩写。
    规则：信任用户输入的特殊大小写（如 FePc），但修正全小写的专有名词（如 sql -> SQL）。
    """
    if not nlp_model:
        return title

    # 定义虚词的词性标签
    MINOR_POS = {"ADP", "DET", "CCONJ", "SCONJ", "PART"}
    
    # 显式定义应该小写的虚词列表（作为spaCy词性分析的补充）
    MINOR_WORDS = {
        # 冠词
        'a', 'an', 'the',
        # 常见介词（3个字母及以下，根据芝加哥格式规则）
        'at', 'by', 'for', 'in', 'of', 'on', 'to', 'up', 'as', 'or',
        # 常见连词
        'and', 'but', 'nor', 'yet', 'so',
        # 从属连词
        'if', 'as', 'or',
        # 其他常见虚词
        'is', 'be', 'am', 'are', 'was', 'were'
    }
    
    def process_word(word, is_first, is_last):
        if not word:
            return ""
        
        # 提取核心词和前后标点
        match = re.match(r'^([^\w]*)(\S+)([^\w]*)$', word)
        if not match:
            return word

        prefix, core_word, suffix = match.groups()
        
        # 规则1：如果单词已包含大写字母（非首字母），则视为特殊格式（如化学式FePc），予以保留
        if any(c.isupper() for c in core_word[1:]):
            return prefix + core_word + suffix

        # 处理连词符复合词
        if '-' in core_word and len(core_word) > 1:
            parts = core_word.split('-')
            processed_parts = [process_word(part, is_first and i == 0, is_last and i == len(parts) - 1) for i, part in enumerate(parts)]
            return prefix + '-'.join(processed_parts) + suffix
            
        # 使用spaCy进行词性分析
        doc = nlp_model(core_word.lower())
        token = doc[0] if doc else None
        
        # 规则2：首词和末词大写
        if is_first or is_last:
            # 专有名词（如SQL）应全大写
            if core_word.upper() in ['SQL', 'API', 'DNA', 'RNA']: # 可扩展的专有名词列表
                return prefix + core_word.upper() + suffix
            return prefix + core_word.capitalize() + suffix
        
        # 规则3：虚词小写（结合spaCy词性分析和显式词汇表）
        core_word_lower = core_word.lower()
        if (token and token.pos_ in MINOR_POS) or core_word_lower in MINOR_WORDS:
            return prefix + core_word.lower() + suffix
            
        # 规则4：其他实词大写
        if core_word.upper() in ['SQL', 'API', 'DNA', 'RNA']:
            return prefix + core_word.upper() + suffix
        return prefix + core_word.capitalize() + suffix

    words = title.split(' ')
    first_word_idx = next((i for i, w in enumerate(words) if w), -1)
    last_word_idx = next((len(words) - 1 - i for i, w in enumerate(reversed(words)) if w), -1)

    processed_words = [process_word(word, i == first_word_idx, i == last_word_idx) for i, word in enumerate(words)]
    return ' '.join(processed_words)

def check_chicago_title_case(title, tpl=None):
    """
    使用spaCy进行芝加哥格式的标题检查
    返回 (is_correct, bad_tokens, corrected_title)
    """
    if not SPACY_AVAILABLE or not nlp:
        # spaCy不可用，无法进行芝加哥格式检查
        return False, ["需要spaCy进行智能检查"], title
    
    try:
        corrected_title = apply_chicago_title_case(title, nlp)
        is_correct = (title.strip() == corrected_title.strip())
        
        if is_correct:
            return True, [], corrected_title
        else:
            # 分析具体哪些词有问题
            original_words = title.split()
            corrected_words = corrected_title.split()
            bad_tokens = []
            
            if len(original_words) == len(corrected_words):
                for orig, corr in zip(original_words, corrected_words):
                    if orig != corr:
                        bad_tokens.append(f"'{orig}' (应为 '{corr}')")
            
            if not bad_tokens: # Fallback if split logic fails
                bad_tokens.append("标题大小写格式不符合芝加哥风格")

            return False, bad_tokens, corrected_title
    except Exception as e:
        print(f"spaCy处理出错: {e}")
        return False, [f"spaCy处理错误: {str(e)}"], title

def check_title_section(extracted, tpl):
    """
    检查标题相关的所有内容
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    if tpl.get('rules', {}).get('title_case', False):
        use_chicago_style = tpl.get('rules', {}).get('chicago_style', False)
        
        if use_chicago_style and SPACY_AVAILABLE and nlp:
            is_correct, bad_tokens, corrected_title = check_chicago_title_case(extracted['title'], tpl)
            if not is_correct:
                report['ok'] = False
                msg_prefix = tpl.get('messages', {}).get('title_bad_prefix', '标题格式问题: ')
                report['messages'].append(f"{msg_prefix}{bad_tokens}")
                suggest_msg = tpl.get('messages', {}).get('title_suggestion')
                if suggest_msg:
                    try:
                        report['messages'].append(suggest_msg.format(corrected=corrected_title))
                    except:
                        report['messages'].append(f"建议修正为: {corrected_title}")
            else:
                ok_msg = tpl.get('messages', {}).get('title_ok')
                if ok_msg:
                    report['messages'].append(ok_msg)
        else:
            # spaCy不可用或未启用chicago_style时的提示
            report['ok'] = False
            report['messages'].append("无法进行智能标题检查：请安装spaCy并确保模板中启用了chicago_style")
            report['messages'].append("安装命令: pip install spacy && python -m spacy download en_core_web_sm")
    
    return report

def check_authors_section(extracted, tpl):
    """
    检查作者相关的所有内容
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    # 作者格式检查（完全由 JSON 规则驱动）
    author_warnings = []
    author_rules = tpl.get('author_warning_rules', [])
    for a in extracted['authors_struct']:
        for rule in author_rules:
            field = str(rule.get('field','')).strip()
            pattern = rule.get('pattern')
            must_match = bool(rule.get('must_match', True))
            when_present_only = bool(rule.get('when_present_only', True))
            message_template = rule.get('message', '')
            if not field or not pattern:
                continue
            value = a.get(field, '')
            if when_present_only and not value:
                continue
            try:
                matched = re.search(pattern, value) is not None
            except Exception:
                matched = False
            violation = (must_match and not matched) or ((not must_match) and matched)
            if violation:
                try:
                    msg = message_template.format(
                        raw=a.get('raw',''),
                        field=field,
                        value=value,
                        surname=a.get('surname',''),
                        given_en=a.get('given_en',''),
                        given_cn=a.get('given_cn','')
                    )
                except Exception:
                    msg = message_template or f"作者字段 {field} 未满足规则"
                author_warnings.append(msg)
    
    if author_warnings:
        report['ok'] = False
        report['messages'].extend(author_warnings)
    else:
        ok_msg = tpl.get('messages', {}).get('authors_pass')
        if ok_msg:
            report['messages'].append(ok_msg)
    
    return report

def check_affiliations_section(extracted, tpl):
    """
    检查单位相关的所有内容
    返回 {'ok': bool, 'messages': []}
    """
    report = {'ok': True, 'messages': []}
    
    # 单位引用检查
    used_affs = set()
    for a in extracted['authors_struct']:
        used_affs.update(a.get('affs', []))
    tpl_aff_keys = set()
    try:
        tpl_aff_keys = set(int(k) for k in tpl.get('affiliations_example', {}).keys())
    except Exception:
        tpl_aff_keys = set()
    doc_aff_keys = set(extracted.get('affiliations', {}).keys())
    valid_affs = tpl_aff_keys.union(doc_aff_keys)
    missing = [x for x in used_affs if x not in valid_affs]
    if missing:
        report['ok'] = False
        miss_prefix = tpl.get('messages', {}).get('affiliations_missing_prefix')
        if miss_prefix:
            report['messages'].append(f"{miss_prefix}{missing}")
    else:
        ok_msg = tpl.get('messages', {}).get('affiliations_ok')
        if ok_msg:
            report['messages'].append(ok_msg)

    # unique-aff detection（消息外置到 JSON）
    if used_affs:
        unique_ids = sorted(list(set(used_affs)))
        if len(unique_ids) == 1:
            single_id = unique_ids[0]
            doc_has_numbered_aff = (single_id in doc_aff_keys)
            tpl_has_numbered_aff = (single_id in tpl_aff_keys)
            if doc_has_numbered_aff or tpl_has_numbered_aff:
                report['ok'] = False
                msg1_tpl = tpl.get('messages', {}).get('affiliations_all_same')
                if msg1_tpl:
                    try:
                        msg1 = msg1_tpl.format(id=single_id)
                    except Exception:
                        msg1 = msg1_tpl
                    report['messages'].append(msg1)
                msg2_tpl = tpl.get('messages', {}).get('affiliations_all_same_advice')
                if msg2_tpl:
                    report['messages'].append(msg2_tpl)
    else:
        if len(doc_aff_keys) == 1:
            single_id = next(iter(doc_aff_keys))
            report['ok'] = False
            msg_tpl = tpl.get('messages', {}).get('affiliations_doc_single_warning')
            if msg_tpl:
                try:
                    msg = msg_tpl.format(id=single_id)
                except Exception:
                    msg = msg_tpl
                report['messages'].append(msg)
        if len(tpl_aff_keys) == 1:
            single_tpl_id = next(iter(tpl_aff_keys))
            report['ok'] = False
            msg_tpl2 = tpl.get('messages', {}).get('affiliations_tpl_single_warning')
            if msg_tpl2:
                try:
                    msg2 = msg_tpl2.format(id=single_tpl_id)
                except Exception:
                    msg2 = msg_tpl2
                report['messages'].append(msg2)
    
    return report

def check_format_section(doc_path, tpl):
    """
    检查格式相关的所有内容，并处理单位段落的特殊间距要求。
    """
    report = {'ok': True, 'messages': []}

    doc = Document(doc_path)
    nonempty_paragraphs = [p for p in doc.paragraphs if p.text and p.text.strip()]
    print("=== 开始格式检查 ===")
    
    if nonempty_paragraphs:
        # --- 标题格式检查 ---
        title_paragraph = nonempty_paragraphs[0]
        tr = tpl.get('format_rules', {}).get('title')
        if tr:
            title_format_issues = check_paragraph_format(title_paragraph,
                expected_font_size_pt=float(tr.get('font_size_pt')),
                expected_font_name=str(tr.get('font_name')),
                expected_bold=bool(tr.get('bold')),
                expected_italic=bool(tr.get('italic')),
                expected_space_before=float(tr.get('space_before')),
                expected_space_after=float(tr.get('space_after')),
                expected_alignment=tr.get('alignment'),
                tpl=tpl)
            if title_format_issues:
                report['ok'] = False
                header = tpl.get('messages', {}).get('format_title_issue_header', "标题格式问题：")
                report['messages'].append(header)
                report['messages'].extend([f"  - {i}" for i in title_format_issues])
        
        # --- 作者格式检查 ---
        if len(nonempty_paragraphs) > 1:
            author_paragraph = nonempty_paragraphs[1]
            ar = tpl.get('format_rules', {}).get('authors')
            if ar:
                author_format_issues = check_paragraph_format(author_paragraph,
                    expected_font_size_pt=float(ar.get('font_size_pt')),
                    expected_font_name=str(ar.get('font_name')),
                    expected_bold=bool(ar.get('bold')),
                    expected_italic=bool(ar.get('italic')),
                    expected_space_before=float(ar.get('space_before')),
                    expected_space_after=float(ar.get('space_after')),
                    expected_alignment=ar.get('alignment'),
                    tpl=tpl)
                if author_format_issues:
                    report['ok'] = False
                    header = tpl.get('messages', {}).get('format_authors_issue_header', "作者格式问题：")
                    report['messages'].append(header)
                    report['messages'].extend([f"  - {i}" for i in author_format_issues])
                else:
                    actual_size_pt, actual_font_name, actual_bold, actual_italic, _ = detect_font_for_run(
                        author_paragraph.runs[0] if author_paragraph.runs else None, 
                        paragraph=author_paragraph
                    )
                    report['messages'].append(f"作者格式检查通过")
                    report['messages'].append(f"  - 字体大小: {get_font_size(actual_size_pt, tpl)} 符合要求")
                    report['messages'].append(f"  - 字体名称: {actual_font_name} 符合要求")
                    report['messages'].append(f"  - 加粗设置: {'是' if actual_bold else '否'} 符合要求")
                    report['messages'].append(f"  - 斜体设置: {'是' if actual_italic else '否'} 符合要求（正体）")

        # --- 单位格式检查 (特殊逻辑) ---
        aff_detection_rules = tpl.get('check_rules', {}).get('affiliation_detection', {})
        numbered_pattern = aff_detection_rules.get('numbered_pattern', r'^\s*\d+[\.\、\s:-]')
        keywords = aff_detection_rules.get('institution_keywords', [])
        keyword_pattern = r'\b(' + '|'.join(re.escape(kw) for kw in keywords) + r')\b' if keywords else None
        
        # 章节标题模式（需要排除）
        section_pattern = re.compile(r'^\s*\d+(\s+\d+)*\s+[A-Z]')
        
        # 停止标记（遇到这些标记时停止检测单位段落）
        stop_markers = re.compile(r'^\s*(Abstract|Keywords|CLC|Introduction|摘要|关键词)[\s:：]', re.IGNORECASE)

        affiliation_paragraphs = []
        if len(nonempty_paragraphs) > 2:
            for p in nonempty_paragraphs[2:]:
                text = p.text.strip()
                
                # 遇到停止标记时，停止检测单位段落
                if stop_markers.match(text):
                    break
                
                # 排除章节标题
                if section_pattern.match(text):
                    continue
                    
                # 检查是否是单位段落
                if re.match(numbered_pattern, text) or (keyword_pattern and re.search(keyword_pattern, text, flags=re.IGNORECASE)):
                    affiliation_paragraphs.append(p)
        
        num_affs = len(affiliation_paragraphs)
        if num_affs > 0:
            afr = tpl.get('format_rules', {}).get('affiliations')
            if afr:
                for i, paragraph in enumerate(affiliation_paragraphs):
                    is_last = (i == num_affs - 1)
                    
                    expected_space_after = float(afr.get('space_after', 0)) if is_last else 0.0

                    affiliation_format_issues = check_paragraph_format(paragraph,
                        expected_font_size_pt=float(afr.get('font_size_pt')),
                        expected_font_name=str(afr.get('font_name')),
                        expected_bold=bool(afr.get('bold')),
                        expected_italic=bool(afr.get('italic')),
                        expected_space_before=float(afr.get('space_before')),
                        expected_space_after=expected_space_after,
                        tpl=tpl)
                    
                    if affiliation_format_issues:
                        report['ok'] = False
                        header_tpl = tpl.get('messages', {}).get('format_affiliation_issue_header', "单位格式问题 (段落 {index})：")
                        header = header_tpl.format(index=nonempty_paragraphs.index(paragraph) + 1)
                        report['messages'].append(header)
                        report['messages'].extend([f"  - {it}" for it in affiliation_format_issues])

    print("=== 格式检查完成 ===")
    return report

# ---------- 段落格式检查 ----------

def check_paragraph_format(paragraph, expected_font_size_pt, expected_font_name, expected_bold, expected_italic, expected_space_before, expected_space_after, expected_alignment=None, tpl=None):
    issues = []
    print(f"检查段落: '{paragraph.text[:30]}...'")
    
    main_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if not main_run:
        issues.append("段落没有可供检查的文本内容")
        return issues
        
    actual_size_pt, actual_font_name, actual_bold, actual_italic, extra_info = detect_font_for_run(main_run, paragraph)
    actual_font_eastasia = extra_info.get('font_eastasia', '宋体')

    # 字体大小
    if expected_font_size_pt is not None:
        actual_size_name = get_font_size(actual_size_pt, tpl)
        expected_size_name = get_font_size(expected_font_size_pt, tpl)
        print(f"字体大小: {actual_size_name}（{actual_size_pt}pt）(期望: {expected_size_name}（{expected_font_size_pt}pt）)")
        if abs(actual_size_pt - expected_font_size_pt) > 0.5:
            issues.append(f"字体大小应为{expected_size_name} ({expected_font_size_pt}pt)，实际为{actual_size_name} ({actual_size_pt}pt)")

    # 字体名称（英文字体）
    if expected_font_name is not None:
        print(f"英文字体: {actual_font_name}, 中文字体: {actual_font_eastasia} (期望: {expected_font_name})")
        if actual_font_name != expected_font_name:
            issues.append(f"英文字体应为{expected_font_name}，实际为{actual_font_name}")

    # 加粗
    if expected_bold is not None:
        print(f"加粗: {'是' if actual_bold else '否'} (期望: {'是' if expected_bold else '否'})")
        if actual_bold != bool(expected_bold):
            issues.append(f"字体应为{'加粗' if expected_bold else '不加粗'}，实际为{'加粗' if actual_bold else '不加粗'}")

    # 斜体
    if expected_italic is not None:
        print(f"斜体: {'是' if actual_italic else '否'} (期望: {'是' if expected_italic else '否'})")
        if actual_italic != bool(expected_italic):
            issues.append(f"字体应为{'斜体' if expected_italic else '正体'}，实际为{'斜体' if actual_italic else '正体'}")

    # 段前间距
    if expected_space_before is not None:
        actual_space_before = paragraph.paragraph_format.space_before
        actual_lines = actual_space_before.pt / 12.0 if actual_space_before and actual_space_before.pt else 0.0
        print(f"段前间距: {actual_lines:.1f}行 (期望: {expected_space_before}行)")
        # 特殊处理：1行可能显示为1.3行
        if expected_space_before == 1.0 and 1.0 <= actual_lines <= 1.35:
            pass  # 认为是正确的
        elif abs(actual_lines - expected_space_before) > 0.2:
            issues.append(f"段前间距应为{expected_space_before}行，实际为{actual_lines:.1f}行")

    # 段后间距 (增加智能容差)
    if expected_space_after is not None:
        actual_space_after = paragraph.paragraph_format.space_after
        actual_lines = actual_space_after.pt / 12.0 if actual_space_after and actual_space_after.pt else 0.0
        print(f"段后间距: {actual_lines:.1f}行 (期望: {expected_space_after}行)")
        
        # 特殊处理：1行可能显示为1.3行
        if expected_space_after == 1.0 and 1.0 <= actual_lines <= 1.35:
            pass  # 认为是正确的
        elif abs(actual_lines - expected_space_after) > 0.2:
            issues.append(f"段后间距应为{expected_space_after}行，实际为{actual_lines:.1f}行")
    
    # 对齐方式
    if expected_alignment is not None:
        # 获取段落对齐方式
        actual_alignment = paragraph.paragraph_format.alignment
        if actual_alignment is None:
            # 尝试从样式获取
            if paragraph.style and paragraph.style.paragraph_format.alignment is not None:
                actual_alignment = paragraph.style.paragraph_format.alignment
            else:
                actual_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 默认左对齐
        
        alignment_names = {
            0: '左对齐',
            1: '居中对齐',
            2: '右对齐',
            3: '两端对齐',
            WD_PARAGRAPH_ALIGNMENT.LEFT: '左对齐',
            WD_PARAGRAPH_ALIGNMENT.CENTER: '居中对齐',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: '右对齐',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: '两端对齐'
        }
        
        actual_alignment_name = alignment_names.get(actual_alignment, '未知')
        print(f"对齐方式: {actual_alignment_name} (期望: {expected_alignment})")
        
        # 判断是否符合要求
        if expected_alignment == 'justify':
            expected_val = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            if actual_alignment != expected_val and actual_alignment != 3:
                issues.append(f"对齐方式应为两端对齐，实际为{actual_alignment_name}")
        elif expected_alignment == 'center':
            expected_val = WD_PARAGRAPH_ALIGNMENT.CENTER
            if actual_alignment != expected_val and actual_alignment != 1:
                issues.append(f"对齐方式应为居中对齐，实际为{actual_alignment_name}")
        elif expected_alignment == 'left':
            expected_val = WD_PARAGRAPH_ALIGNMENT.LEFT
            if actual_alignment != expected_val and actual_alignment != 0:
                issues.append(f"对齐方式应为左对齐，实际为{actual_alignment_name}")
            
    print(f"发现 {len(issues)} 个格式问题")
    print("---")
    return issues

def check_doc_with_template(doc_path, template_identifier):
    """
    主检查函数：调用各个独立的检查函数
    """
    tpl = load_template(template_identifier)
    extracted = extract_from_docx(doc_path, tpl)

    # 调用各个独立的检查函数
    title_report = check_title_section(extracted, tpl)
    authors_report = check_authors_section(extracted, tpl)
    affiliations_report = check_affiliations_section(extracted, tpl)
    format_report = check_format_section(doc_path, tpl)

    # 组装最终报告
    report = {
        'title': title_report,
        'authors': authors_report,
        'affiliations': affiliations_report,
        'format': format_report,
        'summary': []
    }

    # 生成总结
    all_ok = (title_report['ok'] and authors_report['ok'] and 
              affiliations_report['ok'] and format_report['ok'])
    summary_tpl = tpl.get('messages', {}).get('summary_overall')
    if summary_tpl:
        try:
            report['summary'].append(summary_tpl.format(ok=all_ok))
        except Exception:
            report['summary'].append(str(summary_tpl))
    
    report['extracted'] = extracted
    return report

# ---------- 报表与 CLI ----------
def categorize_format_messages(format_messages):
    """
    将格式检查消息按类别分组
    返回 {'title': [], 'authors': [], 'affiliations': [], 'other': []}
    """
    categorized = {
        'title': [],
        'authors': [],
        'affiliations': [],
        'other': []
    }
    
    current_category = None
    
    for msg in format_messages:
        msg_lower = msg.lower()
        msg_stripped = msg.strip()
        
        # 识别分类头部消息
        if '标题格式问题' in msg or 'format_title' in msg_lower:
            current_category = 'title'
            categorized['title'].append(msg)
        elif '作者格式问题' in msg or 'format_authors' in msg_lower:
            current_category = 'authors'
            categorized['authors'].append(msg)
        elif '单位格式问题' in msg or 'format_affiliation' in msg_lower or '单位段落' in msg:
            current_category = 'affiliations'
            categorized['affiliations'].append(msg)
        # 识别通过消息
        elif '标题格式检查通过' in msg:
            current_category = 'title'
            categorized['title'].append(msg)
        elif '作者格式检查通过' in msg:
            current_category = 'authors'
            categorized['authors'].append(msg)
        elif '单位段落' in msg and '格式检查通过' in msg:
            current_category = 'affiliations'
            categorized['affiliations'].append(msg)
        # 识别详细问题消息（通常以"  - "开头）
        elif msg_stripped.startswith('- ') and current_category:
            categorized[current_category].append(msg)
        else:
            categorized['other'].append(msg)
    
    return categorized

def print_report(report):
    print("=== Check Report ===")
    
    # 分类格式消息
    format_messages = report.get('format', {}).get('messages', [])
    categorized_format = categorize_format_messages(format_messages)
    
    # 计算每个部分的综合OK状态（包含格式检查）
    # 如果有格式问题分类到某个部分，该部分就不OK
    title_format_has_issues = any('问题' in msg for msg in categorized_format['title'])
    authors_format_has_issues = any('问题' in msg for msg in categorized_format['authors'])
    affiliations_format_has_issues = any('问题' in msg for msg in categorized_format['affiliations'])
    
    title_ok = report['title']['ok'] and not title_format_has_issues
    authors_ok = report['authors']['ok'] and not authors_format_has_issues
    affiliations_ok = report['affiliations']['ok'] and not affiliations_format_has_issues
    
    # 显示标题部分（合并标题格式消息）
    print("--- TITLE ---")
    print(" OK:", title_ok)
    for m in report['title']['messages']:
            print("  -", m)
    for m in categorized_format['title']:
        print("  -", m)
    
    # 显示作者部分（合并作者格式消息）
    print("--- AUTHORS ---")
    print(" OK:", authors_ok)
    for m in report['authors']['messages']:
        print("  -", m)
    for m in categorized_format['authors']:
        print("  -", m)
    
    # 显示单位部分（合并单位格式消息）
    print("--- AFFILIATIONS ---")
    print(" OK:", affiliations_ok)
    for m in report['affiliations']['messages']:
        print("  -", m)
    for m in categorized_format['affiliations']:
        print("  -", m)
    
    # 如果有无法分类的格式消息，单独显示
    if categorized_format['other']:
        print("--- FORMAT (OTHER) ---")
        print(" OK:", report['format']['ok'])
        for m in categorized_format['other']:
            print("  -", m)
    
    print("--- SUMMARY ---")
    for s in report['summary']:
        print(" ", s)

def print_help():
    print("Usage:")
    print("  python Title_detect.py check <paper.docx> <template.json_or_name>")

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print_help()
        sys.exit(0)
    
    cmd = sys.argv[1]
    if cmd == 'check':
        paper_path = sys.argv[2]
        tpl_id = sys.argv[3]
        if not os.path.isfile(paper_path):
            print(f"论文文件不存在: {paper_path}")
            sys.exit(1)
        try:
            report = check_doc_with_template(paper_path, tpl_id)
        except Exception as e:
            print("检查时出错:", e)
            sys.exit(1)
        print_report(report)
    else:
        print_help()
        sys.exit(0)
'''
python paper_detect\Title_detect.py check template\test.docx templates\Title.json
'''