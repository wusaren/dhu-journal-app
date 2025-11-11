#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
=== 论文格式检测系统 - 集成检测脚本 ===

功能：
1. 调用所有检测模块（Title、Abstract、Keywords、Content、Formula、Table）
2. 生成综合文本报告
3. 在文档副本上添加批注标注所有问题

使用方法：
    python run_all_detections.py <docx文件路径>

输出：
    - <filename>_report.txt - 综合检测报告
    - <filename>_annotated.docx - 带批注的文档副本
"""

import os
import sys
import shutil
import re
from datetime import datetime
from docx import Document

# 模板配置映射：模块名 -> (检测函数所在模块, 检测函数名, 模板JSON路径)
TEMPLATE_MAPPING = {
    'Title': ('paper_detect.Title_detect', 'check_doc_with_template', 'templates/Title.json'),
    'Abstract': ('paper_detect.Abstract_detect', 'check_abstract_with_template', 'templates/Abstract.json'),
    'Keywords': ('paper_detect.Keywords_detect', 'check_keywords_with_template', 'templates/Keywords.json'),
    'Content': ('paper_detect.Content_detect', 'check_content_with_template', 'templates/Content.json'),
    'Formula': ('paper_detect.Formula_detect', 'check_doc_with_template', 'templates/Formula.json'),
    'Figure': ('paper_detect.Figure_detect', 'check_doc_with_template', 'templates/Figure.json'),
    'Table': ('paper_detect.Table_detect', 'check_doc_with_template', 'templates/Table.json'),
}

# 检测模块执行顺序
DETECTION_ORDER = ['Title', 'Abstract', 'Keywords', 'Content', 'Formula', 'Figure', 'Table']


def print_usage():
    """打印使用说明"""
    print(__doc__)
    print("\n使用方法：")
    print("    python run_all_detections.py <docx文件路径> [--enable-figure-api]")
    print("\n参数说明：")
    print("    --enable-figure-api    启用图片内容API检测（会调用API分析图表）")
    print("\n示例：")
    print("    python run_all_detections.py template/test.docx")
    print("    python run_all_detections.py template/test.docx --enable-figure-api")


def parse_arguments():
    """
    解析命令行参数
    返回：(docx文件路径, 是否启用Figure API检测)
    """
    if len(sys.argv) < 2:
        print("错误：参数数量不正确")
        print_usage()
        sys.exit(1)
    
    docx_path = sys.argv[1]
    enable_figure_api = False
    
    # 检查是否有 --enable-figure-api 参数
    if len(sys.argv) >= 3:
        if sys.argv[2] == '--enable-figure-api':
            enable_figure_api = True
            print("注意：已启用图片内容API检测")
        else:
            print(f"警告：未知参数 '{sys.argv[2]}'，将忽略")
    
    # 检查文件是否存在
    if not os.path.isfile(docx_path):
        print(f"错误：文件不存在: {docx_path}")
        sys.exit(1)
    
    # 检查文件扩展名
    if not docx_path.lower().endswith('.docx'):
        print(f"错误：文件必须是.docx格式: {docx_path}")
        sys.exit(1)
    
    return docx_path, enable_figure_api


def import_detection_modules():
    """
    动态导入所有检测模块
    返回：{模块名: 检测函数} 字典
    """
    detection_functions = {}
    
    for module_name, (module_path, func_name, _) in TEMPLATE_MAPPING.items():
        try:
            # 动态导入模块
            module = __import__(module_path, fromlist=[func_name])
            # 获取检测函数
            detection_func = getattr(module, func_name)
            detection_functions[module_name] = detection_func
            print(f"✓ 成功导入 {module_name} 检测模块")
        except ImportError as e:
            print(f"✗ 导入 {module_name} 模块失败: {e}")
            sys.exit(1)
        except AttributeError as e:
            print(f"✗ 在 {module_name} 模块中找不到函数 {func_name}: {e}")
            sys.exit(1)
    
    return detection_functions


def run_all_detections(docx_path, detection_functions, enable_figure_api=False):
    """
    调用所有检测模块并收集报告
    
    参数：
        docx_path: 待检测的文档路径
        detection_functions: 检测函数字典
        enable_figure_api: 是否启用Figure模块的API内容检测
    
    返回：
        {模块名: 报告字典} 的字典
    """
    all_reports = {}
    
    print(f"\n开始检测文档: {docx_path}")
    print("=" * 60)
    
    for module_name in DETECTION_ORDER:
        template_path = TEMPLATE_MAPPING[module_name][2]
        detection_func = detection_functions[module_name]
        
        print(f"\n【{module_name} 检测】")
        print(f"  使用模板: {template_path}")
        
        try:
            # 调用检测函数
            # Figure模块特殊处理：根据参数决定是否启用API内容检测
            if module_name == 'Figure':
                report = detection_func(docx_path, template_path, enable_content_check=enable_figure_api)
            else:
                report = detection_func(docx_path, template_path)
            all_reports[module_name] = report
            
            # 简要显示检测结果
            if isinstance(report, dict):
                # 统计ok状态
                ok_count = sum(1 for key, value in report.items() 
                              if isinstance(value, dict) and value.get('ok', False))
                total_count = sum(1 for key, value in report.items() 
                                 if isinstance(value, dict) and 'ok' in value)
                print(f"  结果: {ok_count}/{total_count} 项检测通过")
            else:
                print(f"  结果: 已完成")
                
        except Exception as e:
            print(f"  ✗ 检测失败: {e}")
            # 记录错误报告
            all_reports[module_name] = {
                'error': True,
                'error_message': str(e),
                'summary': [f'{module_name}检测失败: {e}']
            }
    
    print("\n" + "=" * 60)
    print("所有检测模块执行完成\n")
    
    return all_reports


def generate_comprehensive_report(all_reports):
    """
    生成综合文本报告
    
    参数：
        all_reports: {模块名: 报告字典} 的字典
    
    返回：
        格式化的报告文本字符串
    """
    lines = []
    
    # 报告头部
    lines.append("=" * 80)
    lines.append("论文格式检测综合报告")
    lines.append("=" * 80)
    lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    
    # 统计总体情况
    total_ok = 0
    total_checks = 0
    
    for module_name in DETECTION_ORDER:
        if module_name not in all_reports:
            continue
        
        report = all_reports[module_name]
        
        if report.get('error', False):
            continue
        
        # 统计该模块的检测项
        if isinstance(report, dict):
            # Table模块特殊处理
            if module_name == 'Table':
                # 统计numbering
                numbering = report.get('numbering', {})
                if isinstance(numbering, dict) and 'ok' in numbering:
                    total_checks += 1
                    if numbering.get('ok', False):
                        total_ok += 1
                
                # 统计每个表格的检测项
                tables = report.get('tables', [])
                for table_report in tables:
                    for key in ['caption_format', 'table_style', 'table_alignment']:
                        value = table_report.get(key, {})
                        if isinstance(value, dict) and 'ok' in value:
                            total_checks += 1
                            if value.get('ok', False):
                                total_ok += 1
            else:
                # 其他模块常规统计
                for key, value in report.items():
                    if isinstance(value, dict) and 'ok' in value:
                        total_checks += 1
                        if value.get('ok', False):
                            total_ok += 1
    
    # 总体评估
    lines.append("【总体评估】")
    lines.append(f"总检测项数: {total_checks}")
    lines.append(f"通过项数: {total_ok}")
    lines.append(f"失败项数: {total_checks - total_ok}")
    if total_checks > 0:
        pass_rate = (total_ok / total_checks) * 100
        lines.append(f"通过率: {pass_rate:.1f}%")
    lines.append("")
    
    # 各模块详细报告
    for module_name in DETECTION_ORDER:
        if module_name not in all_reports:
            continue
        
        report = all_reports[module_name]
        
        lines.append("-" * 80)
        lines.append(f"【{module_name} 检测报告】")
        lines.append("-" * 80)
        
        # 如果检测出错
        if report.get('error', False):
            lines.append(f"✗ 检测失败: {report.get('error_message', '未知错误')}")
            lines.append("")
            continue
        
        # Table模块需要特殊处理
        if module_name == 'Table':
            # 处理表格编号检查
            numbering = report.get('numbering', {})
            if isinstance(numbering, dict) and 'ok' in numbering:
                ok_status = "✓ 通过" if numbering.get('ok', False) else "✗ 失败"
                lines.append(f"\n  [Numbering] {ok_status}")
                messages = numbering.get('messages', [])
                if messages:
                    for msg in messages:
                        lines.append(f"    • {msg}")
            
            # 处理每个表格
            tables = report.get('tables', [])
            for i, table_report in enumerate(tables, 1):
                caption_info = table_report.get('caption', {})
                caption_text = caption_info.get('text', f'Table {i}')
                lines.append(f"\n  [表格 {i}: {caption_text[:40]}{'...' if len(caption_text) > 40 else ''}]")
                
                # 标题格式
                caption_format = table_report.get('caption_format', {})
                if isinstance(caption_format, dict) and 'ok' in caption_format:
                    ok_status = "✓" if caption_format.get('ok', False) else "✗"
                    lines.append(f"    标题格式: {ok_status}")
                    messages = caption_format.get('messages', [])
                    if messages and not caption_format.get('ok', False):
                        for msg in messages:
                            lines.append(f"      • {msg}")
                
                # 表格样式
                table_style = table_report.get('table_style', {})
                if isinstance(table_style, dict) and 'ok' in table_style:
                    ok_status = "✓" if table_style.get('ok', False) else "✗"
                    lines.append(f"    表格样式: {ok_status}")
                    messages = table_style.get('messages', [])
                    if messages and not table_style.get('ok', False):
                        for msg in messages:
                            lines.append(f"      • {msg}")
                
                # 表格对齐
                table_alignment = table_report.get('table_alignment', {})
                if isinstance(table_alignment, dict) and 'ok' in table_alignment:
                    ok_status = "✓" if table_alignment.get('ok', False) else "✗"
                    lines.append(f"    内容对齐: {ok_status}")
                    messages = table_alignment.get('messages', [])
                    if messages and not table_alignment.get('ok', False):
                        for msg in messages:
                            lines.append(f"      • {msg}")
        else:
            # 其他模块的常规处理
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and 'ok' in section_value:
                    # 检测项标题
                    section_title = section_key.replace('_', ' ').title()
                    ok_status = "✓ 通过" if section_value.get('ok', False) else "✗ 失败"
                    lines.append(f"\n  [{section_title}] {ok_status}")
                    
                    # 检测项消息
                    messages = section_value.get('messages', [])
                    if messages:
                        for msg in messages:
                            # 格式化消息（添加缩进）
                            if msg.strip().startswith('-'):
                                lines.append(f"      {msg.strip()}")
                            else:
                                lines.append(f"    • {msg}")
        
        # 添加总结
        if 'summary' in report and report['summary']:
            lines.append("\n  【总结】")
            for summary_item in report['summary']:
                lines.append(f"    {summary_item}")
        
        lines.append("")
    
    # 报告尾部
    lines.append("=" * 80)
    lines.append("报告结束")
    lines.append("=" * 80)
    
    return "\n".join(lines)


def save_report_to_file(report_text, output_path):
    """
    保存报告到文件
    
    参数：
        report_text: 报告文本
        output_path: 输出文件路径
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_text)
        print(f"✓ 报告已保存到: {output_path}")
        return True
    except Exception as e:
        print(f"✗ 保存报告失败: {e}")
        return False


def create_document_copy(original_path):
    """
    创建文档副本
    
    参数：
        original_path: 原始文档路径
    
    返回：
        副本文件路径
    """
    # 将副本文件放在与原文件相同的目录
    dir_path = os.path.dirname(original_path)
    base_name = os.path.splitext(os.path.basename(original_path))[0]
    copy_filename = f"{base_name}_annotated.docx"
    copy_path = os.path.join(dir_path, copy_filename) if dir_path else copy_filename
    
    try:
        shutil.copy(original_path, copy_path)
        print(f"✓ 文档副本已创建: {copy_path}")
        return copy_path
    except Exception as e:
        print(f"✗ 创建文档副本失败: {e}")
        return None


def find_paragraph_by_keyword(doc, keyword, case_sensitive=False):
    """
    通过关键字查找段落
    
    参数：
        doc: Document对象
        keyword: 关键字
        case_sensitive: 是否区分大小写
    
    返回：
        匹配的段落对象，未找到返回None
    """
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
        
        text = paragraph.text if case_sensitive else paragraph.text.lower()
        search_key = keyword if case_sensitive else keyword.lower()
        
        if search_key in text:
            return paragraph
    
    return None


def find_paragraph_by_index(doc, index):
    """
    通过索引查找段落
    
    参数：
        doc: Document对象
        index: 段落索引
    
    返回：
        段落对象，索引无效返回None
    """
    try:
        if 0 <= index < len(doc.paragraphs):
            return doc.paragraphs[index]
    except Exception:
        pass
    
    return None


def find_paragraph_by_text(doc, text_fragment, threshold=0.7):
    """
    通过文本片段查找段落（模糊匹配）
    
    参数：
        doc: Document对象
        text_fragment: 文本片段
        threshold: 相似度阈值（0-1）
    
    返回：
        最匹配的段落对象，未找到返回None
    """
    best_match = None
    best_score = 0
    
    search_text = text_fragment.lower().strip()
    if len(search_text) < 5:
        # 文本太短，使用精确匹配
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text.lower():
                return paragraph
        return None
    
    for paragraph in doc.paragraphs:
        if not paragraph.text or len(paragraph.text.strip()) < 5:
            continue
        
        para_text = paragraph.text.lower().strip()
        
        # 简单的相似度计算：检查搜索文本是否在段落中
        if search_text in para_text:
            score = len(search_text) / len(para_text)
            if score > best_score:
                best_score = score
                best_match = paragraph
    
    if best_score >= threshold:
        return best_match
    
    return None


def add_comment_to_paragraph(doc, paragraph, comment_text, author="论文检测系统", initials="PDS"):
    """
    为段落添加批注（覆盖整个段落）
    
    参数：
        doc: Document对象
        paragraph: 段落对象
        comment_text: 批注内容
        author: 批注作者
        initials: 批注作者缩写
    
    返回：
        成功返回True，失败返回False
    """
    try:
        # 确保段落有run
        if not paragraph.runs:
            # 如果段落没有run，添加一个
            paragraph.add_run("")
        
        # 策略：为了让批注覆盖整个段落，我们选择所有的runs
        # 如果段落有多个runs，收集它们；如果只有一个run，就用那个
        runs_to_annotate = paragraph.runs if len(paragraph.runs) > 1 else paragraph.runs[0]
        
        # 添加批注
        doc.add_comment(runs=runs_to_annotate, text=comment_text, author=author, initials=initials)
        return True
    except Exception as e:
        # 如果出错，尝试备用方案：只使用第一个run
        try:
            run = paragraph.runs[0]
            doc.add_comment(runs=run, text=comment_text, author=author, initials=initials)
            print(f"  ⚠ 批注仅覆盖部分文本（备用方案）")
            return True
        except Exception as e2:
            print(f"  ✗ 添加批注失败: {e2}")
            return False


def extract_paragraph_number_from_message(message):
    """
    从错误消息中提取段落编号
    例如："正文段落 1 字体大小..." -> 返回 1
    """
    match = re.search(r'正文段落\s*(\d+)', message)
    if match:
        return int(match.group(1))
    
    match = re.search(r'段落\s*(\d+)', message)
    if match:
        return int(match.group(1))
    
    return None


def group_messages_by_paragraph(messages):
    """
    将消息按段落分组
    
    返回：
        {段落编号: [消息列表], None: [无段落编号的消息]}
    """
    grouped = {}
    
    for msg in messages:
        para_num = extract_paragraph_number_from_message(msg)
        if para_num not in grouped:
            grouped[para_num] = []
        grouped[para_num].append(msg)
    
    return grouped


def group_messages_by_title(messages, titles):
    """
    将标题格式/大小写消息按标题文本分组
    
    参数：
        messages: 消息列表
        titles: 标题信息列表
    
    返回：
        {标题文本: [消息列表]}
    """
    import re
    grouped = {}
    
    for msg in messages:
        # 从消息中提取标题文本
        # 消息格式：标题 'xxx' 字体大小应为...
        # 或：一级标题 'xxx' 大小写不正确...
        title_match = re.search(r"标题\s+'([^']+)'", msg)
        if title_match:
            title_text = title_match.group(1)
            if title_text not in grouped:
                grouped[title_text] = []
            grouped[title_text].append(msg)
        else:
            # 消息头（如"标题格式问题："）或无法提取标题的消息，归到None
            if None not in grouped:
                grouped[None] = []
            grouped[None].append(msg)
    
    # 移除None键（这些是消息头）
    if None in grouped:
        del grouped[None]
    
    return grouped


def parse_issues_from_reports(all_reports):
    """
    从报告中提取问题列表和定位信息
    
    参数：
        all_reports: {模块名: 报告字典} 的字典
    
    返回：
        问题列表 [{
            'module': 模块名,
            'section': 检测项名称,
            'messages': 问题消息列表,
            'locate_method': 定位方法（'keyword', 'index', 'text', 'paragraph_object'），
            'locate_data': 定位数据（关键字、索引、文本片段、段落对象）
        }]
    """
    issues = []
    
    for module_name in DETECTION_ORDER:
        if module_name not in all_reports:
            continue
        
        report = all_reports[module_name]
        
        # 跳过错误报告
        if report.get('error', False):
            continue
        
        # 为不同模块设计定位策略
        if module_name == 'Title':
            # Title模块：使用关键字定位
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if not messages:
                        continue
                    
                    # 根据section类型确定关键字
                    if 'title' in section_key.lower():
                        locate_method = 'index'
                        locate_data = 0  # 标题通常是第一个段落
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': locate_method,
                            'locate_data': locate_data
                        })
                    elif 'author' in section_key.lower():
                        locate_method = 'index'
                        locate_data = 1  # 作者通常是第二个段落
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': locate_method,
                            'locate_data': locate_data
                        })
                    elif 'affiliation' in section_key.lower():
                        locate_method = 'keyword'
                        locate_data = 'College'  # 单位通常包含College关键字
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': locate_method,
                            'locate_data': locate_data
                        })
                    elif 'format' in section_key.lower():
                        # format问题需要按消息内容分组
                        # 将消息按类型分组：标题、作者、单位
                        title_msgs = []
                        author_msgs = []
                        affiliation_msgs = []
                        current_group = None  # 跟踪当前消息组
                        
                        for msg in messages:
                            msg_lower = msg.lower()
                            msg_stripped = msg.strip()
                            
                            # 判断是否是标题行（非子项）
                            is_header = not msg_stripped.startswith('- ')
                            
                            if is_header:
                                # 这是一个标题行，确定它属于哪个组
                                if '标题格式' in msg or 'title' in msg_lower:
                                    current_group = 'title'
                                    title_msgs.append(msg)
                                elif '作者格式' in msg or 'author' in msg_lower:
                                    current_group = 'author'
                                    author_msgs.append(msg)
                                elif '单位格式' in msg or '单位段落' in msg or 'affiliation' in msg_lower:
                                    current_group = 'affiliation'
                                    affiliation_msgs.append(msg)
                                else:
                                    # 默认归类到标题
                                    current_group = 'title'
                                    title_msgs.append(msg)
                            else:
                                # 这是子项，归入当前组
                                if current_group == 'title':
                                    title_msgs.append(msg)
                                elif current_group == 'author':
                                    author_msgs.append(msg)
                                elif current_group == 'affiliation':
                                    affiliation_msgs.append(msg)
                                else:
                                    # 如果没有当前组，默认归类到标题
                                    title_msgs.append(msg)
                        
                        # 为每组消息创建独立的issue
                        if title_msgs:
                            issues.append({
                                'module': module_name,
                                'section': section_key + '_title',
                                'messages': title_msgs,
                                'locate_method': 'index',
                                'locate_data': 0
                            })
                        
                        if author_msgs:
                            issues.append({
                                'module': module_name,
                                'section': section_key + '_authors',
                                'messages': author_msgs,
                                'locate_method': 'index',
                                'locate_data': 1
                            })
                        
                        if affiliation_msgs:
                            # 尝试按段落分组单位消息
                            # 消息格式如："单位格式问题（第4段）："
                            affiliation_by_para = {}
                            general_affiliation_msgs = []
                            current_para_idx = None  # 跟踪当前段落索引
                            
                            for msg in affiliation_msgs:
                                msg_stripped = msg.strip()
                                is_header = not msg_stripped.startswith('- ')
                                
                                if is_header:
                                    # 标题行，尝试提取段落编号
                                    para_match = re.search(r'第(\d+)段', msg)
                                    if para_match:
                                        current_para_idx = int(para_match.group(1)) - 1  # 转换为0-based索引
                                        if current_para_idx not in affiliation_by_para:
                                            affiliation_by_para[current_para_idx] = []
                                        affiliation_by_para[current_para_idx].append(msg)
                                    else:
                                        current_para_idx = None
                                        general_affiliation_msgs.append(msg)
                                else:
                                    # 子项，归入当前段落
                                    if current_para_idx is not None:
                                        affiliation_by_para[current_para_idx].append(msg)
                                    else:
                                        general_affiliation_msgs.append(msg)
                            
                            # 为每个单位段落创建独立的批注
                            for para_idx, msgs in affiliation_by_para.items():
                                issues.append({
                                    'module': module_name,
                                    'section': section_key + f'_affiliation_para{para_idx+1}',
                                    'messages': msgs,
                                    'locate_method': 'index',
                                    'locate_data': para_idx
                                })
                            
                            # 如果有通用的单位消息（没有段落编号），定位到第一个单位
                            if general_affiliation_msgs:
                                issues.append({
                                    'module': module_name,
                                    'section': section_key + '_affiliations',
                                    'messages': general_affiliation_msgs,
                                    'locate_method': 'keyword',
                                    'locate_data': 'College'
                                })
                    else:
                        # 其他section默认定位到第一个段落
                        locate_method = 'index'
                        locate_data = 0
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': locate_method,
                            'locate_data': locate_data
                        })
        
        elif module_name == 'Abstract':
            # Abstract模块：使用关键字定位
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if messages:
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': 'Abstract:'
                        })
        
        elif module_name == 'Keywords':
            # Keywords模块：使用关键字定位
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if messages:
                        # 区分Keywords和CLC、Footnote
                        if 'clc' in section_key.lower():
                            locate_data = 'CLC number'
                        elif 'footnote' in section_key.lower():
                            # 脚注问题：定位到CLC行（脚注引用通常在这一行）
                            # 因为Word脚注在文档底部，无法直接定位，所以定位到引用位置
                            locate_data = 'CLC number'
                        else:
                            locate_data = 'Keywords:'
                        
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': locate_data
                        })
        
        elif module_name == 'Content':
            # Content模块：需要识别具体的正文段落
            hierarchy_report = report.get('hierarchy', {})
            titles = report.get('titles', [])  # 获取标题信息用于定位
            
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted', 'details', 'hierarchy', 'titles']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if not messages:
                        continue
                    
                    # 特殊处理content_format：按段落分组
                    if section_key == 'content_format':
                        # 将消息按段落编号分组
                        grouped_messages = group_messages_by_paragraph(messages)
                        
                        # 为每个段落创建独立的批注issue
                        for para_num, para_messages in grouped_messages.items():
                            if para_num is not None:
                                # 有段落编号，需要定位到具体的正文段落
                                # 使用特殊的定位方法：content_paragraph
                                issues.append({
                                    'module': module_name,
                                    'section': f'{section_key}_para{para_num}',
                                    'messages': para_messages,
                                    'locate_method': 'content_paragraph',
                                    'locate_data': para_num,  # 段落编号（1-based）
                                    'extra': {
                                        'hierarchy_report': hierarchy_report  # 传递标题信息
                                    }
                                })
                            # 忽略没有段落编号的消息（如标题行）
                    
                    elif section_key in ['format', 'case']:
                        # 标题format和case问题：按标题分组定位
                        # 从messages中提取标题名称，定位到具体标题段落
                        grouped_title_messages = group_messages_by_title(messages, titles)
                        
                        for title_text, title_messages in grouped_title_messages.items():
                            if title_text and title_messages:
                                # 在titles列表中查找对应的标题
                                title_info = next((t for t in titles if t.get('text') == title_text), None)
                                if title_info and 'paragraph_index' in title_info:
                                    # 直接使用paragraph_index定位，更准确
                                    issues.append({
                                        'module': module_name,
                                        'section': section_key,
                                        'messages': title_messages,
                                        'locate_method': 'index',
                                        'locate_data': title_info.get('paragraph_index')
                                    })
                                elif title_text:
                                    # 如果没有索引，尝试用完整标题文本（包括编号）定位
                                    full_title_text = title_info.get('full_text', title_text) if title_info else title_text
                                    issues.append({
                                        'module': module_name,
                                        'section': section_key,
                                        'messages': title_messages,
                                        'locate_method': 'text',
                                        'locate_data': full_title_text
                                    })
                    else:
                        # 其他section（如hierarchy）使用Introduction定位
                        issues.append({
                            'module': module_name,
                            'section': section_key,
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': 'Introduction'
                        })
        
        elif module_name == 'Formula':
            # Formula模块：根据报告中的公式段落信息进行定位
            for section_key, section_value in report.items():
                if section_key in ['summary', 'extracted']:
                    continue
                
                if isinstance(section_value, dict) and not section_value.get('ok', False):
                    messages = section_value.get('messages', [])
                    if not messages:
                        continue
                    
                    # 尝试从details中获取公式段落信息
                    details = report.get('details', {})
                    formula_paragraphs = details.get('formula_paragraphs', [])
                    
                    if formula_paragraphs:
                        # 如果有公式段落信息，使用第一个公式段落的文本片段定位
                        first_formula = formula_paragraphs[0]
                        text_preview = first_formula.get('text_preview', '')
                        
                        # 从文本中提取公式编号，如 (1)、(2) 等
                        # 公式编号通常在文本末尾
                        number_match = re.search(r'\((\d+)\)\s*$', text_preview)
                        
                        if number_match:
                            # 使用公式编号定位
                            formula_number = number_match.group(0).strip()  # 如 "(2)"
                            issues.append({
                                'module': module_name,
                                'section': section_key,
                                'messages': messages,
                                'locate_method': 'formula_number',  # 新的定位方法
                                'locate_data': formula_number,
                                'extra': {
                                    'text_preview': text_preview  # 保存完整预览用于备用
                                }
                            })
                        else:
                            # 如果没有公式编号，尝试使用文本末尾的部分
                            # 只取最后的一小段（可能包含编号或特征）
                            locate_text = text_preview[-20:].strip() if len(text_preview) > 20 else text_preview
                            
                            if locate_text and len(locate_text) >= 3:
                                issues.append({
                                    'module': module_name,
                                    'section': section_key,
                                    'messages': messages,
                                    'locate_method': 'text',
                                    'locate_data': locate_text
                                })
                            else:
                                print(f"  ! Formula模块无法定位批注（文本不足）")
                    else:
                        # 如果没有找到公式段落，说明文档中可能没有公式
                        # 不添加批注，因为无法准确定位
                        print(f"  ! Formula模块未检测到公式段落，跳过批注")
        
        elif module_name == 'Table':
            # Table模块：定位到表格标题段落
            # 1. 处理numbering问题（表格编号连续性）
            numbering_report = report.get('numbering', {})
            if isinstance(numbering_report, dict) and not numbering_report.get('ok', False):
                messages = numbering_report.get('messages', [])
                if messages:
                    # 定位到第一个表格标题
                    details = report.get('details', {})
                    caption_info_list = details.get('caption_info', [])
                    if caption_info_list:
                        first_caption = caption_info_list[0]
                        caption_text = first_caption.get('text', 'Table')
                        issues.append({
                            'module': module_name,
                            'section': 'numbering',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]  # 使用标题前20个字符
                        })
            
            # 2. 处理每个表格的问题
            tables_report = report.get('tables', [])
            for i, table_report in enumerate(tables_report):
                caption_info = table_report.get('caption', {})
                caption_text = caption_info.get('text', f'Table {i+1}')
                
                # 检查标题格式
                caption_format = table_report.get('caption_format', {})
                if isinstance(caption_format, dict) and not caption_format.get('ok', False):
                    messages = caption_format.get('messages', [])
                    if messages:
                        issues.append({
                            'module': module_name,
                            'section': f'table{i+1}_caption',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]
                        })
                
                # 检查表格样式
                table_style = table_report.get('table_style', {})
                if isinstance(table_style, dict) and not table_style.get('ok', False):
                    messages = table_style.get('messages', [])
                    if messages:
                        issues.append({
                            'module': module_name,
                            'section': f'table{i+1}_style',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]
                        })
                
                # 检查表格对齐
                table_alignment = table_report.get('table_alignment', {})
                if isinstance(table_alignment, dict) and not table_alignment.get('ok', False):
                    messages = table_alignment.get('messages', [])
                    if messages:
                        issues.append({
                            'module': module_name,
                            'section': f'table{i+1}_alignment',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]
                        })
        
        elif module_name == 'Figure':
            # Figure模块：定位到图片标题段落（类似Table模块）
            # 1. 处理numbering问题（图片编号连续性）
            numbering_report = report.get('numbering', {})
            if isinstance(numbering_report, dict) and not numbering_report.get('ok', False):
                messages = numbering_report.get('messages', [])
                if messages:
                    # 定位到第一个图片标题
                    captions = report.get('captions', [])
                    if captions:
                        first_caption = captions[0]
                        caption_text = first_caption.get('full_text', 'Fig.')
                        issues.append({
                            'module': module_name,
                            'section': 'numbering',
                            'messages': messages,
                            'locate_method': 'keyword',
                            'locate_data': caption_text[:20]  # 使用标题前20个字符
                        })
            
            # 2. 处理每张图片的问题（支持新的报告结构）
            figures = report.get('figures', [])
            for fig_report in figures:
                fig_idx = fig_report.get('figure_index', 0)
                para_idx = fig_report.get('paragraph_index', 0)
                
                # 获取标题信息（如果有）
                has_caption = fig_report.get('has_caption', False)
                caption_info = fig_report.get('caption_info') if has_caption else None
                
                # 标题相关的问题 → 批注在标题段落
                caption_messages = []
                format_check = fig_report.get('format_check', {})
                if isinstance(format_check, dict) and not format_check.get('ok', False):
                    caption_messages.extend(format_check.get('messages', []))
                
                if caption_messages and has_caption:
                    # 批注在标题上
                    caption_text = caption_info.get('full_text', f'Fig.{caption_info.get("number", fig_idx)}')[:30]
                    issues.append({
                        'module': module_name,
                        'section': f'figure{fig_idx}_caption',
                        'messages': caption_messages,
                        'locate_method': 'keyword',
                        'locate_data': caption_text
                    })
                
                # 图片本身的问题 → 批注在图片段落
                picture_messages = []
                
                # 图片对齐等问题
                picture_check = fig_report.get('picture_check', {})
                if isinstance(picture_check, dict) and not picture_check.get('ok', False):
                    picture_messages.extend(picture_check.get('messages', []))
                
                # 图表内容问题
                content_check = fig_report.get('content_check', {})
                if isinstance(content_check, dict) and not content_check.get('ok', False):
                    if content_check.get('is_chart', False):
                        picture_messages.extend(content_check.get('messages', []))
                
                if picture_messages:
                    # 批注在图片段落（使用段落索引）
                    issues.append({
                        'module': module_name,
                        'section': f'figure{fig_idx}_picture',
                        'messages': picture_messages,
                        'locate_method': 'index',
                        'locate_data': para_idx
                    })
    
    return issues


def find_content_paragraph_by_number(doc, para_number, hierarchy_report):
    """
    根据段落编号定位具体的正文段落
    
    参数：
        doc: Document对象
        para_number: 段落编号（1-based，如"正文段落 1"中的1）
        hierarchy_report: 标题层级报告（包含标题段落索引信息）
    
    返回：
        对应的段落对象，未找到返回None
    """
    # 获取所有标题段落的索引
    title_paragraph_indices = set()
    introduction_index = None
    
    if hierarchy_report and hierarchy_report.get('titles'):
        for title_info in hierarchy_report['titles']:
            if 'paragraph_index' in title_info:
                idx = title_info['paragraph_index']
                title_paragraph_indices.add(idx)
                # 找到Introduction段落索引
                if title_info['level'] == 0 and title_info['text'] == 'Introduction':
                    introduction_index = idx
    
    # 如果没有找到Introduction，尝试用关键字查找
    if introduction_index is None:
        for i, para in enumerate(doc.paragraphs):
            if 'Introduction' in para.text:
                introduction_index = i
                break
    
    if introduction_index is None:
        return None
    
    # 从Introduction之后开始查找正文段落
    content_paragraph_count = 0
    
    for i in range(introduction_index + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        
        # 排除标题段落
        if i in title_paragraph_indices:
            continue
        
        # 检查是否为有效正文段落（与Content_detect.py逻辑一致）
        text = paragraph.text.strip() if paragraph.text else ""
        if not text or len(text) <= 20:
            continue
        
        # 排除表格标题、图片标题等
        text_lower = text.lower()
        if (text_lower.startswith('table ') or 
            text_lower.startswith('figure ') or 
            text_lower.startswith('fig.') or
            text_lower.startswith('图 ') or 
            text_lower.startswith('表 ') or
            re.match(r'^(表|图|Table|Figure)\s*\d+', text, re.IGNORECASE)):
            continue
        
        # 排除居中对齐的短段落（可能是标题或表格标题）
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        alignment = paragraph.alignment
        if alignment is not None and alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and len(text) < 80:
            continue
        
        # 这是有效的正文段落
        content_paragraph_count += 1
        
        # 找到目标段落
        if content_paragraph_count == para_number:
            return paragraph
    
    return None


def add_all_comments(doc_path, copy_path, issues_list):
    """
    在文档副本上添加所有批注
    
    参数：
        doc_path: 原始文档路径（用于对比）
        copy_path: 副本文档路径
        issues_list: 问题列表
    
    返回：
        成功添加的批注数量
    """
    try:
        doc = Document(copy_path)
        comment_count = 0
        
        print(f"\n正在添加批注...")
        print(f"  共有 {len(issues_list)} 个问题需要批注")
        
        for issue in issues_list:
            module_name = issue['module']
            section_name = issue['section']
            messages = issue['messages']
            locate_method = issue['locate_method']
            locate_data = issue['locate_data']
            extra = issue.get('extra', {})
            
            # 根据定位方法查找段落
            paragraph = None
            if locate_method == 'keyword' and locate_data:
                paragraph = find_paragraph_by_keyword(doc, locate_data)
            elif locate_method == 'index':
                paragraph = find_paragraph_by_index(doc, locate_data)
            elif locate_method == 'text':
                paragraph = find_paragraph_by_text(doc, locate_data)
            elif locate_method == 'paragraph_object':
                paragraph = locate_data  # 直接使用paragraph对象
            elif locate_method == 'content_paragraph':
                # 定位到具体的正文段落
                hierarchy_report = extra.get('hierarchy_report', {})
                paragraph = find_content_paragraph_by_number(doc, locate_data, hierarchy_report)
            elif locate_method == 'formula_number':
                # 通过公式编号定位（如 "(2)"）
                paragraph = find_paragraph_by_keyword(doc, locate_data)
            
            if paragraph:
                # 构建批注内容
                comment_text = f"[{module_name}-{section_name}]\n"
                for msg in messages:
                    comment_text += f"• {msg}\n"
                
                # 添加批注
                if add_comment_to_paragraph(doc, paragraph, comment_text.strip()):
                    comment_count += 1
                    print(f"  ✓ 已添加批注: {module_name}-{section_name}")
            else:
                print(f"  ✗ 无法定位段落: {module_name}-{section_name} (方法:{locate_method}, 数据:{locate_data})")
        
        # 保存文档
        doc.save(copy_path)
        print(f"\n✓ 成功添加 {comment_count} 个批注")
        print(f"✓ 批注文档已保存: {copy_path}")
        
        return comment_count
    
    except Exception as e:
        print(f"\n✗ 添加批注过程出错: {e}")
        return 0


def main():
    """主函数"""
    print("=" * 60)
    print("论文格式检测系统 - 集成检测工具")
    print("=" * 60)
    
    # 解析命令行参数
    docx_path, enable_figure_api = parse_arguments()
    print(f"\n待检测文件: {docx_path}")
    
    if not enable_figure_api:
        print("注意：图片内容API检测已禁用（使用 --enable-figure-api 启用）")
    
    # 导入检测模块
    print("\n正在加载检测模块...")
    detection_functions = import_detection_modules()
    
    # 执行所有检测
    all_reports = run_all_detections(docx_path, detection_functions, enable_figure_api)
    
    # 生成综合报告
    print("\n正在生成综合报告...")
    report_text = generate_comprehensive_report(all_reports)
    
    # 保存报告（放在与原文件相同的目录）
    dir_path = os.path.dirname(docx_path)
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    report_filename = f"{base_name}_report.txt"
    report_path = os.path.join(dir_path, report_filename) if dir_path else report_filename
    save_report_to_file(report_text, report_path)
    
    # 创建文档副本
    print("\n正在创建文档副本...")
    copy_path = create_document_copy(docx_path)
    
    if copy_path:
        # 从报告中提取问题
        print("\n正在分析问题...")
        issues_list = parse_issues_from_reports(all_reports)
        print(f"  共识别出 {len(issues_list)} 个问题")
        
        # 添加批注
        if issues_list:
            comment_count = add_all_comments(docx_path, copy_path, issues_list)
            if comment_count > 0:
                print(f"\n✓ 批注添加完成！共添加 {comment_count} 个批注")
        else:
            print("\n✓ 未发现问题，无需添加批注")
    
    print("\n" + "=" * 60)
    print("检测流程完成")
    print("=" * 60)
    print(f"\n输出文件：")
    print(f"  1. 检测报告: {report_path}")
    if copy_path:
        print(f"  2. 批注文档: {copy_path}")
    print("")


if __name__ == '__main__':
    main()

# python run_all_detections.py template/test.docx --enable-figure-api